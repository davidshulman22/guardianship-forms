#!/usr/bin/env python3
"""
Create Broward County local form .docx templates with {tags}
for use with docxtemplater in the GS Court Forms app.

Forms generated:
  BW-0010  Affidavit Regarding Criminal History
  BW-0020  Mandatory Checklist — Formal Admin Testate
  BW-0030  Mandatory Checklist — Formal Admin Intestate
  BW-0040  Mandatory Checklist — Summary Admin Testate
  BW-0050  Mandatory Checklist — Summary Admin Intestate
  BW-0060  Affidavit of Heirs
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')


def set_run_font(run, size=10, bold=False, font_name='Arial'):
    run.font.size = Pt(size)
    run.font.name = font_name
    run.bold = bold


def add_styled_paragraph(doc, text, size=10, bold=False, alignment=None,
                         space_after=None, space_before=None):
    p = doc.add_paragraph()
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    return p


# ---------------------------------------------------------------------------
# Shared helpers for checklists
# ---------------------------------------------------------------------------

def _checklist_header(doc, title):
    """Standard court header + checklist title used by all BW checklists."""
    add_styled_paragraph(doc,
        'IN THE CIRCUIT COURT OF THE SEVENTEENTH JUDICIAL CIRCUIT,',
        size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc,
        'IN AND FOR BROWARD COUNTY, FLORIDA',
        size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc,
        'PROBATE DIVISION',
        size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_styled_paragraph(doc, title,
        size=10, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)


def _checklist_instructions(doc):
    add_styled_paragraph(doc,
        'This Checklist must be completed and e-filed with your Petition. '
        'Review and sign the applicable certification clause at the end of '
        'the checklist prior to submitting it with your Petition. If any of '
        'the items below are not checked, please complete ' + '"' + 'Certification '
        'B.' + '"' + ' Completing and e-filing this Checklist does not obviate any '
        'additional obligations imposed by rule or statute.',
        size=9, bold=True, space_after=12)


def _hearing_section(doc):
    add_styled_paragraph(doc, 'HEARING:', size=9, bold=True, space_after=6)
    add_styled_paragraph(doc,
        '{hearing_exparte_check}  At the time of filing this Petition, I intend '
        'to pursue this Petition on ex-parte, motion, or special set calendar.',
        size=9, space_after=3)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{hearing_no_hearing_check}  At the time of filing this Petition, I intend '
        'to have this Petition submitted to the Judge without a hearing.',
        size=9, space_after=12)


def _case_info(doc):
    p = doc.add_paragraph()
    run = p.add_run('CASE NUMBER: PRC-{file_no}          In Re Estate of: {decedent_name}')
    set_run_font(run, size=9)
    p.paragraph_format.space_after = Pt(12)


def _certifications(doc, petition_type):
    """Add Certification A and B blocks with petition-type-specific language."""
    add_styled_paragraph(doc,
        'Please complete the Certification that applies to your filing (either '
        'Certification A or Certification B). If Petitioner is represented by '
        'counsel, only counsel must complete the applicable Certification Clause. '
        'If Petitioner is pro se then the applicable Certification must be '
        'completed by Petitioner.',
        size=9, bold=True, space_after=12)

    # Certification A
    add_styled_paragraph(doc, 'CERTIFICATION A:', size=9, bold=True, space_after=6)
    add_styled_paragraph(doc,
        'The undersigned Petitioner {petitioner_name} / Attorney {attorney_name} '
        'certifies that he/she has reviewed the information necessary to support '
        'the ' + petition_type + '. The Petitioner / Attorney further certifies '
        'that all the required information was previously filed or filed '
        'concurrently with the Petition. The Petitioner / Attorney acknowledges '
        'that the Petition will not be reviewed by Court staff until the necessary '
        'information has been accepted into the e-filing system. The Petitioner / '
        'Attorney further acknowledges that a hearing may be required to process '
        'the Petition.',
        size=9, space_after=6)

    add_styled_paragraph(doc, 'Petitioner' + "'" + 's signature: ______________________________',
        size=9, space_after=0)
    add_styled_paragraph(doc, 'Signed on: ________________, 20______',
        size=9, space_after=6)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc, 'Attorney' + "'" + 's signature: ______________________________',
        size=9, space_after=0)
    add_styled_paragraph(doc, 'Signed on: ________________, 20______',
        size=9, space_after=18)

    # Certification B
    add_styled_paragraph(doc, 'CERTIFICATION B:', size=9, bold=True, space_after=6)
    add_styled_paragraph(doc,
        'The undersigned Petitioner {petitioner_name} / Attorney {attorney_name} '
        'certifies that he/she has reviewed the information necessary to support '
        'the ' + petition_type + '. The Petitioner / Attorney certifies that, '
        'after a diligent search and reasonable effort, the Petitioner / Attorney '
        'was unable to submit the following information for the following reasons:',
        size=9, space_after=6)
    add_styled_paragraph(doc, '{cert_b_missing_items}', size=9, space_after=6)
    add_styled_paragraph(doc,
        'The Petitioner / Attorney acknowledges that a hearing may be required '
        'concerning the deficiency.',
        size=9, space_after=6)

    add_styled_paragraph(doc, 'Petitioner' + "'" + 's signature: ______________________________',
        size=9, space_after=0)
    add_styled_paragraph(doc, 'Signed on: ________________, 20______',
        size=9, space_after=6)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc, 'Attorney' + "'" + 's signature: ______________________________',
        size=9, space_after=0)
    add_styled_paragraph(doc, 'Signed on: ________________, 20______',
        size=9, space_after=0)


def _notary_block(doc, affiant_field='{affiant_name}'):
    """Standard notary block used by affidavits."""
    p = doc.add_paragraph()
    run = p.add_run('State of {notary_state}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('County of {notary_county}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(12)

    add_styled_paragraph(doc,
        'Sworn to (or affirmed) and subscribed before me by means of '
        '{notary_means} physical presence or ____ online notarization, '
        'this {notary_day} day of {notary_month}, 20{notary_year_short}, '
        'by ' + affiant_field + '.',
        size=10, space_after=18)

    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Notary Public or Deputy Clerk', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    add_styled_paragraph(doc, '{notary_personally_known_check} Personally known',
        size=10, space_after=6)
    add_styled_paragraph(doc, '{notary_produced_id_check} Produced identification',
        size=10, space_after=12)

    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Print, Type, or Stamp Commissioned', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Name of Notary Public or Deputy Clerk', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    add_styled_paragraph(doc, 'Type of identification: {notary_identification}',
        size=10, space_after=0)


def _new_doc(top=0.75, bottom=0.75, left=1, right=1, font_size=9):
    """Create a new Document with standard margins and font."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(font_size)
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)
    return doc


# ---------------------------------------------------------------------------
# BW-0010: Affidavit Regarding Criminal History
# ---------------------------------------------------------------------------

def create_criminal_history_affidavit():
    """Required for ALL Personal Representatives and Petitioners."""
    doc = _new_doc(top=1, bottom=1, left=1.25, right=1.25, font_size=10)

    # Header
    add_styled_paragraph(doc,
        'IN THE SEVENTEENTH JUDICIAL CIRCUIT, IN AND FOR',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc,
        'BROWARD COUNTY, FLORIDA',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Case caption
    p = doc.add_paragraph()
    run = p.add_run('In Re: Estate of {decedent_name}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('Case No.: {file_no}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('Division: {division}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(12)

    # Title
    add_styled_paragraph(doc,
        'AFFIDAVIT CONCERNING CRIMINAL HISTORY',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    add_styled_paragraph(doc,
        'This affidavit must be filed by all Personal Representatives and '
        'Petitioners in all testate and intestate cases, formal and summary '
        'administrations.',
        size=10, space_after=12)

    add_styled_paragraph(doc, 'I, {affiant_name}, swear or affirm:',
        size=10, space_after=12)

    # Checkbox 1 - no felony
    add_styled_paragraph(doc,
        '{no_felony_check}  I certify that I have not been convicted of a felony.',
        size=10, space_after=6)
    p = doc.add_paragraph()
    run = p.add_run('(initial)')
    set_run_font(run, size=8)
    p.paragraph_format.space_after = Pt(12)

    # Checkbox 2 - felony
    add_styled_paragraph(doc,
        '{has_felony_check}  I certify that I have been convicted of a felony. '
        'List offense, date of conviction, court and case number, state and county '
        'of the court, regardless of whether adjudication was entered or withheld.',
        size=10, space_after=12)

    add_styled_paragraph(doc, 'Offense(s): {felony_offenses}', size=10, space_after=6)
    add_styled_paragraph(doc, 'Date(s) of conviction: {felony_dates}', size=10, space_after=6)
    add_styled_paragraph(doc, 'Court & case number: {felony_court_case}', size=10, space_after=6)
    add_styled_paragraph(doc, 'State & County of the court: {felony_state_county}',
        size=10, space_after=18)

    # Perjury statement
    add_styled_paragraph(doc,
        'UNDER PENALTY OF PERJURY, I SWEAR OR AFFIRM THAT I HAVE READ THE '
        'FOREGOING AFFIDAVIT CONCERNING CRIMINAL HISTORY AND THE FACTS STATED '
        'HEREIN ARE TRUE AND COMPLETE TO THE BEST OF MY KNOWLEDGE.',
        size=10, bold=True, space_after=24)

    # Signature block
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, "Affiant" + "'" + "s signature", size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Print name and address of Affiant', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    _notary_block(doc)

    path = os.path.join(TEMPLATE_DIR, 'BW-0010.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------
# BW-0020: Checklist — Formal Administration of Testate Estate
# ---------------------------------------------------------------------------

def create_checklist_formal_testate():
    """
    Source: reference/Broward-Checklist-Formal-Admin-Testate.pdf (4 pages)
    """
    doc = _new_doc()

    _checklist_header(doc, 'CHECKLIST FOR PETITION FOR FORMAL ADMINISTRATION OF TESTATE ESTATE')
    _checklist_instructions(doc)
    _hearing_section(doc)
    _case_info(doc)

    items = [
        ('{cl_death_cert_check}',
         'A death certificate was filed.'),
        ('{cl_criminal_history_check}',
         'The Petitioner filed an Affidavit Regarding Criminal History (form '
         'available on the Seventeenth Judicial Circuit' + "'" + 's Webpage).'),
        ('{cl_residence_check}',
         'If the decedent was a Florida resident, the death certificate reflects '
         'a Broward County residence. If the decedent was not a Florida resident, '
         'the decedent owned property in Broward County, and the situs of the '
         'property is reflected in the Petition for Administration.'),
        ('{cl_petition_verified_check}',
         'The Petition is verified, signed by the Petitioner, and signed by an '
         'attorney of record.'),
        ('{cl_interest_address_check}',
         'The Petitioner' + "'" + 's interest in estate and the Petitioner' + "'"
         + 's address are listed in the Petition.'),
        ('{cl_will_filed_check}',
         'A copy of the original will or codicil was e-filed and the original '
         'will / codicil was deposited with the Broward County Clerk of Court;'),
    ]

    for tag, text in items:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{cl_lost_will_check}  The original will / codicil cannot be located, a '
        'Petition to Establish a Lost or Destroyed Will / Codicil was filed, and '
        'those who would take but for the will / codicil have consented to the '
        'Petition to Establish a Lost or Destroyed Will / Codicil.',
        size=9, space_after=12)

    # Self-proved will with OR clauses
    add_styled_paragraph(doc,
        '{cl_self_proved_check}  The decedent was a Florida resident and the '
        'will / codicil is self-proven under the laws of Florida. If the will / '
        'codicil is not self-proven, an oath of witness was executed in front of '
        'a Clerk of the Court, Commissioner, or Judge and the oath was filed with '
        'the Petition. (NOTE: a notary stamp is insufficient.)',
        size=9, space_after=6)

    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{cl_will_conformity_check}  The decedent was not a Florida resident, and '
        'an affidavit was filed demonstrating that the will / codicil was executed '
        'in conformity with the laws of the state or country where the will was '
        'executed.',
        size=9, space_after=12)

    items2 = [
        ('{cl_petitioner_qualified_check}',
         'The Petitioner is not a convicted felon and the Petitioner is a Florida '
         'resident. If the Petitioner is not a Florida Resident, the Petitioner is '
         'related to the decedent within the statutorily required degree.'),
        ('{cl_beneficiaries_check}',
         'The correct beneficiaries are listed in the Petition with the birthdates '
         'of the minor beneficiaries, if any.'),
        ('{cl_assets_check}',
         'The assets of the estate and the approximate values of the assets are '
         'listed in the Petition.'),
        ('{cl_preference_check}',
         'The proposed personal representative has preference of appointment for '
         'testate estates. If the Petitioner is not the first personal representative '
         'nominated in the will, the Petitioner has filed the necessary renunciations '
         'or death certificates that sufficiently demonstrate the proposed personal '
         'representative' + "'" + 's preference of appointment.'),
        ('{cl_oath_filed_check}',
         'An oath of personal representative and designation of resident agent '
         'were filed, and they comply with the applicable probate rules.'),
        ('{cl_order_filed_check}',
         'A proposed order admitting will to probate and appointing personal '
         'representative was filed, and the signature page contains at least four '
         '(4) lines of text and has the case number on it.'),
        ('{cl_letters_filed_check}',
         'Proposed letters of administration were filed and the signature page '
         'contains at least four (4) lines of text and has the case number on it.'),
    ]

    for tag, text in items2:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    # Trust beneficiary section
    add_styled_paragraph(doc,
        '{cl_no_trust_check}  A trust is not a beneficiary of the decedent.',
        size=9, space_after=6, bold=True)

    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{cl_trust_beneficiary_check}  If a trust of the decedent is a beneficiary '
        'of the will offered for probate: A disclosure of qualified trust '
        'beneficiaries is contained in the Petition or in a separate notice.',
        size=9, space_after=18)

    _certifications(doc, 'Petition for Formal Administration of Testate Estate')

    path = os.path.join(TEMPLATE_DIR, 'BW-0020.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------
# BW-0030: Checklist — Formal Administration of Intestate Estate
# ---------------------------------------------------------------------------

def create_checklist_formal_intestate():
    """
    Source: reference/Broward-Checklist-Formal-Admin-Intestate.pdf (3 pages)
    """
    doc = _new_doc()

    _checklist_header(doc, 'CHECKLIST FOR PETITION FOR FORMAL ADMINISTRATION OF INTESTATE ESTATE')
    _checklist_instructions(doc)
    _hearing_section(doc)
    _case_info(doc)

    items = [
        ('{cl_death_cert_check}',
         'A death certificate was filed.'),
        ('{cl_criminal_history_check}',
         'An Affidavit Concerning Criminal History was filed (form available on '
         'the Seventeenth Judicial Circuit' + "'" + 's Webpage).'),
        ('{cl_affidavit_heirs_check}',
         'An Affidavit of Heirs was filed (form available on the Seventeenth '
         'Judicial Circuit' + "'" + 's Webpage).'),
        ('{cl_petition_verified_check}',
         'The Petition is verified.'),
        ('{cl_petition_signed_interested_check}',
         'The Petition is signed by the interested person(s).'),
        ('{cl_petition_signed_attorney_check}',
         'The Petition is signed by an attorney of record.'),
        ('{cl_petitioner_relationship_check}',
         'The Petition includes the Petitioner' + "'" + 's relationship to '
         'decedent and the Petitioner' + "'" + 's residence.'),
        ('{cl_petitioner_qualified_check}',
         'The Petitioner is not a convicted felon and the Petitioner is a Florida '
         'resident. If the Petitioner is not a Florida resident, the Petitioner '
         'is related to the decedent within the statutorily required degree.'),
        ('{cl_beneficiaries_check}',
         'The correct beneficiaries are listed in the Petition with the birthdates '
         'of the minor beneficiaries, if any.'),
        ('{cl_pr_intestate_preference_check}',
         'The proposed personal representative has preference of appointment in '
         'an intestate administration.'),
        ('{cl_assets_check}',
         'The assets of the estate and the approximate value of the assets are '
         'listed in the Petition.'),
        ('{cl_oath_filed_check}',
         'An oath of personal representative and designation of resident agent '
         'were filed, and they comply with the applicable probate rules.'),
        ('{cl_order_filed_check}',
         'A proposed order appointing personal representative was filed, and the '
         'order provides space for the Court to enter a bond in its discretion.'),
        ('{cl_order_sig_page_check}',
         'The signature page of the proposed order contains at least four (4) '
         'lines of text and has the case number on it.'),
        ('{cl_letters_filed_check}',
         'Proposed letters of administration were filed and contain at least (4) '
         'lines of text on the signature page.'),
    ]

    for tag, text in items:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    _certifications(doc, 'Petition for Formal Administration of Intestate Estate')

    path = os.path.join(TEMPLATE_DIR, 'BW-0030.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------
# BW-0040: Checklist — Summary Administration of Testate Estate
# ---------------------------------------------------------------------------

def create_checklist_summary_testate():
    """
    Source: reference/Broward-Checklist-Summary-Admin-Testate.pdf (4 pages, rev 12/9/2025)
    """
    doc = _new_doc()

    _checklist_header(doc, 'CHECKLIST FOR PETITION FOR SUMMARY ADMINISTRATION OF TESTATE ESTATE')
    _checklist_instructions(doc)
    _hearing_section(doc)
    _case_info(doc)

    # Death cert
    add_styled_paragraph(doc,
        '{scl_death_cert_check}  A copy of the decedent' + "'" + 's death certificate was filed.',
        size=9, space_after=6)

    # Medical bills with OR
    add_styled_paragraph(doc,
        '{scl_medical_bills_check}  The Petitioner submitted proof of payment of the '
        'decedent' + "'" + 's reasonable and necessary medical bills from the last '
        '60 days of the decedent' + "'" + 's last illness.',
        size=9, space_after=3)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_no_medical_check}  If there are no such expenses, Petitioner has '
        'stated so in the Petition.',
        size=9, space_after=6)

    # Funeral expenses
    add_styled_paragraph(doc,
        '{scl_funeral_expenses_check}  The Petitioner submitted proof of payment of '
        'the decedent' + "'" + 's reasonable funeral expenses.',
        size=9, space_after=6)

    # Criminal history
    add_styled_paragraph(doc,
        '{scl_criminal_history_check}  The Petitioner filed an Affidavit Regarding '
        'Criminal History (form available on the Seventeenth Judicial Circuit' + "'"
        + 's Website).',
        size=9, space_after=6)

    # Beneficiary info
    add_styled_paragraph(doc,
        '{scl_beneficiary_info_check}  The Petition includes: (a) the name and '
        'address of the decedent' + "'" + 's surviving spouse (if any); (b) the '
        'names and addresses of the decedent' + "'" + 's beneficiaries and their '
        'relationship to the decedent; (c) if any beneficiary is a minor, the year '
        'of birth is included.',
        size=9, space_after=6)

    items = [
        ('{scl_venue_check}',
         'The Petition includes a statement showing venue.'),
        ('{scl_domiciliary_proceedings_check}',
         'The Petition includes a statement specifying whether there are domiciliary '
         'or principal proceedings from another state or country.'),
        ('{scl_summary_eligible_check}',
         'The Petition demonstrates the eligibility for summary administration, '
         '(i.e., the decedent died over 2 years ago or the value of the estate, '
         'less exempt property, does not exceed $75,000).'),
        ('{scl_petitioner_qualified_check}',
         'The Petitioner is a beneficiary or a person nominated as personal '
         'representative in the decedent' + "'" + 's will offered for probate.'),
        ('{scl_will_no_admin_check}',
         'The Petition includes a statement that the decedent' + "'" + 's will does '
         'not direct administration as required by chapter 733, Florida Statutes.'),
        ('{scl_assets_described_check}',
         'The Petition specifically describes the assets to be distributed, and '
         'includes values for each asset (e.g., name and address of the financial '
         'institution and the associated account number(s), legal property '
         'descriptions, etc.).'),
        ('{scl_creditor_search_check}',
         'The Petition states that a diligent search for creditors was conducted '
         'and acknowledges the penalty for failing to make a diligent search.'),
    ]
    for tag, text in items:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    # Claims section
    add_styled_paragraph(doc,
        '{scl_no_claims_check}  No claims have been filed against the estate.',
        size=9, space_after=6, bold=True)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_claims_filed_check}  If claims have been filed against the estate:',
        size=9, space_after=3, bold=True)
    add_styled_paragraph(doc,
        '        {scl_claims_barred_check}  The creditor' + "'" + 's claims are '
        'otherwise barred by statute;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_claims_paid_check}  Provision for payment has been made to '
        'the extent that assets are available;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_claims_insufficient_check}  There are insufficient assets to '
        'satisfy the outstanding claims, and formal notice of the Petition was '
        'served on the outstanding creditors.',
        size=9, space_after=12)

    # Beneficiary notice
    add_styled_paragraph(doc,
        '{scl_formal_notice_check}  All beneficiaries under the will offered for '
        'probate have received formal notice of the petition and the proposed '
        'distribution;',
        size=9, space_after=6)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_consents_filed_check}  The Petitioner has filed consents from all '
        'the beneficiaries under the will offered for probate.',
        size=9, space_after=6)

    items2 = [
        ('{scl_distribution_correct_check}',
         'The proposed order includes the correct distribution of assets as '
         'directed by the will offered for probate.'),
        ('{scl_will_admitted_check}',
         'A proposed order admitting will to probate was filed.'),
        ('{scl_order_sig_page_check}',
         'The signature page of the proposed order contains at least four (4) '
         'lines of text and has the case number on it.'),
    ]
    for tag, text in items2:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    # Trust section
    add_styled_paragraph(doc,
        '{scl_no_trust_check}  A trust is not a beneficiary of the decedent.',
        size=9, space_after=6, bold=True)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_trust_beneficiary_check}  If a trust of the decedent is a beneficiary '
        'of the will offered for probate:',
        size=9, space_after=3, bold=True)
    add_styled_paragraph(doc,
        '        {scl_trustee_is_petitioner_check}  Every trustee of the decedent' + "'"
        + 's trust is also a Petitioner for summary administration, and the '
        'Petitioners filed a disclosure of qualified trust beneficiaries and served '
        'each qualified beneficiary of the trust formal notice of the Petition.',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_trustee_not_petitioner_check}  At least one trustee of the '
        'decedent' + "'" + 's trust is not a Petitioner for summary administration.',
        size=9, space_after=12)

    # Real property section
    add_styled_paragraph(doc,
        '{scl_no_real_property_check}  The estate contains no real property of the '
        'decedent.',
        size=9, space_after=6, bold=True)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_has_real_property_check}  The estate contains real property of the '
        'decedent, and:',
        size=9, space_after=3, bold=True)
    add_styled_paragraph(doc,
        '        {scl_homestead_filed_check}  The Petitioner is claiming the '
        'decedent' + "'" + 's real property is homestead, and a Petition for '
        'Homestead has been filed;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_no_homestead_check}  The Petitioner is not claiming homestead '
        'protection for the real property.',
        size=9, space_after=18)

    _certifications(doc, 'Petition for Summary Administration of Testate Estate')

    path = os.path.join(TEMPLATE_DIR, 'BW-0040.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------
# BW-0050: Checklist — Summary Administration of Intestate Estate
# ---------------------------------------------------------------------------

def create_checklist_summary_intestate():
    """
    Source: reference/Broward-Checklist-Summary-Admin-Intestate.pdf (4 pages, rev 12/9/2025)
    """
    doc = _new_doc()

    _checklist_header(doc, 'CHECKLIST FOR PETITION FOR SUMMARY ADMINISTRATION OF INTESTATE ESTATE')
    _checklist_instructions(doc)
    _hearing_section(doc)
    _case_info(doc)

    # Death cert
    add_styled_paragraph(doc,
        '{scl_death_cert_check}  A copy of the decedent' + "'" + 's death certificate was filed.',
        size=9, space_after=6)

    # Medical bills with OR
    add_styled_paragraph(doc,
        '{scl_medical_bills_check}  The Petitioner submitted proof of payment of the '
        'decedent' + "'" + 's reasonable and necessary medical bills from the last '
        '60 days of the decedent' + "'" + 's last illness.',
        size=9, space_after=3)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_no_medical_check}  If there are no such expenses, Petitioner has '
        'stated so in the Petition.',
        size=9, space_after=6)

    # Funeral expenses
    add_styled_paragraph(doc,
        '{scl_funeral_expenses_check}  The Petitioner submitted proof of payment of '
        'the decedent' + "'" + 's reasonable funeral expenses.',
        size=9, space_after=6)

    # Criminal history
    add_styled_paragraph(doc,
        '{scl_criminal_history_check}  The Petitioner filed an Affidavit Regarding '
        'Criminal History (form available on the Seventeenth Judicial Circuit' + "'"
        + 's Webpage).',
        size=9, space_after=6)

    # Affidavit of Heirs (intestate-specific)
    add_styled_paragraph(doc,
        '{scl_affidavit_heirs_check}  The Petitioner filed an Affidavit of Heirs '
        '(form available on the Seventeenth Judicial Circuit' + "'" + 's Webpage).',
        size=9, space_after=6)

    items = [
        ('{scl_venue_check}',
         'The Petition includes a statement showing venue.'),
        ('{scl_domiciliary_proceedings_check}',
         'The Petition includes a statement specifying whether there are domiciliary '
         'or principal proceedings from another state or country.'),
        ('{scl_summary_eligible_check}',
         'The Petition demonstrates the eligibility for summary administration, '
         '(i.e., the decedent died over 2 years ago or the value of the estate, '
         'less exempt property, does not exceed $75,000).'),
        ('{scl_petitioner_is_beneficiary_check}',
         'The Petitioner is a beneficiary of the estate.'),
        ('{scl_beneficiary_info_intestate_check}',
         'The Petition includes: (a) the name and address of the decedent' + "'"
         + 's surviving spouse (if any); (b) the names and addresses of the '
         'beneficiaries and their relationship to the decedent. If any beneficiary '
         'is a minor, the year of birth is included.'),
        ('{scl_no_wills_check}',
         'The Petition states that, after a reasonably diligent search, the '
         'Petitioner is unaware of any unrevoked wills or codicils.'),
        ('{scl_assets_described_check}',
         'The Petition specifically describes the assets to be distributed, and '
         'includes values for each asset (e.g., name and address of the financial '
         'institution and the associated account number(s), legal property '
         'descriptions, etc.).'),
    ]
    for tag, text in items:
        add_styled_paragraph(doc, tag + '  ' + text, size=9, space_after=6)

    # Beneficiary notice (intestate)
    add_styled_paragraph(doc,
        '{scl_formal_notice_check}  All intestate beneficiaries have received '
        'formal notice of the Petition and the proposed distribution;',
        size=9, space_after=6)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_consents_filed_check}  The Petitioner has filed consents from all '
        'the intestate beneficiaries.',
        size=9, space_after=6)

    # Creditor search
    add_styled_paragraph(doc,
        '{scl_creditor_search_check}  The Petition states that a diligent search '
        'for creditors was conducted and acknowledges the penalty for failing to '
        'make a diligent search.',
        size=9, space_after=6)

    # Claims section
    add_styled_paragraph(doc,
        '{scl_no_claims_check}  No claims have been filed against the estate.',
        size=9, space_after=6, bold=True)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_claims_filed_check}  If claims have been filed against the estate:',
        size=9, space_after=3, bold=True)
    add_styled_paragraph(doc,
        '        {scl_claims_barred_check}  The creditor' + "'" + 's claims are '
        'otherwise barred by statute;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_claims_paid_check}  Provision for payment has been made to '
        'the extent that assets are available;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_claims_insufficient_check}  There are insufficient assets to '
        'satisfy the outstanding claims, and notice of the Petition was served on '
        'the outstanding creditors.',
        size=9, space_after=12)

    # Distribution + order
    add_styled_paragraph(doc,
        '{scl_distribution_intestacy_check}  The proposed order includes a correct '
        'distribution of assets under the laws governing intestacy.',
        size=9, space_after=6)
    add_styled_paragraph(doc,
        '{scl_order_sig_page_check}  The signature page of the proposed order '
        'contains at least four (4) lines of text and has the case number on it.',
        size=9, space_after=6)

    # Real property section (intestate uses "Petition to Determine Homestead")
    add_styled_paragraph(doc,
        '{scl_no_real_property_check}  The estate does not contain real property '
        'of the decedent.',
        size=9, space_after=6, bold=True)
    add_styled_paragraph(doc, 'OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '{scl_has_real_property_check}  The estate contains real property of the '
        'decedent, and:',
        size=9, space_after=3, bold=True)
    add_styled_paragraph(doc,
        '        {scl_homestead_filed_check}  The Petitioner is claiming the '
        'decedent' + "'" + 's real property is homestead, and a Petition to '
        'Determine Homestead has been filed;',
        size=9, space_after=3)
    add_styled_paragraph(doc, '        OR', size=9, bold=True, space_after=3)
    add_styled_paragraph(doc,
        '        {scl_no_homestead_check}  The Petitioner is not claiming homestead '
        'protection for the real property.',
        size=9, space_after=18)

    _certifications(doc, 'Petition for Summary Administration of Intestate Estate')

    path = os.path.join(TEMPLATE_DIR, 'BW-0050.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------
# BW-0060: Affidavit of Heirs
# ---------------------------------------------------------------------------

def create_affidavit_of_heirs():
    """
    Source: reference/Broward-Affidavit-of-Heirs.pdf (4 pages)
    Required for ALL intestate cases (formal and summary).
    """
    doc = _new_doc(top=1, bottom=1, left=1.25, right=1.25, font_size=10)

    # Header
    add_styled_paragraph(doc,
        'IN THE SEVENTEENTH JUDICIAL CIRCUIT, IN AND FOR',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc,
        'BROWARD COUNTY, FLORIDA',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    add_styled_paragraph(doc,
        'PROBATE DIVISION',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Case caption
    p = doc.add_paragraph()
    run = p.add_run('In Re: Estate of {decedent_name}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('Case No.: {file_no}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(0)

    p = doc.add_paragraph()
    run = p.add_run('Judge: {judge}')
    set_run_font(run, size=10)
    p.paragraph_format.space_after = Pt(12)

    # Title
    add_styled_paragraph(doc,
        'AFFIDAVIT OF HEIRS',
        size=11, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    # Instructions
    add_styled_paragraph(doc,
        'For purposes of this affidavit, you must list ALL RELATIVES of the '
        'Decedent, including yourself, if applicable. Please include even the '
        'names of relatives who were deceased at the time of the Decedent' + "'"
        + 's death, indicating that they are deceased and specifying the date of '
        'death. If the Decedent never had a relative within a particular category '
        '(i.e. the decedent was the only child, and therefore had no siblings), '
        'please indicate ' + '"' + 'None' + '"' + ' in that category. If the '
        'Decedent' + "'" + 's relatives in a particular category are unknown '
        'please specify ' + '"' + 'Unknown.' + '"' + ' When applicable, please '
        'indicate if the relationship is that of a half-relative (i.e. half-brother '
        'or half-sister).',
        size=10, space_after=12)

    # Item 1 - Undersigned
    add_styled_paragraph(doc,
        '1. The undersigned, {affiant_name},',
        size=10, space_after=6)
    add_styled_paragraph(doc,
        '     {aoh_has_interest_check}  has an interest in this estate, OR',
        size=10, space_after=3)
    add_styled_paragraph(doc,
        '     {aoh_no_interest_check}  does not have an interest in this estate.',
        size=10, space_after=6)
    add_styled_paragraph(doc,
        'I am {aoh_is_related_check} / am not {aoh_not_related_check} related to '
        'the Decedent as follows: {aoh_relationship}.',
        size=10, space_after=6)
    add_styled_paragraph(doc,
        'I have known the Decedent for {aoh_years_known} years.',
        size=10, space_after=12)

    # Items 2-9: Family information (textarea fields)
    family_items = [
        ('2a', 'Spouse of the Decedent. (Provide the name, age, and address. If the '
         'spouse is deceased, provide the name and date of death.)',
         '{aoh_spouse_info}'),
        ('2b', 'Decedent' + "'" + 's former spouse(s) (due to death or divorce). '
         '(Provide the name, age, and address. If the former spouse is deceased, '
         'provide the name and date of death. If Decedent and former spouse were '
         'divorced provide the name of former spouse and date of divorce.)',
         '{aoh_former_spouses}'),
        ('3a', 'Children of the Decedent, or descendants of deceased children. '
         '(Provide the name, age, and address. If any of the children are deceased, '
         'provide the name and date of death. In addition, indicate if Decedent has '
         'any grandchildren from the predeceased children, include the grandchild' + "'"
         + 's name here and provide further information at 4.)',
         '{aoh_children}'),
        ('3b', 'If any of the children are not biologically related to both the '
         'Decedent and Decedent' + "'" + 's spouse at the time of Decedent' + "'"
         + 's death, provide that child' + "'" + 's name here and the name of that '
         'particular child' + "'" + 's other biological parent. If the surviving '
         'spouse has children who are not the children of the Decedent provide '
         'their names.',
         '{aoh_non_biological_children}'),
        ('4', 'Grandchildren of the Decedent. (Provide the name, age, and address. '
         'If the grandchild is deceased, indicate the name and date of death.)',
         '{aoh_grandchildren}'),
        ('5', 'Parents of the Decedent. (Provide the name, age, and address. If the '
         'parents are deceased, indicate the name and date of death.)',
         '{aoh_parents}'),
        ('6', 'Brothers and sisters of the Decedent, or descendants of deceased '
         'brothers or sisters. (You must specify if the relationship is that of a '
         'half-relative, i.e., half-brother or half-sister. You must provide the '
         'name, age, and current address of the Decedent' + "'" + 's brothers or '
         'sisters or half-siblings. If any of the brothers or sisters or '
         'half-siblings are deceased, you must provide the name and date of death. '
         'In addition, you must list the children of any predeceased brothers or '
         'sisters or half-siblings, if any, along with their current addresses.)',
         '{aoh_siblings}'),
        ('7', 'Nephews and nieces of the Decedent. (Provide the name, age, and '
         'address. If the nephew or niece is deceased, indicate the name and date '
         'of death.)',
         '{aoh_nephews_nieces}'),
        ('8', 'Grandparents of the Decedent. (Provide the name, age, and address. '
         'If the grandparents are deceased, indicate the name and date of death.)',
         '{aoh_grandparents}'),
        ('9', 'If there are any relatives who have survived the Decedent and are not '
         'listed in the categories specified above, provide the name, relationship '
         'to the Decedent, age, and address (i.e., aunts, uncles, cousins, if '
         'applicable). Attach additional pages if necessary.',
         '{aoh_other_relatives}'),
    ]

    for num, desc, tag in family_items:
        add_styled_paragraph(doc, num + '. ' + desc, size=10, space_after=6)
        add_styled_paragraph(doc, tag, size=10, space_after=12)

    # Perjury statement
    add_styled_paragraph(doc,
        'UNDER PENALTIES OF PERJURY, I DECLARE THAT I HAVE READ THE FOREGOING '
        'AFFIDAVIT OF HEIRS AND THE FACTS STATED HEREIN ARE TRUE AND COMPLETE TO '
        'THE BEST OF MY KNOWLEDGE AND BELIEF.',
        size=10, bold=True, space_after=24)

    # Signature block
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Affiant' + "'" + 's Signature', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)

    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, '______________________________', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    add_styled_paragraph(doc, 'Print name and address of Affiant', size=10,
        alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=18)

    _notary_block(doc)

    path = os.path.join(TEMPLATE_DIR, 'BW-0060.docx')
    doc.save(path)
    print('Created: ' + path)
    return path


# ---------------------------------------------------------------------------

if __name__ == '__main__':
    create_criminal_history_affidavit()
    create_checklist_formal_testate()
    create_checklist_formal_intestate()
    create_checklist_summary_testate()
    create_checklist_summary_intestate()
    create_affidavit_of_heirs()
    print('\nDone. 6 Broward templates created in templates/ directory.')
