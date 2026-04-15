#!/usr/bin/env python3
"""
Update multi-petitioner FLSSI templates to use docxtemplater loop tags.

Replaces {petitioner_names} with {#petitioners}{pet_name}, {pet_address}{/petitioners}
in P2-0205, P2-0215, P2-0220, P2-0225.
"""

from docx import Document
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')

TEMPLATES = ['P2-0205.docx', 'P2-0215.docx', 'P2-0220.docx', 'P2-0225.docx']


def update_template(filename):
    path = os.path.join(TEMPLATE_DIR, filename)
    if not os.path.exists(path):
        print('SKIP (not found): ' + path)
        return False

    doc = Document(path)
    replaced = False

    for paragraph in doc.paragraphs:
        if '{petitioner_names}' in paragraph.text:
            # Replace in runs
            for run in paragraph.runs:
                if '{petitioner_names}' in run.text:
                    run.text = run.text.replace(
                        '{petitioner_names}',
                        '{#petitioners}{pet_name}, {pet_address}\n{/petitioners}'
                    )
                    replaced = True

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if '{petitioner_names}' in paragraph.text:
                        for run in paragraph.runs:
                            if '{petitioner_names}' in run.text:
                                run.text = run.text.replace(
                                    '{petitioner_names}',
                                    '{#petitioners}{pet_name}, {pet_address}\n{/petitioners}'
                                )
                                replaced = True

    if replaced:
        doc.save(path)
        print('Updated: ' + path)
    else:
        print('No {petitioner_names} tag found in: ' + filename)

    return replaced


if __name__ == '__main__':
    count = 0
    for tmpl in TEMPLATES:
        if update_template(tmpl):
            count += 1
    print('\nDone. Updated ' + str(count) + ' of ' + str(len(TEMPLATES)) + ' templates.')
