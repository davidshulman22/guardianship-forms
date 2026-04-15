#!/usr/bin/env python3
"""
tag_probate_templates.py — Tags all 19 summary administration probate forms
(P2-0204 through P2-0650) with docxtemplater placeholders.

Reads .docx files from the FLSSI source folder, inserts {field_name} tags,
and saves the tagged versions to templates/.

Run from project root:
    python3 tag_probate_templates.py

Uses python-docx (pip install python-docx).
"""

import os
import re
import shutil
from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_DIR = os.path.join(
    os.path.dirname(PROJECT_DIR),
    'FODPROBWD2025', 'Converted DOCX'
)
TEMPLATES_DIR = os.path.join(PROJECT_DIR, 'templates')

# ---------------------------------------------------------------------------
# Form definitions: source filename → template filename
# ---------------------------------------------------------------------------

FORMS = {
    'P2-0204': 'P2-0204 Petition for Summary Administration testate single petitioner.docx',
    'P2-0205': 'P2-0205 Petition for Summary Administration multiple petitioners.docx',
    'P2-0214': 'P2-0214 Petition for Summary Administration intestate single petitioner.docx',
    'P2-0215': 'P2-0215 Petition for Summary Administration intestate multiple petitioners.docx',
    'P2-0219': 'P2-0219 Petition for Summary Administration testate nonrisident decedent single petitioner.docx',
    'P2-0220': 'P2-0220 Petition for Summary Administration testate nonrisident decedent multiole petitioners.docx',
    'P2-0224': 'P2-0224 Petition for Summary Administration intestate nonrisident decedent single petitioner.docx',
    'P2-0225': 'P2-0225 Petition for Summary Administration intestate nonrisident decedent multiple petitioners.docx',
    'P2-0300': 'P2-0300 Order of Summary Administration testate.docx',
    'P2-0310': 'P2-0310 Order of Summary Administration intestate.docx',
    'P2-0320': 'P2-0320 Order of Summary Administration testate nonresident decedent.docx',
    'P2-0322': 'P2-0322 Order of Summary Administration testate nonresident decedent auth copy will.docx',
    'P2-0325': 'P2-0325Order of Summary Administration intestate nonresident decedent.docx',
    'P2-0355': 'P2-0355 Notice to Creditors summary administration.docx',
    'P2-0500': 'P2-0500 Order Admitting Will to Probate and of Summary Administration.docx',
    'P2-0600': 'P2-0600 Order Admitting Will to Probate Summary Administration.docx',
    'P2-0610': 'P2-0610 Order Admitting Will to Probate Summary Administration self proved.docx',
    'P2-0630': 'P2-0630 Order Admitting Will of Nonresident to Probate Summary Administration auth copy will.docx',
    'P2-0650': 'P2-0650 Order Admitting Will to Probate Summary Administration nonresident self proved.docx',
}

# ---------------------------------------------------------------------------
# Utility functions
# ---------------------------------------------------------------------------

def find_blank_groups(para, skip_leading=True):
    """
    Find groups of consecutive whitespace-only runs in a paragraph.
    Returns list of (start_idx, end_idx) tuples (inclusive).
    A blank group is consecutive runs where run.text contains only
    spaces, tabs, or underscores.

    If skip_leading=True, blank groups that appear before any text run
    are excluded (these are formatting indentation, not fill-in blanks).
    """
    groups = []
    in_blank = False
    start = None
    seen_text = False

    for i, run in enumerate(para.runs):
        text = run.text
        is_ws = text and not text.strip() and not text.strip('_')
        # Also treat underscore-only runs as blank
        is_underscore = text and bool(re.match(r'^[_ \t]+$', text))
        is_blank = is_ws or is_underscore

        if is_blank and not in_blank:
            in_blank = True
            start = i
        elif not is_blank and in_blank:
            in_blank = False
            # Only include blank groups that appear AFTER real text
            if not skip_leading or seen_text:
                groups.append((start, i - 1))

        # Update seen_text AFTER processing blank group boundaries
        # so leading blanks (before any text) are correctly skipped
        if not is_blank and text and text.strip():
            seen_text = True

    if in_blank:
        if not skip_leading or seen_text:
            groups.append((start, len(para.runs) - 1))

    return groups


def replace_blank_group(para, group, tag):
    """Replace a blank group (start, end) with the given tag text."""
    start, end = group
    para.runs[start].text = tag
    for i in range(start + 1, end + 1):
        para.runs[i].text = ''


def replace_nth_blank(para, n, tag):
    """Replace the nth blank group (0-indexed) in a paragraph with tag."""
    groups = find_blank_groups(para)
    if n < len(groups):
        replace_blank_group(para, groups[n], tag)
        return True
    return False


def set_run_text(para, run_idx, text):
    """Set the text of a specific run."""
    if run_idx < len(para.runs):
        para.runs[run_idx].text = text


def append_text_to_run(para, run_idx, text):
    """Append text to a specific run."""
    if run_idx < len(para.runs):
        para.runs[run_idx].text += text


def remove_paragraph(para):
    """Remove a paragraph from the document."""
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def is_empty_para(para):
    """Check if a paragraph is empty or whitespace-only."""
    return not para.text.strip()


def find_para_starting_with(doc, text, start_from=0):
    """Find paragraph index where text starts with given string."""
    for i in range(start_from, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith(text):
            return i
    return -1


def find_para_containing(doc, text, start_from=0):
    """Find paragraph index containing given text."""
    for i in range(start_from, len(doc.paragraphs)):
        if text in doc.paragraphs[i].text:
            return i
    return -1


def clear_para_and_set_text(para, text):
    """Clear all runs in a paragraph and set a single run with the text."""
    if not para.runs:
        return
    # Keep the formatting of the first run
    first_run = para.runs[0]
    first_run.text = text
    for run in para.runs[1:]:
        run.text = ''


def remove_paras_between(doc, start_idx, end_idx):
    """Remove paragraphs from start_idx to end_idx (exclusive)."""
    # Remove in reverse order to preserve indices
    to_remove = []
    for i in range(start_idx, end_idx):
        if i < len(doc.paragraphs):
            to_remove.append(doc.paragraphs[i])
    for para in to_remove:
        remove_paragraph(para)


# ---------------------------------------------------------------------------
# Case header tagging (common to all 19 forms)
# ---------------------------------------------------------------------------

def tag_case_header(doc, report):
    """Tag county, decedent_name, file_no, division in the case header."""

    # County: "IN THE CIRCUIT COURT FOR [blank] COUNTY"
    idx = find_para_containing(doc, 'CIRCUIT COURT FOR')
    if idx >= 0:
        para = doc.paragraphs[idx]
        if replace_nth_blank(para, 0, '{county}'):
            report['county'] = True

    # Decedent name: "IN RE: ESTATE OF" — append {decedent_name} to this line
    idx = find_para_containing(doc, 'IN RE:')
    if idx >= 0:
        para = doc.paragraphs[idx]
        # Find the run containing "ESTATE OF" and append the tag after it
        for run in para.runs:
            if run.text.rstrip().endswith('OF'):
                run.text = run.text.rstrip() + ' {decedent_name}'
                report['decedent_name'] = True
                break
            elif 'ESTATE OF' in run.text:
                # "IN RE:  ESTATE OF" might be in a single run, possibly
                # followed by tabs and "PROBATE DIVISION" in the same paragraph
                run.text = run.text.replace('ESTATE OF', 'ESTATE OF {decedent_name}')
                report['decedent_name'] = True
                break

    # File No: "File No." followed by tab/blank
    idx = find_para_starting_with(doc, 'File No')
    if idx >= 0:
        para = doc.paragraphs[idx]
        if replace_nth_blank(para, 0, '{file_no}'):
            report['file_no'] = True

    # Division: "Division" followed by tab/blank
    idx = find_para_starting_with(doc, 'Division')
    if idx >= 0:
        para = doc.paragraphs[idx]
        # Some have "Division " + tab
        if para.runs and len(para.runs) > 1:
            groups = find_blank_groups(para)
            if groups:
                replace_blank_group(para, groups[0], '{division}')
                report['division'] = True


# ---------------------------------------------------------------------------
# Petition tagging (P2-0204 through P2-0225)
# ---------------------------------------------------------------------------

def tag_petition(doc, form_id, report):
    """Tag a summary administration petition form."""

    is_testate = form_id in ('P2-0204', 'P2-0205', 'P2-0219', 'P2-0220')
    is_multiple = form_id in ('P2-0205', 'P2-0215', 'P2-0220', 'P2-0225')
    is_nonresident = form_id in ('P2-0219', 'P2-0220', 'P2-0224', 'P2-0225')

    # --- Petitioner name ---
    pet_word = 'Petitioners' if is_multiple else 'Petitioner'
    idx = find_para_containing(doc, pet_word + ',')
    if idx >= 0 and 'allege' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        # Find the run with underscores or the blank after petitioner
        for j, run in enumerate(para.runs):
            if '_____' in run.text or (run.text.startswith(', ') and '___' in run.text):
                run.text = re.sub(r'_+', '{petitioner_names}' if is_multiple else '{petitioner_name}', run.text)
                report['petitioner_name'] = True
                break
            elif run.text.strip().startswith(',') and j + 1 < len(para.runs):
                # Check next run for underscores
                next_run = para.runs[j + 1]
                if '___' in next_run.text:
                    next_run.text = re.sub(r'_+', '{petitioner_names}' if is_multiple else '{petitioner_name}', next_run.text)
                    report['petitioner_name'] = True
                    break

    # --- Paragraph 1: petitioner interest ---
    idx = find_para_containing(doc, 'interest in the above estate as')
    if idx >= 0:
        para = doc.paragraphs[idx]
        # The blank is after "as" - find blank groups
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{petitioner_interest}')
            report['petitioner_interest'] = True

    # --- Paragraph 2: Decedent info ---
    idx = find_para_containing(doc, 'Decedent,')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)

        # Map blank groups to fields based on surrounding text
        # Build a map: for each blank group, determine what field it is
        # by looking at the text run before it
        field_map = []
        for g_start, g_end in groups:
            # Find the text run before this blank group
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break

            if 'decedent,' in before_text or before_text.endswith('decedent,'):
                field_map.append('{decedent_full_name}')
            elif 'address was' in before_text:
                field_map.append('{decedent_address}')
            elif 'are' == before_text or before_text.endswith('are'):
                field_map.append('{decedent_ssn_last4}')
            elif 'died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',':
                # Could be death year or other comma-delimited field
                # Check if previous blank was death_date
                if field_map and field_map[-1] == '{decedent_death_date}':
                    field_map.append('{decedent_death_year}')
                else:
                    field_map.append(None)
            elif 'at' == before_text or before_text.endswith(', at'):
                field_map.append('{decedent_death_place}')
            elif 'domiciled in' in before_text:
                field_map.append('{decedent_domicile}')
            else:
                field_map.append(None)

        # Apply the field map
        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # --- Beneficiaries table ---
    idx = find_para_containing(doc, 'NAME')
    if idx >= 0 and 'ADDRESS' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#beneficiaries}{ben_name}\t{ben_address}\t{ben_relationship}\t{ben_year_of_birth}{/beneficiaries}')
        report['beneficiaries_loop'] = True

        # Remove the "[if Minor]" line if it follows
        if idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[idx + 1]
            if 'Minor' in next_para.text:
                remove_paragraph(next_para)

    # --- Venue ---
    idx = find_para_containing(doc, 'Venue of this proceeding')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{venue_reason}')
            report['venue_reason'] = True
        # Check if venue continues on next line with just "."
        if idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[idx + 1]
            if next_para.text.strip() == '.':
                # Merge with venue paragraph
                if para.runs:
                    para.runs[-1].text += '.'
                remove_paragraph(next_para)

    # --- Will info (testate only) ---
    if is_testate:
        if is_nonresident:
            # Nonresident testate: "authenticated copy of the decedent's last will, dated"
            idx = find_para_containing(doc, "last will, dated")
            if idx < 0:
                idx = find_para_containing(doc, "decedent's last will, dated")
        else:
            idx = find_para_containing(doc, "original of the decedent")
        if idx >= 0:
            para = doc.paragraphs[idx]
            groups = find_blank_groups(para)
            if len(groups) >= 1:
                replace_blank_group(para, groups[0], '{will_date}')
                report['will_date'] = True
            if len(groups) >= 2:
                replace_blank_group(para, groups[1], '{will_year}')
                report['will_year'] = True
            if not is_nonresident and len(groups) >= 3:
                replace_blank_group(para, groups[2], '{codicil_dates}')
                report['codicil_dates'] = True
            if not is_nonresident and len(groups) >= 4:
                replace_blank_group(para, groups[3], '{codicil_year}')
                report['codicil_year'] = True

    # --- Domiciliary proceedings ---
    idx = find_para_containing(doc, 'Domiciliary')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        # The blank after "Letters have been issued by" area
        # This paragraph may extend to next lines
        # Check for "the address of which is" on next line
        addr_idx = find_para_containing(doc, 'the address of which is', idx)
        if addr_idx >= 0:
            addr_para = doc.paragraphs[addr_idx]
            addr_groups = find_blank_groups(addr_para)
            if addr_groups:
                replace_blank_group(addr_para, addr_groups[0], '{domiciliary_court_address}')
                report['domiciliary_court_address'] = True

        # "to [blank], whose address is [blank]"
        to_idx = find_para_starting_with(doc, 'to ', addr_idx if addr_idx >= 0 else idx)
        if to_idx >= 0:
            to_para = doc.paragraphs[to_idx]
            to_groups = find_blank_groups(to_para)
            if len(to_groups) >= 1:
                replace_blank_group(to_para, to_groups[0], '{domiciliary_representative}')
                report['domiciliary_representative'] = True
            if len(to_groups) >= 2:
                replace_blank_group(to_para, to_groups[1], '{domiciliary_representative_address}')
                report['domiciliary_representative_address'] = True

    # --- Assets table ---
    idx = find_para_containing(doc, 'Assets')
    if idx >= 0 and 'Estimated Value' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#assets}{asset_description}\t{estimated_value}{/assets}')
        report['assets_loop'] = True

        # Remove blank filler paragraphs until next numbered paragraph
        remove_idx = idx + 1
        while remove_idx < len(doc.paragraphs):
            p = doc.paragraphs[remove_idx]
            text = p.text.strip()
            if text and not re.match(r'^[\s_]*$', text):
                break
            remove_paragraph(p)
            # Don't increment — next para shifts into this position

    # --- Distribution table ---
    idx = find_para_containing(doc, 'It is proposed that all assets')
    if idx >= 0:
        # Find the "Name ... Asset, Share or Amount" header
        dist_idx = find_para_containing(doc, 'Name', idx)
        if dist_idx >= 0 and 'Asset' in doc.paragraphs[dist_idx].text:
            para = doc.paragraphs[dist_idx]
            clear_para_and_set_text(para, '{#distribution}{dist_name}\t{dist_asset_share}{/distribution}')
            report['distribution_loop'] = True

            # Remove blank filler paragraphs
            remove_idx = dist_idx + 1
            while remove_idx < len(doc.paragraphs):
                p = doc.paragraphs[remove_idx]
                text = p.text.strip()
                if text and not re.match(r'^[\s_]*$', text):
                    break
                remove_paragraph(p)

    # --- Signing ---
    idx = find_para_containing(doc, 'Signed on this')
    if idx >= 0:
        para = doc.paragraphs[idx]
        # Replace underscores in runs — may be split across runs or combined
        full_text = para.text
        for j, run in enumerate(para.runs):
            text = run.text
            if '___' in text and 'day' in text:
                # Pattern: "this ________ day of _____________________"
                text = re.sub(r'_+ day of _+', '{signing_day} day of {signing_month}', text)
                run.text = text
                report['signing_date'] = True
                break
        if 'signing_date' not in report:
            # Underscores might be in separate runs — look for "this" run followed by underscore run
            for j, run in enumerate(para.runs):
                if '___' in run.text and j > 0:
                    prev_text = para.runs[j-1].text.lower()
                    if 'this' in prev_text:
                        run.text = re.sub(r'_+', '{signing_day}', run.text, count=1)
                        # Look for next underscore run after "day of"
                        for k in range(j+1, len(para.runs)):
                            if '___' in para.runs[k].text:
                                para.runs[k].text = re.sub(r'_+', '{signing_month}', para.runs[k].text, count=1)
                                break
                        report['signing_date'] = True
                        break
        # Year blank (spaces after comma, last blank group)
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[-1], ' {signing_year}')

    # --- Attorney block ---
    # Email addresses
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.text.strip() == '_________________________________________':
            # First one is primary email, second is secondary
            if 'email1' not in report:
                clear_para_and_set_text(para, '{attorney_email}')
                report['email1'] = True
            elif 'email2' not in report:
                clear_para_and_set_text(para, '{attorney_email_secondary}')
                report['email2'] = True

    # Florida Bar No
    idx = find_para_containing(doc, 'Florida Bar No')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{attorney_bar_no}')
            report['attorney_bar_no'] = True

    # Telephone
    idx = find_para_containing(doc, 'Telephone:')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{attorney_phone}')
            report['attorney_phone'] = True


# ---------------------------------------------------------------------------
# Order of Summary Administration tagging (P2-0300 through P2-0325)
# ---------------------------------------------------------------------------

def tag_order(doc, form_id, report):
    """Tag an Order of Summary Administration form."""

    is_testate = form_id in ('P2-0300', 'P2-0320', 'P2-0322')

    # Find the main "On the petition of" paragraph
    idx = find_para_containing(doc, 'On the petition of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)

        # Map blanks by surrounding text
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break

            if 'petition of' in before_text:
                field_map.append('{petitioner_name}')
            elif 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'died on' in before_text or 'decedent died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',':
                if field_map and field_map[-1] == '{decedent_death_date}':
                    field_map.append('{decedent_death_year}')
                elif field_map and 'will_date' in (field_map[-1] or ''):
                    field_map.append('{will_year}')
                else:
                    field_map.append(None)
            elif 'will dated' in before_text:
                field_map.append('{will_date}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Will date in adjudged section (for some forms)
    idx2 = find_para_containing(doc, 'The will dated')
    if idx2 < 0:
        idx2 = find_para_containing(doc, 'will dated', idx + 1 if idx >= 0 else 0)
    if idx2 >= 0 and idx2 != idx:
        para = doc.paragraphs[idx2]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{will_date}')
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{will_year}')
        # Witnesses
        witness_idx = find_para_containing(doc, 'attested by', idx2)
        if witness_idx < 0:
            # Might be in same paragraph
            if 'attested by' in para.text:
                for gi in range(len(groups)):
                    g_start, g_end = groups[gi]
                    before_text = ''
                    for j in range(g_start - 1, -1, -1):
                        if para.runs[j].text.strip():
                            before_text = para.runs[j].text.strip().lower()
                            break
                    if 'attested by' in before_text:
                        replace_blank_group(para, groups[gi], '{witnesses}')
                        report['witnesses'] = True

    # Distribution table
    idx = find_para_containing(doc, 'Name')
    if idx >= 0 and 'Address' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#distribution}{dist_name}\t{dist_address}\t{dist_asset_share}{/distribution}')
        report['distribution_loop'] = True

        # Remove blank filler paragraphs
        remove_idx = idx + 1
        while remove_idx < len(doc.paragraphs):
            p = doc.paragraphs[remove_idx]
            text = p.text.strip()
            if text and not re.match(r'^[\s_]*$', text):
                break
            remove_paragraph(p)

    # Order date
    idx = find_para_containing(doc, 'ORDERED on')
    if idx < 0:
        idx = find_para_containing(doc, 'Ordered on')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{order_date}')
            report['order_date'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{order_year}')
            report['order_year'] = True

    # Circuit Judge — judge name goes on the blank line above "Circuit Judge"
    idx = find_para_containing(doc, 'Circuit Judge')
    if idx >= 0:
        # Check previous paragraph for blank (judge name on separate line)
        if idx > 0 and 'judge_name' not in report:
            prev = doc.paragraphs[idx - 1]
            if is_empty_para(prev):
                clear_para_and_set_text(prev, '{judge_name}')
                report['judge_name'] = True
        # If no blank line above, check for blank within the Circuit Judge paragraph
        if 'judge_name' not in report:
            para = doc.paragraphs[idx]
            groups = find_blank_groups(para)
            if groups:
                replace_blank_group(para, groups[0], '{judge_name}')
                report['judge_name'] = True


# ---------------------------------------------------------------------------
# Notice to Creditors tagging (P2-0355)
# ---------------------------------------------------------------------------

def tag_notice(doc, report):
    """Tag the Notice to Creditors (summary administration) form."""

    # "estate of [blank], deceased"
    idx = find_para_containing(doc, 'estate of')
    if idx >= 0 and 'deceased' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)

        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break

            if 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'number' in before_text or 'file number' in before_text:
                field_map.append('{file_no}')
            elif 'address of which is' in before_text:
                field_map.append('{court_address}')
            elif 'death was' in before_text or 'date of death' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == '$' or 'value' in before_text:
                field_map.append('{estate_value}')
            elif 'county' in before_text:
                field_map.append('{county_name}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Distribution table: "Name Address"
    idx = find_para_containing(doc, 'Name')
    if idx >= 0 and 'Address' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#distribution}{dist_name}\t{dist_address}{/distribution}')
        report['distribution_loop'] = True

    # Publication date
    idx = find_para_containing(doc, 'date of first publication')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{publication_date}')
            report['publication_date'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{publication_year}')
            report['publication_year'] = True

    # Attorney info
    idx = find_para_containing(doc, 'Florida Bar No')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{attorney_bar_no}')
            report['attorney_bar_no'] = True

    # Email addresses (underscores)
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        if para.text.strip().startswith('___') and len(para.text.strip()) > 10:
            if 'email1' not in report:
                clear_para_and_set_text(para, '{attorney_email}')
                report['email1'] = True
            elif 'email2' not in report:
                clear_para_and_set_text(para, '{attorney_email_secondary}')
                report['email2'] = True


# ---------------------------------------------------------------------------
# Order Admitting Will tagging (P2-0500 through P2-0650)
# ---------------------------------------------------------------------------

def tag_will_order(doc, form_id, report):
    """Tag an Order Admitting Will to Probate form."""

    # Find main paragraph with "last will of"
    idx = find_para_containing(doc, 'last will of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)

        # These forms have fewer blanks — decedent name, death date, etc.
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break

            if 'will of' in before_text or 'last will of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'oath of' in before_text:
                field_map.append('{oath_witness}')
            elif 'died on' in before_text or 'decedent died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif 'petition of' in before_text:
                field_map.append('{petitioner_name}')
            elif 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'state of' in before_text:
                field_map.append('{execution_state}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # "On the petition of" paragraph (P2-0500, P2-0630)
    idx = find_para_containing(doc, 'On the petition of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'petition of' in before_text:
                field_map.append('{petitioner_name}')
            elif 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',':
                if field_map and field_map[-1] == '{decedent_death_date}':
                    field_map.append('{decedent_death_year}')
                elif field_map and 'will_date' in (field_map[-1] or ''):
                    field_map.append('{will_year}')
                else:
                    field_map.append(None)
            elif 'will of' in before_text:
                field_map.append('{decedent_full_name}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # ADJUDGED paragraph with will date and witnesses
    idx = find_para_containing(doc, 'will dated')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'will dated' in before_text or 'dated' in before_text:
                field_map.append('{will_date}')
            elif before_text == ',':
                if field_map and field_map[-1] == '{will_date}':
                    field_map.append('{will_year}')
                else:
                    field_map.append(None)
            elif 'attested by' in before_text:
                field_map.append('{witnesses}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Distribution table (P2-0500 only)
    if form_id == 'P2-0500':
        idx = find_para_containing(doc, 'Name')
        if idx >= 0 and 'Address' in doc.paragraphs[idx].text:
            para = doc.paragraphs[idx]
            clear_para_and_set_text(para, '{#distribution}{dist_name}\t{dist_address}\t{dist_asset_share}{/distribution}')
            report['distribution_loop'] = True

            remove_idx = idx + 1
            while remove_idx < len(doc.paragraphs):
                p = doc.paragraphs[remove_idx]
                text = p.text.strip()
                if text and not re.match(r'^[\s_]*$', text):
                    break
                remove_paragraph(p)

    # Order date
    idx = find_para_containing(doc, 'ORDERED on')
    if idx < 0:
        idx = find_para_containing(doc, 'Ordered on')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{order_date}')
            report['order_date'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{order_year}')
            report['order_year'] = True

    # Circuit Judge — judge name on blank line above
    idx = find_para_containing(doc, 'Circuit Judge')
    if idx >= 0 and 'judge_name' not in report:
        if idx > 0:
            prev = doc.paragraphs[idx - 1]
            if is_empty_para(prev):
                clear_para_and_set_text(prev, '{judge_name}')
                report['judge_name'] = True
        if 'judge_name' not in report:
            para = doc.paragraphs[idx]
            groups = find_blank_groups(para)
            if groups:
                replace_blank_group(para, groups[0], '{judge_name}')
                report['judge_name'] = True


# ---------------------------------------------------------------------------
# Main processing
# ---------------------------------------------------------------------------

def process_form(form_id, source_filename):
    """Process a single form: copy, tag, save."""
    source_path = os.path.join(SOURCE_DIR, source_filename)
    dest_path = os.path.join(TEMPLATES_DIR, form_id + '.docx')

    if not os.path.exists(source_path):
        print(f'  ERROR: Source file not found: {source_filename}')
        return None

    # Copy source to templates
    shutil.copy2(source_path, dest_path)

    # Open and tag
    doc = Document(dest_path)
    report = {}

    # Tag case header (all forms)
    tag_case_header(doc, report)

    # Tag form-specific fields
    if form_id.startswith('P2-02'):
        # Petitions
        tag_petition(doc, form_id, report)
    elif form_id in ('P2-0300', 'P2-0310', 'P2-0320', 'P2-0322', 'P2-0325'):
        # Orders of Summary Administration
        tag_order(doc, form_id, report)
    elif form_id == 'P2-0355':
        # Notice to Creditors
        tag_notice(doc, report)
    elif form_id.startswith('P2-0'):
        # Order Admitting Will
        tag_will_order(doc, form_id, report)

    # Save
    doc.save(dest_path)
    return report


def main():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    print('=' * 70)
    print('Tagging 19 Summary Administration Probate Templates')
    print('=' * 70)
    print(f'Source: {SOURCE_DIR}')
    print(f'Output: {TEMPLATES_DIR}')
    print()

    for form_id, filename in sorted(FORMS.items()):
        print(f'--- {form_id}: {filename[:60]} ---')
        report = process_form(form_id, filename)
        if report:
            tagged = [k for k, v in report.items() if v]
            print(f'  Tagged: {", ".join(tagged)}')
        else:
            print(f'  FAILED')
        print()

    print('Done! Tagged templates saved to templates/')


if __name__ == '__main__':
    main()
