#!/usr/bin/env python3
"""
tag_formal_admin_templates.py — Tags 10 core formal administration probate forms
with docxtemplater placeholders.

Reads .docx files from the FLSSI source folder, inserts {field_name} tags,
and saves the tagged versions to templates/.

Run from project root:
    python3 tag_formal_admin_templates.py

Uses python-docx (pip install python-docx).
"""

import os
import re
import shutil
from docx import Document
from docx.oxml.ns import qn

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------

PROJECT_DIR = os.path.dirname(os.path.abspath(__file__))
SOURCE_DIR = os.path.join(
    os.path.dirname(PROJECT_DIR),
    'FODPROBWD2025', 'Converted DOCX'
)
TEMPLATES_DIR = os.path.join(PROJECT_DIR, 'templates')

# ---------------------------------------------------------------------------
# Form definitions
# ---------------------------------------------------------------------------

FORMS = {
    'P3-0100': 'P3-0100 Petition for Administration testate FL resident single petitioner.docx',
    'P3-0120': 'P3-0120 Petition for Administration intestate FL resident single petitioner.docx',
    'P3-0420': 'P3-0420 Order Admitting Will to Probate and Appointing Personal Representative self proved single.docx',
    'P3-0440': 'P3-0440 Order Appointing Personal Representative intestate single.docx',
    'P3-0600': 'P3-0600 Oat of Personal Representative and Designation of Resident Agent.docx',
    'P3-0700': 'P3-0700 Letters of Administration single personal representative.docx',
    'P3-0740': 'P3-0740 Notice to Creditors.docx',
    'P1-0900': 'P1-0900 Notice of Designation of Email Addresses for Service of Documents.docx',
    'P3-0900': 'P3-0900 Inventory.docx',
    'P5-0400': 'P5-0400 Petition for Discharge single personal representative.docx',
    'P5-0800': 'P5-0800 Order of Discharge single PR.docx',
}

# ---------------------------------------------------------------------------
# Utility functions (same as tag_probate_templates.py)
# ---------------------------------------------------------------------------

def find_blank_groups(para, skip_leading=True):
    groups = []
    in_blank = False
    start = None
    seen_text = False

    for i, run in enumerate(para.runs):
        text = run.text
        is_ws = text and not text.strip() and not text.strip('_')
        is_underscore = text and bool(re.match(r'^[_ \t]+$', text))
        is_blank = is_ws or is_underscore

        if is_blank and not in_blank:
            in_blank = True
            start = i
        elif not is_blank and in_blank:
            in_blank = False
            if not skip_leading or seen_text:
                groups.append((start, i - 1))

        if not is_blank and text and text.strip():
            seen_text = True

    if in_blank:
        if not skip_leading or seen_text:
            groups.append((start, len(para.runs) - 1))

    return groups


def replace_blank_group(para, group, tag):
    start, end = group
    para.runs[start].text = tag
    for i in range(start + 1, end + 1):
        para.runs[i].text = ''


def replace_nth_blank(para, n, tag):
    groups = find_blank_groups(para)
    if n < len(groups):
        replace_blank_group(para, groups[n], tag)
        return True
    return False


def clear_para_and_set_text(para, text):
    if not para.runs:
        return
    para.runs[0].text = text
    for run in para.runs[1:]:
        run.text = ''


def remove_paragraph(para):
    p = para._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def is_empty_para(para):
    return not para.text.strip()


def find_para_containing(doc, text, start_from=0):
    for i in range(start_from, len(doc.paragraphs)):
        if text in doc.paragraphs[i].text:
            return i
    return -1


def find_para_starting_with(doc, text, start_from=0):
    for i in range(start_from, len(doc.paragraphs)):
        if doc.paragraphs[i].text.strip().startswith(text):
            return i
    return -1


def replace_blanks_by_context(para, report):
    """Replace blank groups based on surrounding text context."""
    groups = find_blank_groups(para)
    for g_start, g_end in groups:
        before_text = ''
        for j in range(g_start - 1, -1, -1):
            if para.runs[j].text.strip():
                before_text = para.runs[j].text.strip().lower()
                break

        tag = None
        if 'petition of' in before_text or 'petitioner,' in before_text:
            tag = '{petitioner_name}'
        elif 'estate of' in before_text and 'deceased' not in before_text:
            tag = '{decedent_full_name}'
        elif 'died on' in before_text or 'decedent died on' in before_text:
            tag = '{decedent_death_date}'
        elif 'will dated' in before_text or 'dated' == before_text:
            tag = '{will_date}'
        elif 'attested by' in before_text:
            tag = '{witnesses}'
        elif 'address was' in before_text:
            tag = '{decedent_address}'
        elif before_text == 'are' or before_text.endswith('are'):
            tag = '{decedent_ssn_last4}'
        elif 'domiciled in' in before_text:
            tag = '{decedent_domicile}'
        elif 'address is' in before_text and 'whose address is' in before_text:
            tag = '{pr_address}'
        elif 'bond in the sum of' in before_text or 'sum of' in before_text:
            tag = '{bond_amount}'
        elif 'reason of' in before_text:
            tag = '{pr_entitlement_reason}'

        if tag:
            replace_blank_group(para, (g_start, g_end), tag)
            report[tag.strip('{}')] = True


# ---------------------------------------------------------------------------
# Case header tagging (shared)
# ---------------------------------------------------------------------------

def tag_case_header(doc, report):
    idx = find_para_containing(doc, 'CIRCUIT COURT FOR')
    if idx >= 0:
        para = doc.paragraphs[idx]
        if replace_nth_blank(para, 0, '{county}'):
            report['county'] = True

    idx = find_para_containing(doc, 'IN RE:')
    if idx >= 0:
        para = doc.paragraphs[idx]
        for run in para.runs:
            if run.text.rstrip().endswith('OF'):
                run.text = run.text.rstrip() + ' {decedent_name}'
                report['decedent_name'] = True
                break
            elif 'ESTATE OF' in run.text:
                run.text = run.text.replace('ESTATE OF', 'ESTATE OF {decedent_name}')
                report['decedent_name'] = True
                break

    idx = find_para_starting_with(doc, 'File No')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{file_no}')
            report['file_no'] = True

    idx = find_para_starting_with(doc, 'Division')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{division}')
            report['division'] = True


def tag_signing_block(doc, report):
    idx = find_para_containing(doc, 'Signed on this')
    if idx >= 0:
        para = doc.paragraphs[idx]
        for j, run in enumerate(para.runs):
            text = run.text
            if '___' in text and 'day' in text:
                text = re.sub(r'_+ day of _+', '{signing_day} day of {signing_month}', text)
                run.text = text
                report['signing_date'] = True
                break
        if 'signing_date' not in report:
            for j, run in enumerate(para.runs):
                if '___' in run.text and j > 0:
                    prev_text = para.runs[j-1].text.lower()
                    if 'this' in prev_text:
                        run.text = re.sub(r'_+', '{signing_day}', run.text, count=1)
                        for k in range(j+1, len(para.runs)):
                            if '___' in para.runs[k].text:
                                para.runs[k].text = re.sub(r'_+', '{signing_month}', para.runs[k].text, count=1)
                                break
                        report['signing_date'] = True
                        break
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[-1], ' {signing_year}')


def tag_attorney_block(doc, report):
    for i in range(len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        if text.startswith('___') and len(text) > 10:
            if 'email1' not in report:
                clear_para_and_set_text(para, '{attorney_email}')
                report['email1'] = True
            elif 'email2' not in report:
                clear_para_and_set_text(para, '{attorney_email_secondary}')
                report['email2'] = True

    idx = find_para_containing(doc, 'Florida Bar No')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{attorney_bar_no}')
            report['attorney_bar_no'] = True

    idx = find_para_containing(doc, 'Telephone:')
    if idx < 0:
        idx = find_para_containing(doc, 'Telephone')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{attorney_phone}')
            report['attorney_phone'] = True


def tag_order_footer(doc, report):
    idx = find_para_containing(doc, 'ORDERED on')
    if idx < 0:
        idx = find_para_containing(doc, 'Ordered on')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{order_date}')
            report['order_date'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{order_year}')
            report['order_year'] = True

    idx = find_para_containing(doc, 'Circuit Judge')
    if idx >= 0 and 'judge_name' not in report:
        if idx > 0:
            prev = doc.paragraphs[idx - 1]
            if is_empty_para(prev):
                clear_para_and_set_text(prev, '{judge_name}')
                report['judge_name'] = True
        if 'judge_name' not in report:
            para = doc.paragraphs[idx]
            groups = find_blank_groups(para)
            if groups:
                replace_blank_group(para, groups[0], '{judge_name}')
                report['judge_name'] = True


# ---------------------------------------------------------------------------
# Petition for Administration (P3-0100, P3-0120)
# ---------------------------------------------------------------------------

def tag_petition_for_admin(doc, form_id, report):
    is_testate = form_id == 'P3-0100'

    # Petitioner name
    idx = find_para_containing(doc, 'alleges:')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{petitioner_name}')
            report['petitioner_name'] = True

    # Paragraph 1: interest + address
    idx = find_para_containing(doc, 'interest in the above estate as')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{petitioner_interest}')
            report['petitioner_interest'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{petitioner_address}')
            report['petitioner_address'] = True

    # Paragraph 2: Decedent info
    idx = find_para_containing(doc, 'Decedent,')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'decedent,' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'address was' in before_text:
                field_map.append('{decedent_address}')
            elif before_text.endswith('are'):
                field_map.append('{decedent_ssn_last4}')
            elif 'died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',':
                if field_map and field_map[-1] == '{decedent_death_date}':
                    field_map.append('{decedent_death_year}')
                else:
                    field_map.append(None)
            elif before_text == 'at' or before_text.endswith(', at'):
                field_map.append('{decedent_death_place}')
            elif 'domiciled in' in before_text:
                field_map.append('{decedent_domicile}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Beneficiaries table
    idx = find_para_containing(doc, 'NAME')
    if idx >= 0 and 'ADDRESS' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#beneficiaries}{ben_name}\t{ben_address}\t{ben_relationship}\t{ben_year_of_birth}{/beneficiaries}')
        report['beneficiaries_loop'] = True
        if idx + 1 < len(doc.paragraphs):
            next_para = doc.paragraphs[idx + 1]
            if 'Minor' in next_para.text:
                remove_paragraph(next_para)

    # Venue
    idx = find_para_containing(doc, 'Venue of this proceeding')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[-1], '{venue_reason}')
            report['venue_reason'] = True

    # Paragraph 5: Proposed PR info
    idx = find_para_containing(doc, 'qualified to serve as personal representative')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_list = ['{pr_name}', '{pr_address}']
        # Additional blanks may be: relationship (for non-resident), etc.
        for i, g in enumerate(groups):
            if i < len(field_list):
                replace_blank_group(para, g, field_list[i])
                report[field_list[i].strip('{}')] = True
            elif i == len(groups) - 1:
                # Last blank is likely the relationship for non-resident PR
                replace_blank_group(para, g, '{pr_relationship}')
                report['pr_relationship'] = True

    # Will info (testate only)
    if is_testate:
        idx = find_para_containing(doc, "last will dated")
        if idx >= 0:
            para = doc.paragraphs[idx]
            groups = find_blank_groups(para)
            if len(groups) >= 1:
                replace_blank_group(para, groups[0], '{will_date}')
                report['will_date'] = True
            if len(groups) >= 2:
                replace_blank_group(para, groups[1], '{will_year}')
                report['will_year'] = True
            if len(groups) >= 3:
                replace_blank_group(para, groups[2], '{codicil_dates}')
                report['codicil_dates'] = True

    # "requests that... be appointed" paragraph — PR name
    idx = find_para_containing(doc, 'requests that')
    if idx >= 0 and 'appointed personal representative' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{pr_name}')

    # Domiciliary proceedings
    idx = find_para_containing(doc, 'Domiciliary')
    if idx < 0:
        idx = find_para_containing(doc, 'domiciliary')
    if idx >= 0:
        # "the address of which is" line
        addr_idx = find_para_containing(doc, 'the address of which is', idx)
        if addr_idx >= 0:
            para = doc.paragraphs[addr_idx]
            groups = find_blank_groups(para)
            if groups:
                replace_blank_group(para, groups[0], '{domiciliary_court_address}')
                report['domiciliary_court_address'] = True

        to_idx = find_para_starting_with(doc, 'to ', addr_idx if addr_idx >= 0 else idx)
        if to_idx >= 0:
            para = doc.paragraphs[to_idx]
            groups = find_blank_groups(para)
            if len(groups) >= 1:
                replace_blank_group(para, groups[0], '{domiciliary_representative}')
                report['domiciliary_representative'] = True
            if len(groups) >= 2:
                replace_blank_group(para, groups[1], '{domiciliary_representative_address}')
                report['domiciliary_representative_address'] = True

    tag_signing_block(doc, report)
    tag_attorney_block(doc, report)


# ---------------------------------------------------------------------------
# Order Admitting Will / Appointing PR (P3-0420, P3-0440)
# ---------------------------------------------------------------------------

def tag_order_appointing_pr(doc, form_id, report):
    is_testate = form_id == 'P3-0420'

    # Main paragraph(s) — find "petition of" or "instrument presented"
    for idx in range(len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        text = para.text.lower()
        if 'petition of' in text or 'instrument presented' in text or 'decedent died on' in text:
            groups = find_blank_groups(para)
            field_map = []
            for g_start, g_end in groups:
                before_text = ''
                for j in range(g_start - 1, -1, -1):
                    if para.runs[j].text.strip():
                        before_text = para.runs[j].text.strip().lower()
                        break

                if 'petition of' in before_text:
                    field_map.append('{petitioner_name}')
                elif 'estate of' in before_text or 'will of' in before_text or 'last will of' in before_text:
                    field_map.append('{decedent_full_name}')
                elif 'died on' in before_text:
                    field_map.append('{decedent_death_date}')
                elif before_text == ',' and field_map:
                    prev = field_map[-1] if field_map[-1] else ''
                    if 'death_date' in prev:
                        field_map.append('{decedent_death_year}')
                    elif 'will_date' in prev:
                        field_map.append('{will_year}')
                    else:
                        field_map.append(None)
                elif 'entitled' in before_text or 'qualified' in before_text:
                    field_map.append('{pr_name}')
                elif 'reason of' in before_text:
                    field_map.append('{pr_entitlement_reason}')
                elif 'will dated' in before_text or 'dated' == before_text:
                    field_map.append('{will_date}')
                elif 'attested by' in before_text:
                    field_map.append('{witnesses}')
                elif 'sum of' in before_text:
                    field_map.append('{bond_amount}')
                else:
                    field_map.append(None)

            for (g_start, g_end), tag in zip(groups, field_map):
                if tag:
                    replace_blank_group(para, (g_start, g_end), tag)
                    report[tag.strip('{}')] = True

    # ADJUDGED paragraphs — PR appointment
    for idx in range(len(doc.paragraphs)):
        para = doc.paragraphs[idx]
        if 'ADJUDGED that' in para.text and 'appointed personal representative' in para.text:
            groups = find_blank_groups(para)
            for g_start, g_end in groups:
                before_text = ''
                for j in range(g_start - 1, -1, -1):
                    if para.runs[j].text.strip():
                        before_text = para.runs[j].text.strip().lower()
                        break
                if 'adjudged that' in before_text:
                    replace_blank_group(para, (g_start, g_end), '{pr_name}')
                    report['pr_name'] = True
                elif 'sum of' in before_text:
                    replace_blank_group(para, (g_start, g_end), '{bond_amount}')
                    report['bond_amount'] = True

    tag_order_footer(doc, report)


# ---------------------------------------------------------------------------
# Oath of PR (P3-0600)
# ---------------------------------------------------------------------------

def tag_oath_pr(doc, report):
    # STATE OF
    idx = find_para_starting_with(doc, 'STATE OF')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{oath_state}')
            report['oath_state'] = True

    # COUNTY OF
    idx = find_para_starting_with(doc, 'COUNTY OF')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{oath_county}')
            report['oath_county'] = True

    # "I, (Affiant)" — PR name
    idx = find_para_containing(doc, 'Affiant')
    if idx >= 0 and 'state under oath' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        # Blank before "state under oath" is PR name; blank after "estate of" is decedent
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'estate of' in before_text:
                replace_blank_group(para, (g_start, g_end), '{decedent_full_name}')
                report['decedent_full_name'] = True

    # Paragraph 3: residence and PO address
    idx = find_para_containing(doc, 'My place of residence')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{pr_residence}')
            report['pr_residence'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{pr_po_address}')
            report['pr_po_address'] = True

    # Paragraph 6: Agent designation
    idx = find_para_containing(doc, 'I hereby designate')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_list = ['{agent_name}', '{agent_county}', '{agent_address}', '{agent_po_address}']
        for i, g in enumerate(groups):
            if i < len(field_list):
                replace_blank_group(para, g, field_list[i])
                report[field_list[i].strip('{}')] = True

    # Notary section
    idx = find_para_containing(doc, 'Sworn to')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        # Blanks: date, identification
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{notary_date}')
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{notary_identification}')

    idx = find_para_containing(doc, 'Notary Public State of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{notary_state}')

    idx = find_para_containing(doc, 'Commission Expires')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{notary_commission_expires}')

    idx = find_para_containing(doc, 'Commission Number')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{notary_commission_number}')

    # ACCEPTANCE section
    idx = find_para_containing(doc, 'I CERTIFY that I am a permanent resident')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{agent_acceptance_county}')

    # "Signed on" in acceptance
    idx = find_para_containing(doc, 'Signed on')
    if idx >= 0 and 'Signed on this' not in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{acceptance_date}')
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{acceptance_year}')


# ---------------------------------------------------------------------------
# Letters of Administration (P3-0700)
# ---------------------------------------------------------------------------

def tag_letters(doc, report):
    # WHEREAS paragraphs with decedent info
    idx = find_para_containing(doc, 'WHEREAS,')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'whereas,' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'resident of' in before_text:
                field_map.append('{decedent_residence}')
            elif 'died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',':
                if field_map and field_map[-1] == '{decedent_death_date}':
                    field_map.append('{decedent_death_year}')
                else:
                    field_map.append(None)
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Second WHEREAS — PR name
    idx2 = find_para_containing(doc, 'WHEREAS,', idx + 1 if idx >= 0 else 0)
    if idx2 >= 0 and 'appointed personal representative' in doc.paragraphs[idx2].text:
        para = doc.paragraphs[idx2]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{pr_name}')
            report['pr_name'] = True

    # NOW THEREFORE — PR name and decedent
    idx = find_para_containing(doc, 'NOW, THEREFORE')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{pr_name}')
        # "estate of" blank is decedent
        for g_start, g_end in groups[1:]:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'estate of' in before_text:
                replace_blank_group(para, (g_start, g_end), '{decedent_full_name}')

    tag_order_footer(doc, report)


# ---------------------------------------------------------------------------
# Notice to Creditors formal admin (P3-0740)
# ---------------------------------------------------------------------------

def tag_notice_creditors(doc, report):
    idx = find_para_containing(doc, 'administration of the estate')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'death was' in before_text or 'date of death' in before_text:
                field_map.append('{decedent_death_date}')
            elif before_text == ',' and field_map and field_map[-1] == '{decedent_death_date}':
                field_map.append('{decedent_death_year}')
            elif 'county' in before_text:
                field_map.append('{court_county}')
            elif 'address of which is' in before_text:
                field_map.append('{court_address}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Publication date
    idx = find_para_containing(doc, 'date of first publication')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if len(groups) >= 1:
            replace_blank_group(para, groups[0], '{publication_date}')
            report['publication_date'] = True
        if len(groups) >= 2:
            replace_blank_group(para, groups[1], '{publication_year}')
            report['publication_year'] = True

    tag_attorney_block(doc, report)


# ---------------------------------------------------------------------------
# Inventory (P3-0900)
# ---------------------------------------------------------------------------

def tag_inventory(doc, report):
    # "estate of [blank], deceased, who died on"
    idx = find_para_containing(doc, 'personal representative of the estate of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'estate of' in before_text:
                replace_blank_group(para, (g_start, g_end), '{decedent_full_name}')
                report['decedent_full_name'] = True
            elif 'died on' in before_text:
                replace_blank_group(para, (g_start, g_end), '{decedent_death_date}')
                report['decedent_death_date'] = True
            elif before_text == ',':
                replace_blank_group(para, (g_start, g_end), '{decedent_death_year}')
                report['decedent_death_year'] = True

    tag_signing_block(doc, report)
    tag_attorney_block(doc, report)


# ---------------------------------------------------------------------------
# Petition for Discharge (P5-0400)
# ---------------------------------------------------------------------------

def tag_petition_discharge(doc, report):
    # "Petitioner, [blank], as personal representative"
    idx = find_para_containing(doc, 'as personal representative')
    if idx >= 0 and 'alleges' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        if groups:
            replace_blank_group(para, groups[0], '{pr_name}')
            report['pr_name'] = True

    # Paragraph 1: decedent info
    idx = find_para_containing(doc, 'The decedent,')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'decedent,' in before_text:
                field_map.append('{decedent_full_name}')
            elif 'resident of' in before_text:
                field_map.append('{decedent_residence}')
            elif 'died on' in before_text:
                field_map.append('{decedent_death_date}')
            elif 'issued to petitioner on' in before_text or 'issued' in before_text:
                field_map.append('{letters_date}')
            elif before_text == ',':
                if field_map and 'death_date' in (field_map[-1] or ''):
                    field_map.append('{decedent_death_year}')
                elif field_map and 'letters_date' in (field_map[-1] or ''):
                    field_map.append('{letters_year}')
                else:
                    field_map.append(None)
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    # Interested persons table
    idx = find_para_containing(doc, 'NAME')
    if idx >= 0 and 'ADDRESS' in doc.paragraphs[idx].text:
        para = doc.paragraphs[idx]
        clear_para_and_set_text(para, '{#interested_persons}{ip_name}\t{ip_address}{/interested_persons}')
        report['interested_persons_loop'] = True

    tag_signing_block(doc, report)
    tag_attorney_block(doc, report)


# ---------------------------------------------------------------------------
# Order of Discharge (P5-0800)
# ---------------------------------------------------------------------------

def tag_order_discharge(doc, report):
    idx = find_para_containing(doc, 'Petition for Discharge of')
    if idx >= 0:
        para = doc.paragraphs[idx]
        groups = find_blank_groups(para)
        field_map = []
        for g_start, g_end in groups:
            before_text = ''
            for j in range(g_start - 1, -1, -1):
                if para.runs[j].text.strip():
                    before_text = para.runs[j].text.strip().lower()
                    break
            if 'discharge of' in before_text:
                field_map.append('{pr_name}')
            elif 'estate of' in before_text:
                field_map.append('{decedent_full_name}')
            else:
                field_map.append(None)

        for (g_start, g_end), tag in zip(groups, field_map):
            if tag:
                replace_blank_group(para, (g_start, g_end), tag)
                report[tag.strip('{}')] = True

    tag_order_footer(doc, report)


# ---------------------------------------------------------------------------
# Notice of Designation of Email Addresses (P1-0900)
# ---------------------------------------------------------------------------

def tag_email_designation(doc, report):
    # This form is simple — email addresses and signing
    tag_signing_block(doc, report)
    tag_attorney_block(doc, report)


# ---------------------------------------------------------------------------
# Main processing
# ---------------------------------------------------------------------------

def process_form(form_id, source_filename):
    source_path = os.path.join(SOURCE_DIR, source_filename)
    dest_path = os.path.join(TEMPLATES_DIR, form_id + '.docx')

    if not os.path.exists(source_path):
        print(f'  ERROR: Source file not found: {source_filename}')
        return None

    shutil.copy2(source_path, dest_path)
    doc = Document(dest_path)
    report = {}

    tag_case_header(doc, report)

    if form_id in ('P3-0100', 'P3-0120'):
        tag_petition_for_admin(doc, form_id, report)
    elif form_id in ('P3-0420', 'P3-0440'):
        tag_order_appointing_pr(doc, form_id, report)
    elif form_id == 'P3-0600':
        tag_oath_pr(doc, report)
    elif form_id == 'P3-0700':
        tag_letters(doc, report)
    elif form_id == 'P3-0740':
        tag_notice_creditors(doc, report)
    elif form_id == 'P1-0900':
        tag_email_designation(doc, report)
    elif form_id == 'P3-0900':
        tag_inventory(doc, report)
    elif form_id == 'P5-0400':
        tag_petition_discharge(doc, report)
    elif form_id == 'P5-0800':
        tag_order_discharge(doc, report)

    doc.save(dest_path)
    return report


def main():
    os.makedirs(TEMPLATES_DIR, exist_ok=True)

    print('=' * 70)
    print('Tagging 11 Core Formal Administration Probate Templates')
    print('=' * 70)
    print(f'Source: {SOURCE_DIR}')
    print(f'Output: {TEMPLATES_DIR}')
    print()

    for form_id, filename in sorted(FORMS.items()):
        print(f'--- {form_id}: {filename[:60]} ---')
        report = process_form(form_id, filename)
        if report:
            tagged = [k for k, v in report.items() if v]
            print(f'  Tagged: {", ".join(tagged)}')
        else:
            print(f'  FAILED')
        print()

    print('Done! Tagged templates saved to templates/')


if __name__ == '__main__':
    main()
