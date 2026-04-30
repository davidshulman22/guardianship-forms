#!/usr/bin/env python3
"""
Build clean single-column guardianship petition templates that match
Ginsberg Shulman's preferred format (based on Jill's Villareal edits).

Abandons the 2-column FLSSI layout in favor of:
  - Single-column caption via a 2-column table (left: case title, right: division/file)
  - Left-aligned body (no justified text — prevents word-spreading when {field} has line breaks)
  - Real Word table for next-of-kin and property value
  - docxtemplater conditional blocks for has/has-no alternatives, preneed, professional guardian
  - Running header with matter name + Page X of Y on pages 2+
  - Clean signature lines for Petitioner and Attorney

Currently builds:
  G2-010  Petition to Determine Incapacity
  G2-140  Notice of Designation of Email Addresses for Service
  G3-010  Petition for Appointment of Emergency Temporary Guardian
  G3-025  Petition for Appointment of Plenary Guardian of Property
  G3-026  Petition for Appointment of Limited Guardian of Person and Property
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
import re
import shutil
import zipfile
import tempfile

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), '..', 'templates')

FONT = 'Times New Roman'
BODY_SIZE = 12
SMALL_SIZE = 11

# ---------------------------------------------------------------------------
# Word numbering per docx-numbering skill (firm pleading conventions)
# ---------------------------------------------------------------------------
# Level 0: 1. 2. 3.    (body paragraphs)     firstLine=360, left=0, 12pt TNR
# Level 1: a. b. c.    (sub-paragraphs)      hanging=360 at left=720
# Level 2: i. ii. iii. (sub-sub-paragraphs)  hanging=360 at left=1440
# All paragraphs share numId=1 so they're continuous (skill: "do NOT create
# a new reference mid-document — that would restart numbering").

PLEADING_NUMBERING_XML = '''<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:numbering xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
  <w:abstractNum w:abstractNumId="0">
    <w:multiLevelType w:val="hybridMultilevel"/>
    <w:lvl w:ilvl="0">
      <w:start w:val="1"/>
      <w:numFmt w:val="decimal"/>
      <w:lvlText w:val="%1."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="0" w:firstLine="360"/>
      </w:pPr>
      <w:rPr>
        <w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman"/>
        <w:sz w:val="24"/>
      </w:rPr>
    </w:lvl>
    <w:lvl w:ilvl="1">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerLetter"/>
      <w:lvlText w:val="%2."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="720" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
    <w:lvl w:ilvl="2">
      <w:start w:val="1"/>
      <w:numFmt w:val="lowerRoman"/>
      <w:lvlText w:val="%3."/>
      <w:lvlJc w:val="left"/>
      <w:pPr>
        <w:ind w:left="1440" w:hanging="360"/>
      </w:pPr>
    </w:lvl>
  </w:abstractNum>
  <w:num w:numId="1">
    <w:abstractNumId w:val="0"/>
  </w:num>
</w:numbering>
'''

NUMBERING_CONTENT_TYPE = '<Override PartName="/word/numbering.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.numbering+xml"/>'
NUMBERING_REL_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering'


def _inject_numbering_part(docx_path):
    """Post-save: write word/numbering.xml with the firm's pleading numbering
    definition and register it in [Content_Types].xml and
    word/_rels/document.xml.rels.

    IMPORTANT: always OVERWRITES any existing numbering.xml — python-docx
    auto-creates a default one the moment we touch numPr, and it doesn't
    contain our 1./a./i. definitions at numId=1. The firm's definitions
    must win.
    """
    tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx')
    os.close(tmp_fd)
    try:
        with zipfile.ZipFile(docx_path, 'r') as zin, \
             zipfile.ZipFile(tmp_path, 'w', zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                if item.filename == 'word/numbering.xml':
                    # Skip — we'll write our own below.
                    continue
                content = zin.read(item.filename)
                if item.filename == '[Content_Types].xml':
                    text = content.decode('utf-8')
                    if 'numbering.xml' not in text:
                        text = text.replace('</Types>', NUMBERING_CONTENT_TYPE + '</Types>')
                        content = text.encode('utf-8')
                elif item.filename == 'word/_rels/document.xml.rels':
                    text = content.decode('utf-8')
                    if 'numbering.xml' not in text:
                        ids = [int(m) for m in re.findall(r'Id="rId(\d+)"', text)]
                        new_id = max(ids) + 1 if ids else 1
                        rel = f'<Relationship Id="rId{new_id}" Type="{NUMBERING_REL_TYPE}" Target="numbering.xml"/>'
                        text = text.replace('</Relationships>', rel + '</Relationships>')
                        content = text.encode('utf-8')
                zout.writestr(item, content)
            zout.writestr('word/numbering.xml', PLEADING_NUMBERING_XML)
        shutil.move(tmp_path, docx_path)
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)


def _apply_pleading_numbering(paragraph, level=0, num_id=1):
    """Attach <w:numPr><w:ilvl/><w:numId/></w:numPr> to a paragraph so Word
    auto-numbers it via the shared pleading numbering definition.
    """
    pPr = paragraph._p.get_or_add_pPr()
    existing = pPr.find(qn('w:numPr'))
    if existing is not None:
        pPr.remove(existing)
    numPr = OxmlElement('w:numPr')
    ilvl = OxmlElement('w:ilvl')
    ilvl.set(qn('w:val'), str(level))
    numPr.append(ilvl)
    nId = OxmlElement('w:numId')
    nId.set(qn('w:val'), str(num_id))
    numPr.append(nId)
    pPr.append(numPr)


def _set_run(run, *, size=BODY_SIZE, bold=False, italic=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # East Asia + complex script font for Word compat
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)


def _add_para(doc_or_cell, text='', *, size=BODY_SIZE, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None, left_indent=None,
              space_before=0, space_after=6, line_spacing=None, keep_with_next=False):
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if keep_with_next:
        pf.keep_with_next = True
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic)
    return p


def _add_run(p, text, *, size=BODY_SIZE, bold=False, italic=False):
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic)
    return run


def _borderless_table(doc, rows, cols, *, col_widths_in=None):
    """Create a table with no visible borders (for the caption layout)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    # Nuke borders
    tblPr = tbl._element.find(qn('w:tblPr')) or tbl._element.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return tbl


def _table_with_borders(doc, rows, cols, *, col_widths_in=None):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    tblPr = tbl._element.find(qn('w:tblPr')) or tbl._element.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return tbl


def _clear_cell(cell):
    """Remove the default empty paragraph a new cell ships with."""
    p = cell.paragraphs[0]
    if not p.runs and not p.text:
        p._element.getparent().remove(p._element)


def _cell_para(cell, text='', **kwargs):
    """Add a paragraph inside a table cell (works like _add_para)."""
    return _add_para(cell, text, **kwargs)


def _add_page_field(paragraph, field_code):
    """Insert a Word field (PAGE / NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    _set_run(run, size=SMALL_SIZE, italic=True)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' ' + field_code + ' '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def _set_different_first_page_header(section):
    """Enable 'different first page' so page 1 has no header."""
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)


# ---------------------------------------------------------------------------
# G3-025 Petition for Appointment of Plenary Guardian of Property
# ---------------------------------------------------------------------------

def build_g3_025():
    """Rebuilt per docx-numbering skill: real Word numbering, 1.5 line
    spacing on numbered paragraphs, Broward AI certification above signature.
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    # ---- Title ----
    _add_para(doc, 'PETITION FOR APPOINTMENT OF PLENARY GUARDIAN',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '(Incapacity - property)',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- Body ----
    # Intro — unnumbered per skill.
    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        'Petitioner\u2019s residence is {petitioner_residence} and petitioner\u2019s post office address is {petitioner_address}.')

    _pleading_para(doc,
        '{aip_name} (the AIP-Alleged Incapacitated Person) is an alleged incapacitated person who is {aip_age} years of age, whose address is {aip_address}.')

    _pleading_para(doc,
        'The nature of the Ward\u2019s alleged incapacity is {ward_incapacity_nature}.')

    _pleading_para(doc,
        'Petitioner {#has_alternatives}has knowledge, information and belief that there are alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. They are described as follows: {alternatives_description}. However, they are insufficient to meet the needs of the Ward because: {alternatives_insufficient_reason}. Thus, it is necessary that a plenary guardian be appointed to exercise all delegable rights of the Ward.{/has_alternatives}{^has_alternatives}has knowledge, information and belief that there are NO alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. Thus, it is necessary that a plenary guardian be appointed to exercise all delegable rights of the Ward, as no other alternatives exist.{/has_alternatives}')

    _pleading_para(doc,
        'Petitioner {#has_preneed}has knowledge, information and belief that the alleged incapacitated person has a preneed guardian designation. The designated preneed guardian is {preneed_guardian_name}.{/has_preneed}{^has_preneed}has knowledge, information and belief that the alleged incapacitated person has NO preneed guardian designation{#preneed_reason}, {preneed_reason}{/preneed_reason}.{/has_preneed}')

    _pleading_para(doc,
        'The names and addresses of the next of kin of the Ward are:',
        keep_with_next=True)
    _next_of_kin_table(doc)

    _pleading_para(doc,
        'The proposed guardian, {proposed_guardian_name}, whose address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as plenary guardian of the property of the Ward.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian} The relationship and previous association of the proposed guardian to the Ward (including any activities designated in Florida Statutes section 744.446(3)) are {proposed_guardian_relationship}. The proposed guardian should be appointed because {appointment_reason}.')

    _pleading_para(doc,
        'The nature and value of the property subject to guardianship are as follows:',
        keep_with_next=True)
    _property_items_table(doc)

    _pleading_para(doc,
        'Reasonable search has been made for all of the information required by Florida law and by the applicable Florida Probate Rules. Any such information that is not set forth in full above cannot be ascertained without delay that would adversely affect the Ward or the Ward\u2019s property.')

    # Closing — unnumbered per skill.
    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that {proposed_guardian_name} be appointed plenary guardian of the property of the Ward.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Appointment of Plenary Guardian of Property')
    _add_miami_dade_ai_certification(doc, 'Petition for Appointment of Plenary Guardian of Property')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'G3-025.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# Shared builders
# ---------------------------------------------------------------------------

def _apply_page_setup(doc):
    """1" margins, Times New Roman 12pt base style."""
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)


def _apply_running_header(doc, header_title):
    """Right-aligned running header on pages 2+ only:
         header_title (italic)
         Page X of Y
    """
    section = doc.sections[0]
    _set_different_first_page_header(section)
    header = section.header
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    h_p1 = header.add_paragraph()
    h_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_p1.paragraph_format.space_after = Pt(0)
    _add_run(h_p1, header_title, size=SMALL_SIZE, italic=True)
    h_p2 = header.add_paragraph()
    h_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_p2.paragraph_format.space_after = Pt(12)
    _add_run(h_p2, 'Page ', size=SMALL_SIZE, italic=True)
    _add_page_field(h_p2, 'PAGE')
    _add_run(h_p2, ' of ', size=SMALL_SIZE, italic=True)
    _add_page_field(h_p2, 'NUMPAGES')


def _add_guardianship_caption(doc, case_title_line, *, top_line='IN RE: GUARDIANSHIP OF'):
    """
    Standard guardianship caption:
      IN THE CIRCUIT COURT FOR {county_caption} COUNTY, FLORIDA (centered, bold)
      4-row borderless table (every line bold):
        top_line                       |  PROBATE DIVISION
        {aip_name_upper},              |
        case_title_line                |  File No. {file_no}
                                       |  Division {division}

    Pass top_line='IN RE:' for pre-guardianship petitions (e.g. incapacity).

    {county_caption} is set by prepareTemplateData() in app.js to an uppercased
    version of the matter county.
    """
    _add_para(doc, 'IN THE CIRCUIT COURT FOR {county_caption} COUNTY, FLORIDA',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    caption = _borderless_table(doc, rows=4, cols=2, col_widths_in=[3.5, 3.0])
    _clear_cell(caption.cell(0, 0)); _cell_para(caption.cell(0, 0), top_line, bold=True, space_after=0)
    _clear_cell(caption.cell(0, 1)); _cell_para(caption.cell(0, 1), 'PROBATE DIVISION', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 0)); _cell_para(caption.cell(1, 0), '{aip_name_upper},', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 1)); _cell_para(caption.cell(1, 1), '', bold=True, space_after=0)
    _clear_cell(caption.cell(2, 0)); _cell_para(caption.cell(2, 0), case_title_line, bold=True, space_after=0)
    _clear_cell(caption.cell(2, 1)); _cell_para(caption.cell(2, 1), 'File No. {file_no}', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 0)); _cell_para(caption.cell(3, 0), '', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 1)); _cell_para(caption.cell(3, 1), 'Division {division}', bold=True, space_after=0)
    _add_para(doc, '', space_after=18)


def _numbered_para(doc, number, text, *, indent=Inches(0.5), space_after=12, keep_with_next=False):
    """LEGACY: kept for forms that haven't been refactored to Word numbering.
    Produces an indented paragraph with a hardcoded '1.\\t' prefix.
    Prefer _pleading_para() for new work — see docx-numbering skill."""
    p = _add_para(doc, '', first_indent=indent, space_after=space_after, keep_with_next=keep_with_next)
    _add_run(p, f'{number}.\t')
    _add_run(p, text)
    return p


def _pleading_para(doc, text, *, level=0, keep_with_next=False, size=BODY_SIZE):
    """Numbered pleading paragraph per the docx-numbering skill:
      - Uses Word's native numbering (numPr → numId=1)
      - 1.5 line spacing (line=360, auto)
      - No before/after spacing (numbering definition controls layout)
      - 12pt Times New Roman run
    Call _ensure_pleading_numbering(doc) once before any _pleading_para() call
    so the numPr references a real numbering definition after save.
    """
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    # 1.5 line spacing, auto lineRule
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if keep_with_next:
        pf.keep_with_next = True
    _apply_pleading_numbering(p, level=level, num_id=1)
    if text:
        _add_run(p, text, size=size)
    return p


def _ensure_pleading_numbering(doc):
    """Marker no-op during build. The actual numbering part is injected in
    post-save via _inject_numbering_part() because python-docx has no API
    for creating numbering parts. Kept as a hook in case that changes.
    """
    pass


def _next_of_kin_table(doc):
    """Three-column bordered table with docxtemplater repeating row."""
    kin = _table_with_borders(doc, rows=2, cols=3, col_widths_in=[2.3, 2.9, 1.3])
    for cell in kin.rows[0].cells:
        _clear_cell(cell)
    _cell_para(kin.cell(0, 0), 'NAME', bold=True, space_after=0)
    _cell_para(kin.cell(0, 1), 'ADDRESS', bold=True, space_after=0)
    _cell_para(kin.cell(0, 2), 'RELATIONSHIP', bold=True, space_after=0)
    for cell in kin.rows[1].cells:
        _clear_cell(cell)
    _cell_para(kin.cell(1, 0), '{#next_of_kin}{name}', space_after=0)
    _cell_para(kin.cell(1, 1), '{address}', space_after=0)
    _cell_para(kin.cell(1, 2), '{relationship}{/next_of_kin}', space_after=0)
    return kin


def _property_items_table(doc):
    """Two-column bordered table for property inventory with repeating row."""
    prop = _table_with_borders(doc, rows=2, cols=2, col_widths_in=[4.5, 2.0])
    for cell in prop.rows[0].cells:
        _clear_cell(cell)
    _cell_para(prop.cell(0, 0), 'Nature of Property', bold=True, space_after=0)
    _cell_para(prop.cell(0, 1), 'Value', bold=True, space_after=0)
    for cell in prop.rows[1].cells:
        _clear_cell(cell)
    _cell_para(prop.cell(1, 0), '{#property_items}{item_description}', space_after=0)
    _cell_para(prop.cell(1, 1), '{item_value}{/property_items}', space_after=0)
    return prop


def _add_broward_ai_certification(doc, doc_title):
    """Broward 17th Circuit AO 2026-03-Gen certification. Renders only when
    BOTH the per-form 'used_ai' checkbox is on AND the matter county is
    Broward — wrapped in nested docxtemplater conditionals. The user must
    affirmatively check 'Was generative AI used to draft this document?' on
    the questionnaire (default OFF); these are court-published forms with
    deterministic merge-field substitution and most filings will not require
    the disclosure.
    """
    indent = Inches(0.5)
    text = (
        '{#used_ai}{#county_is_broward}The undersigned hereby certifies that generative artificial '
        f'intelligence was used to prepare this {doc_title}. The undersigned has '
        'independently verified the accuracy of every citation to the law and/or the '
        'record, and the accuracy of any language drafted by generative artificial '
        'intelligence, including quotations, citations, paraphrased assertions, facts, '
        'and legal analysis.{/county_is_broward}{/used_ai}'
    )
    _add_para(doc, text, first_indent=indent, space_before=6, space_after=18)


def _add_miami_dade_ai_certification(doc, doc_title):
    """Miami-Dade 11th Circuit AO 26-04 certification. Renders only when
    BOTH the per-form 'used_ai' checkbox is on AND the matter county is
    Miami-Dade — wrapped in nested docxtemplater conditionals. Text is the
    exact verbatim language required by AO 26-04. The {doc_title} parameter
    is intentionally unused (AO 26-04 mandates fixed wording referring to
    'this filing'); kept in the signature so call sites can mirror the
    Broward helper without refactoring.
    """
    del doc_title  # AO 26-04 wording is fixed.
    indent = Inches(0.5)
    text = (
        '{#used_ai}{#county_is_miami_dade}Generative artificial intelligence was used in the '
        'preparation of this filing. The undersigned certifies that all factual '
        'assertions, legal authority, and citations have been independently reviewed '
        'and verified for accuracy and accepts full responsibility for the contents '
        'of this filing.{/county_is_miami_dade}{/used_ai}'
    )
    _add_para(doc, text, first_indent=indent, space_before=6, space_after=18)


def _add_signature_block(doc, *, role='Petitioner', name_field='{petitioner_name}'):
    """Petitioner + Attorney signature lines and attorney contact block.

    Per Phase 1 (probate) and Phase 7b (guardianship), the signing date is
    rendered as a blank line — not asked at drafting time. The signer
    handwrites the date when they sign.
    """
    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=Inches(0.5), space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, f'{name_field}, {role}', space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for ' + role, space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)


# ---------------------------------------------------------------------------
# G3-026 Petition for Appointment of Limited Guardian of Person and Property
# ---------------------------------------------------------------------------

def build_g3_026():
    """Rebuilt per docx-numbering skill: real Word numbering (not hardcoded
    '1.\t'), 1.5 line spacing on numbered paragraphs, no empty spacer
    paragraphs between them, Broward AI certification above signature block.
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    # ---- Title ----
    _add_para(doc, 'PETITION FOR APPOINTMENT OF LIMITED GUARDIAN',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '(Incapacity - person and property)',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- Body ----
    # Intro paragraph — unnumbered per skill.
    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    # Numbered body paragraphs — Word auto-numbers via numPr/numId=1.
    _pleading_para(doc,
        'Petitioner\u2019s residence is {petitioner_residence} and petitioner\u2019s post office address is {petitioner_address}.')

    _pleading_para(doc,
        '{aip_name} (the AIP-Alleged Incapacitated Person) is an alleged incapacitated person who is {aip_age} years of age, whose address is {aip_address}.')

    # 3. Rights removal + delegation — the "keep_with_next" keeps the heading
    # glued to the checkbox list.
    _pleading_para(doc,
        'The Ward is unable to meet the essential requirements for certain aspects of the Ward\u2019s physical health or safety and certain aspects of the management of the Ward\u2019s financial resources. Consequently, the following rights should be removed:',
        keep_with_next=True)

    rights_indent = Inches(1.0)
    for check_name, label in [
        ('remove_marry_check',            'to marry,'),
        ('remove_vote_check',             'to vote,'),
        ('remove_govt_benefits_check',    'to personally apply for government benefits,'),
        ('remove_drivers_license_check',  'to have a driver\u2019s license,'),
        ('remove_travel_check',           'to travel,'),
        ('remove_employment_check',       'to seek or retain employment,'),
    ]:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')

    # Unnumbered blank line after the remove-rights checkbox list.
    _add_para(doc, '', space_after=0)

    _add_para(doc, 'and the following rights should be delegated to a limited guardian of the person and property of the Ward:',
              keep_with_next=True)

    for check_name, label in [
        ('delegate_contract_check',       'to contract,'),
        ('delegate_sue_check',            'to sue and defend lawsuits,'),
        ('delegate_govt_benefits_check',  'to apply for government benefits,'),
        ('delegate_property_check',       'to manage property or to make any gift or disposition of property,'),
        ('delegate_residence_check',      'to determine the Ward\u2019s residence,'),
        ('delegate_medical_check',        'to consent to medical and mental health treatment,'),
        ('delegate_social_check',         'to make decisions about the Ward\u2019s social environment or other social aspects of the Ward\u2019s life.'),
    ]:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')

    # Unnumbered blank line after the delegate-rights checkbox list.
    _add_para(doc, '', space_after=0)

    # 4. Approximate value and description (property table)
    _pleading_para(doc,
        'The approximate value and description of the Ward\u2019s property are as follows:',
        keep_with_next=True)
    _property_items_table(doc)

    # 5. Incapable property
    _pleading_para(doc,
        'The Ward is incapable of managing the following property, responsibility for which should be delegated to the guardian: {incapable_property}.')

    # 6. Alternatives — conditional
    _pleading_para(doc,
        'Petitioner {#has_alternatives}has knowledge, information and belief that there are alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. They are described as follows: {alternatives_description}. However, they are insufficient to meet the needs of the Ward because: {alternatives_insufficient_reason}. Thus, it is necessary that a limited guardian be appointed to exercise the delegable rights of the Ward identified above.{/has_alternatives}{^has_alternatives}has knowledge, information and belief that there are NO alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. Thus, it is necessary that a limited guardian be appointed to exercise the delegable rights of the Ward identified above, as no other alternatives exist.{/has_alternatives}')

    # 7. Preneed — conditional
    _pleading_para(doc,
        'Petitioner {#has_preneed}has knowledge, information and belief that the alleged incapacitated person has a preneed guardian designation. The designated preneed guardian is {preneed_guardian_name}.{/has_preneed}{^has_preneed}has knowledge, information and belief that the alleged incapacitated person has NO preneed guardian designation{#preneed_reason}, {preneed_reason}{/preneed_reason}.{/has_preneed}')

    # 8. Next of kin
    _pleading_para(doc,
        'The names and addresses of the next of kin of the Ward are:',
        keep_with_next=True)
    _next_of_kin_table(doc)

    # 9. Proposed guardian
    _pleading_para(doc,
        'The proposed guardian, {proposed_guardian_name}, whose address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as limited guardian of the person and property of the Ward.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian} The relationship and previous association of the proposed guardian to the Ward (including any activities designated in Florida Statutes section 744.446(3)) are {proposed_guardian_relationship}. The proposed guardian should be appointed because {appointment_reason}.')

    # 10. Reasonable search
    _pleading_para(doc,
        'Reasonable search has been made for all of the information required by Florida law and by the applicable Florida Probate Rules. Any such information that is not set forth in full above cannot be ascertained without delay that would adversely affect the Ward or the Ward\u2019s property.')

    # Closing (WHEREFORE) — unnumbered per skill.
    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that {proposed_guardian_name} be appointed limited guardian of the person and property of the Ward.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    # Broward AI certification — conditional on county (populated by app.js).
    _add_broward_ai_certification(doc, 'Petition for Appointment of Limited Guardian of Person and Property')
    _add_miami_dade_ai_certification(doc, 'Petition for Appointment of Limited Guardian of Person and Property')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'G3-026.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)  # add word/numbering.xml + rels
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-010 Petition for Appointment of Emergency Temporary Guardian
# ---------------------------------------------------------------------------

def build_g3_010():
    """Legacy entry point — produces templates/G3-010.docx."""
    _build_emergency_template('G3-010.docx')


def build_g3_emergency():
    """Smart-template alias for the Emergency Temporary Guardian petition.
    Identical content to G3-010 — separate output path so the wizard's smart-
    template namespace (G3-PETITION / G3-OATH / G3-ORDER / G3-LETTERS /
    G3-EMERGENCY) is consistent. Legacy bundle preset still references G3-010
    so both files coexist.
    """
    _build_emergency_template('G3-EMERGENCY.docx')


def _build_emergency_template(filename):
    """Rebuilt per docx-numbering skill."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name} — ETG')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    # ---- Title ----
    _add_para(doc, 'PETITION FOR APPOINTMENT OF',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'EMERGENCY TEMPORARY GUARDIAN',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Intro — unnumbered.
    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        'Petitioner\u2019s residence is {petitioner_residence}, petitioner\u2019s post office address is {petitioner_address}, and petitioner\u2019s telephone number is {petitioner_phone}.')

    _pleading_para(doc,
        'A Petition to Determine Incapacity has been filed in this Court with respect to {aip_name}, an alleged incapacitated person, but a guardian has not been appointed.')

    _pleading_para(doc,
        'Petitioner is an adult interested in the welfare of the alleged incapacitated person.')

    _pleading_para(doc,
        'There appears to be imminent danger that the physical or mental health or safety of the alleged incapacitated person will be seriously impaired or that the property of that person is in danger of being wasted, misappropriated or lost unless immediate action is taken because: {imminent_danger_reason}.')

    _pleading_para(doc,
        '{aip_name} is an alleged incapacitated person whose date of birth is {aip_dob}, and who is {aip_age} years of age, whose address is {aip_address}.')

    _pleading_para(doc,
        'The nature of the alleged incapacitated person\u2019s alleged incapacity is {aip_incapacity_nature}.')

    _pleading_para(doc,
        'The names and addresses of the next of kin of the alleged incapacitated person are:',
        keep_with_next=True)
    _next_of_kin_table(doc)

    _pleading_para(doc,
        'The proposed emergency temporary guardian, {proposed_guardian_name}, whose address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as guardian of the alleged incapacitated person.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian}')

    _pleading_para(doc,
        'The relationship and previous association of the proposed emergency temporary guardian to the alleged incapacitated person is {proposed_guardian_relationship}.')

    _pleading_para(doc,
        'The proposed emergency temporary guardian should be appointed because: {appointment_reason}.')

    _pleading_para(doc,
        'The nature and value of the property subject to guardianship are as follows:',
        keep_with_next=True)
    _property_items_table(doc)

    _pleading_para(doc,
        'Petitioner {#has_alternatives}has knowledge, information and belief that there are possible alternatives to emergency temporary guardianship including, but not limited to, trust agreements, powers of attorney, designations of health care surrogate, guardian advocate under Florida Statutes section 744.3085, or other advance directives. They are described as follows: {alternatives_description}. However, they are not sufficient to meet the current needs of the alleged incapacitated person for the reasons specified in paragraph 4 above.{/has_alternatives}{^has_alternatives}has knowledge, information and belief that there are NO alternatives to emergency temporary guardianship such as trust agreements, powers of attorney, designations of health care surrogate, guardian advocate under Florida Statutes section 744.3085, or other advance directives. Thus, immediate appointment of an emergency temporary guardian is necessary to protect the alleged incapacitated person for the reasons specified in paragraph 4 above.{/has_alternatives}')

    _pleading_para(doc,
        'Petitioner {#has_preneed}has knowledge, information and belief that the alleged incapacitated person has a preneed guardian designation. The designated preneed guardian is {preneed_guardian_name}.{/has_preneed}{^has_preneed}has knowledge, information and belief that the alleged incapacitated person has NO preneed guardian designation{#preneed_reason}, {preneed_reason}{/preneed_reason}.{/has_preneed}')

    # Closing — unnumbered.
    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that summary proceedings be held upon this petition, that the court appoint an attorney to represent the alleged incapacitated person in these proceedings, and that an emergency temporary guardian of the {guardianship_scope} be appointed for the alleged incapacitated person.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc, 'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
              first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Appointment of Emergency Temporary Guardian')
    _add_miami_dade_ai_certification(doc, 'Petition for Appointment of Emergency Temporary Guardian')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, filename)
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G2-010 Petition to Determine Incapacity
# ---------------------------------------------------------------------------

def build_g2_010():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'In re {aip_name} — Incapacity')

    _ensure_pleading_numbering(doc)
    _add_guardianship_caption(doc, 'An alleged incapacitated person',
                              top_line='IN RE:')

    # ---- Title ----
    _add_para(doc, 'PETITION TO DETERMINE INCAPACITY',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Intro — unnumbered.
    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        'Petitioner is an adult, age {petitioner_age}, whose present address is {petitioner_address}, and whose relationship to the hereafter named alleged incapacitated person is {petitioner_relationship}.')

    _pleading_para(doc,
        'Petitioner believes {aip_name}, age {aip_age}, a resident of {aip_county} County, Florida, whose primary spoken language is {aip_primary_language}, and whose present address is {aip_address}, to be incapacitated, based upon the following factual information: {factual_basis}.')

    _pleading_para(doc,
        'The names and addresses of all persons known to petitioner who have knowledge of such facts through personal observation are: {knowledge_witnesses}.')

    # 4. Rights AIP cannot exercise (checkbox list)
    _pleading_para(doc,
        'The alleged incapacitated person is incapable of exercising the following rights:',
        keep_with_next=True)

    rights_indent = Inches(1.0)
    for check_name, label in [
        ('right_marry_check',            'to marry'),
        ('right_vote_check',             'to vote'),
        ('right_contract_check',         'to contract'),
        ('right_travel_check',           'to travel'),
        ('right_sue_check',              'to sue and defend lawsuits'),
        ('right_drivers_license_check',  'to have a driver\u2019s license'),
        ('right_residency_check',        'to determine his or her residency'),
        ('right_employment_check',       'to seek or retain employment'),
        ('right_medical_check',          'to consent to medical and mental health treatment'),
        ('right_property_check',         'to manage property or to make any gift or disposition of property'),
        ('right_govt_benefits_check',    'to personally apply for government benefits'),
        ('right_social_check',           'to make decisions about his or her social environment or other social aspects of his or her life'),
    ]:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')

    # Unnumbered blank line after the rights checkbox list.
    _add_para(doc, '', space_after=0)

    indent = Inches(0.5)
    _add_para(doc, 'Petitioner has insufficient experience to make judgments concerning the rights the alleged incapacitated person is incapable of exercising (strike if not applicable).',
              first_indent=indent, space_after=12, italic=True)

    # 5. Plenary or limited — numbered, with two inline checkboxes.
    p = _pleading_para(doc, '')
    _add_run(p, '(')
    _add_run(p, '{guardianship_plenary_check}')
    _add_run(p, ') Plenary  (')
    _add_run(p, '{guardianship_limited_check}')
    _add_run(p, ') Limited guardianship is being sought for the alleged incapacitated person (check if known).')

    # 6. Alternatives — checkbox list
    _pleading_para(doc,
        'There are the following possible alternatives to guardianship (check all those of which petitioner is aware):',
        keep_with_next=True)

    for check_name, label in [
        ('alt_trust_check',                'trust agreements;'),
        ('alt_poa_check',                  'powers of attorney;'),
        ('alt_healthcare_surrogate_check', 'designations of health care surrogates;'),
        ('alt_advance_directives_check',   'other advance directives; or'),
        ('alt_supported_decision_check',   'supported decision making agreements or other assistance to the alleged incapacitated person in exercising the person\u2019s rights.'),
    ]:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')

    # Unnumbered blank line after the alternatives checkbox list.
    _add_para(doc, '', space_after=0)

    _add_para(doc,
        'If a guardianship is being sought, explain why the checked possible alternatives to guardianship are inappropriate or insufficient to meet the needs of the alleged incapacitated person or allow the person to independently exercise the person\u2019s rights: {alternatives_explanation}.',
        first_indent=indent, space_after=12)

    # 7. Next of kin
    _pleading_para(doc,
        'The names, addresses and relationships of all known next of kin of the alleged incapacitated person are (give years of birth of any who are minors):',
        keep_with_next=True)
    _next_of_kin_table(doc)

    # 8. Physician
    _pleading_para(doc,
        'The alleged incapacitated person\u2019s attending or family physician is: {physician_name}, {physician_address}, {physician_phone}.')

    # Closing — unnumbered.
    _add_para(doc,
        'Petitioner requests that an examination be made as to the mental and physical condition of the alleged incapacitated person as provided by law, and that an order be entered determining the mental and physical capacity of said person.',
        first_indent=indent, space_before=12, space_after=12)

    _add_para(doc, 'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
              first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition to Determine Incapacity')
    _add_miami_dade_ai_certification(doc, 'Petition to Determine Incapacity')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'G2-010.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G2-140 Notice of Designation of Email Addresses for Service
# ---------------------------------------------------------------------------

def build_g2_140():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    # ---- Title ----
    _add_para(doc, 'NOTICE OF DESIGNATION OF E-MAIL ADDRESSES',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'FOR SERVICE OF DOCUMENTS',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)

    _add_para(doc,
        'Pursuant to Florida Rule of Judicial Administration 2.516(b)(1)(A), the undersigned counsel designates the following e-mail addresses for service in this matter:',
        first_indent=indent, space_after=12)

    # Primary / secondary — pull directly from attorney defaults so Jill's
    # guardianship matters route to maribel@ as secondary automatically.
    p = _add_para(doc, '', left_indent=Inches(0.5), space_after=6)
    _add_run(p, 'Primary E-Mail Address:\t', bold=True)
    _add_run(p, '{attorney_email}')
    p = _add_para(doc, '', left_indent=Inches(0.5), space_after=18)
    _add_run(p, 'Secondary E-Mail Address:\t', bold=True)
    _add_run(p, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}{^attorney_email_secondary}(none){/attorney_email_secondary}')

    # Signed on this ___ day — blank, signer handwrites.
    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=indent, space_after=24)

    # Broward AI certification — only renders in Broward-county matters.
    _add_broward_ai_certification(doc, 'Notice of Designation of E-Mail Addresses for Service')
    _add_miami_dade_ai_certification(doc, 'Notice of Designation of E-Mail Addresses for Service')

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for Petitioner', space_after=18)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=0)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=0)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=18)

    # Certificate of service
    _add_para(doc, 'CERTIFICATE OF SERVICE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    _add_para(doc,
        'I CERTIFY that a true and correct copy of the foregoing has been furnished to {cos_served_to} by {cos_service_method} on this _____ day of __________, 20___.',
        first_indent=indent, space_after=24)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for Petitioner', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'G2-140.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-PETITION  Smart petition for guardian appointment
# ---------------------------------------------------------------------------
#
# Replaces 9 FLSSI forms via matter-flag branching:
#   G-3.020 / .021 / .022   (minor — both / person / property)
#   G-3.023 / .024 / .025   (plenary — both / person / property)
#   G-3.026 / .027 / .028   (limited — both / person / property)
#
# Branch axes (matter.matterData flags set by the Open Guardianship wizard):
#   is_minor              \   exactly one is true
#   is_adult_incapacity   /
#   is_plenary            \   exactly one true (adult only)
#   is_limited            /
#   scope_person          \
#   scope_property         |  exactly one true (minor + adult)
#   scope_both            /
#
# Derived flags computed in prepareTemplateData() so the template stays
# readable:
#   includes_property = scope_property OR scope_both
#   includes_person   = scope_person   OR scope_both
#   guardian_kind_caps     — "PLENARY GUARDIAN" / "LIMITED GUARDIAN" /
#                            "GUARDIAN OF MINOR"
#   scope_subtitle         — "(Incapacity - person)" / "(Guardianship of
#                            Person and Property)" / etc.
#   scope_phrase           — "of the person", "of the property", "of the
#                            person and property"
#   ward_term              — "Ward" (adult) / "minor" (minor)
#   ward_term_lower        — "ward"  (adult) / "minor" (minor)
# ---------------------------------------------------------------------------

def build_g3_petition():
    """Smart petition replacing the 9 G-3.020 through G-3.028 FLSSI forms.

    Variants are driven by pre-computed presentation tokens emitted by
    prepareTemplateData() in app.js, plus a small set of matter-level flags
    set by the Open Guardianship wizard. Token approach minimizes nested
    docxtemplater conditionals and keeps the body text readable.

    Pre-computed tokens (see prepareTemplateData):
      {guardian_kind_caps}        "GUARDIAN OF MINOR" / "PLENARY GUARDIAN" / "LIMITED GUARDIAN"
      {guardian_kind_lower}       "guardian" / "plenary guardian" / "limited guardian"
      {scope_subtitle}            "(Guardianship of Person)" / "(Incapacity - person)" / etc.
      {scope_phrase}              "of the person" / "of the property" / "of the person and property"
      {ward_term}                 "Ward" / "minor"
      {ward_term_lower}           "ward"  / "minor"
      {delegable_rights_phrase}   plenary: "all delegable rights of the Ward"
                                  limited: "the delegable rights of the Ward identified above"
      {limited_aspects_phrase}    text describing what aspects the Ward fails to meet,
                                  varies by scope (person / property / both)

    Matter flags (set by wizardLoadGuardianshipForms):
      is_minor / is_adult_incapacity   — caption + body terminology
      is_limited                       — rights-removal/delegation sections
      includes_property                — property table + incapable-property gate
      show_limited_person_rights       — person-removal rights list (limited + person|both)
      show_limited_property_rights_only — property-only rights-removal list (limited + property)
      show_limited_property_section    — incapable-property paragraph (limited + includes property)
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    # Caption — case-title-line via {#is_minor} for the only place that branches.
    _add_para(doc, 'IN THE CIRCUIT COURT FOR {county_caption} COUNTY, FLORIDA',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    caption = _borderless_table(doc, rows=4, cols=2, col_widths_in=[3.5, 3.0])
    _clear_cell(caption.cell(0, 0)); _cell_para(caption.cell(0, 0), 'IN RE: GUARDIANSHIP OF', bold=True, space_after=0)
    _clear_cell(caption.cell(0, 1)); _cell_para(caption.cell(0, 1), 'PROBATE DIVISION', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 0)); _cell_para(caption.cell(1, 0), '{aip_name_upper},', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 1)); _cell_para(caption.cell(1, 1), '', bold=True, space_after=0)
    _clear_cell(caption.cell(2, 0))
    p = _cell_para(caption.cell(2, 0), '', bold=True, space_after=0)
    _add_run(p, '{#is_minor}A minor{/is_minor}{^is_minor}An alleged incapacitated person{/is_minor}', bold=True)
    _clear_cell(caption.cell(2, 1)); _cell_para(caption.cell(2, 1), 'File No. {file_no}', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 0)); _cell_para(caption.cell(3, 0), '', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 1)); _cell_para(caption.cell(3, 1), 'Division {division}', bold=True, space_after=0)
    _add_para(doc, '', space_after=18)

    _add_para(doc, 'PETITION FOR APPOINTMENT OF {guardian_kind_caps}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{scope_subtitle}',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    # 1. Petitioner residence + address (same across all variants).
    _pleading_para(doc,
        'Petitioner’s residence is {petitioner_residence} and petitioner’s post office address is {petitioner_address}.')

    # 2. Subject (Ward / minor).
    _pleading_para(doc,
        '{#is_minor}{aip_name} is a minor, who is {aip_age} years of age. The residence of the minor is {aip_address}.{/is_minor}'
        '{^is_minor}{aip_name} (the Ward) is an alleged incapacitated person who is {aip_age} years of age, whose address is {aip_address}.{/is_minor}')

    # 3. Adult only: nature of incapacity.
    _pleading_para(doc,
        '{#is_adult_incapacity}The nature of the Ward’s alleged incapacity is {ward_incapacity_nature}.{/is_adult_incapacity}'
        '{#is_minor}The names and addresses of the parents, or if none, the next of kin, of the minor are as follows:{/is_minor}',
        keep_with_next=True)

    # Minor: NOK table renders right after "parents/NOK" sentence (item 3).
    _gated_table(doc, '{#is_minor}', _next_of_kin_table, '{/is_minor}')

    # 4. Limited only: rights-removal lead-in + scoped rights lists.
    _pleading_para(doc,
        '{#is_limited}The Ward is unable to meet the essential requirements for certain aspects of the Ward’s {limited_aspects_phrase}. Consequently, the following rights should be removed:{/is_limited}',
        keep_with_next=True)

    _add_limited_rights_removal_lists(doc)

    # Delegation lead-in.
    _add_para(doc,
        '{#is_limited}and the following rights should be delegated to a limited guardian {scope_phrase} of the Ward:{/is_limited}',
        keep_with_next=True)

    _add_limited_rights_delegation_lists(doc)

    # 5. Property table — minor uses different lead-in than adult.
    _pleading_para(doc,
        '{#includes_property}{#is_minor}The nature and value of property subject to the guardianship are as follows:{/is_minor}'
        '{#is_adult_incapacity}The {limited_property_lead}value and description of the Ward’s property are as follows:{/is_adult_incapacity}{/includes_property}',
        keep_with_next=True)
    _gated_table(doc, '{#includes_property}', _property_items_table, '{/includes_property}')

    # 6. Limited + property scope: incapable-property delegation.
    _pleading_para(doc,
        '{#show_limited_property_section}The Ward is incapable of managing the following property, responsibility for which should be delegated to the guardian: {incapable_property}.{/show_limited_property_section}')

    # 7. Alternatives — minor and adult phrase the conclusion differently.
    _pleading_para(doc,
        '{#is_minor}There {#has_alternatives}are{/has_alternatives}{^has_alternatives}are not{/has_alternatives} possible alternatives to guardianship known to petitioner, such as trust agreements, powers of attorney, surrogates, guardian advocates appointed pursuant to Florida Statutes section 744.3085 or advance directives.{#has_alternatives} They are described as follows: {alternatives_description}. However, they are insufficient to meet the needs of the minor because {alternatives_insufficient_reason}.{/has_alternatives} Thus, it is necessary that a guardian be appointed for the minor.{/is_minor}'
        '{#is_adult_incapacity}Petitioner {#has_alternatives}has knowledge, information and belief that there are alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. They are described as follows: {alternatives_description}. However, they are insufficient to meet the needs of the Ward because: {alternatives_insufficient_reason}. Thus, it is necessary that a {guardian_kind_lower} be appointed to exercise {delegable_rights_phrase}.{/has_alternatives}{^has_alternatives}has knowledge, information and belief that there are NO alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. Thus, it is necessary that a {guardian_kind_lower} be appointed to exercise {delegable_rights_phrase}, as no other alternatives exist.{/has_alternatives}{/is_adult_incapacity}')

    # 8. Adult only: preneed designation.
    _pleading_para(doc,
        '{#is_adult_incapacity}Petitioner {#has_preneed}has knowledge, information and belief that the alleged incapacitated person has a preneed guardian designation. The designated preneed guardian is {preneed_guardian_name}.{/has_preneed}{^has_preneed}has knowledge, information and belief that the alleged incapacitated person has NO preneed guardian designation{#preneed_reason}, {preneed_reason}{/preneed_reason}.{/has_preneed}{/is_adult_incapacity}')

    # 9. Adult NOK table.
    _pleading_para(doc,
        '{#is_adult_incapacity}The names and addresses of the next of kin of the Ward are:{/is_adult_incapacity}',
        keep_with_next=True)
    _gated_table(doc, '{#is_adult_incapacity}', _next_of_kin_table, '{/is_adult_incapacity}')

    # 10. Proposed guardian.
    _pleading_para(doc,
        'The proposed guardian, {proposed_guardian_name}, whose address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as {guardian_kind_lower} {scope_phrase} of the {ward_term_lower}.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian} The relationship and previous association of the proposed guardian to the {ward_term_lower} (including any activities designated in Florida Statutes section 744.446(3)) are {proposed_guardian_relationship}. The proposed guardian should be appointed because {appointment_reason}.')

    # 11. Adult only: reasonable search.
    _pleading_para(doc,
        '{#is_adult_incapacity}Reasonable search has been made for all of the information required by Florida law and by the applicable Florida Probate Rules. Any such information that is not set forth in full above cannot be ascertained without delay that would adversely affect the Ward or the Ward’s property.{/is_adult_incapacity}')

    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that {proposed_guardian_name} be appointed {guardian_kind_lower} {scope_phrase} of the {ward_term_lower}.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Appointment of Guardian')
    _add_miami_dade_ai_certification(doc, 'Petition for Appointment of Guardian')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'G3-PETITION.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


def _gated_table(doc, open_tag, table_builder_fn, close_tag):
    """Wrap a table in docxtemplater open/close tags so the whole table —
    including the repeating-row tags inside — only renders when the gate
    evaluates true."""
    _add_para(doc, open_tag, space_after=0)
    table_builder_fn(doc)
    _add_para(doc, close_tag, space_after=0)


def _add_limited_rights_removal_lists(doc):
    """Two checkbox lists — only one renders, gated by a single matter flag.
    Person-list applies to person-only and both scopes (per FLSSI G-3.026
    and G-3.027). Property-only list per G-3.028 is a different set."""
    rights_indent = Inches(1.0)

    person_rights = [
        ('remove_marry_check',           'to marry,'),
        ('remove_vote_check',            'to vote,'),
        ('remove_govt_benefits_check',   'to personally apply for government benefits,'),
        ('remove_drivers_license_check', 'to have a driver’s license,'),
        ('remove_travel_check',          'to travel,'),
        ('remove_employment_check',      'to seek or retain employment,'),
    ]
    property_only_rights = [
        ('remove_contract_check', 'to contract,'),
        ('remove_sue_check',      'to sue and defend lawsuits,'),
        ('remove_property_check', 'to manage property or to make any gift or disposition of property,'),
    ]

    _add_para(doc, '{#show_limited_person_rights}', space_after=0)
    for check_name, label in person_rights:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')
    _add_para(doc, '{/show_limited_person_rights}', space_after=0)

    _add_para(doc, '{#show_limited_property_rights_only}', space_after=0)
    for check_name, label in property_only_rights:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')
    _add_para(doc, '{/show_limited_property_rights_only}', space_after=0)


def _add_limited_rights_delegation_lists(doc):
    """Three delegation lists; the right one renders by scope. Limited only.
    Person-only matches FLSSI G-3.027; property-only matches G-3.028;
    both-scope matches G-3.026."""
    rights_indent = Inches(1.0)

    person_only = [
        ('delegate_residence_check', 'to determine the Ward’s residence,'),
        ('delegate_medical_check',   'to consent to medical and mental health treatment,'),
        ('delegate_social_check',    'to make decisions about the Ward’s social environment or other social aspects of the Ward’s life.'),
    ]
    property_only = [
        ('delegate_contract_check',      'to contract,'),
        ('delegate_sue_check',           'to sue and defend lawsuits,'),
        ('delegate_govt_benefits_check', 'to apply for government benefits,'),
        ('delegate_property_check',      'to manage property or to make any gift or disposition of property.'),
    ]
    both_scope = [
        ('delegate_contract_check',      'to contract,'),
        ('delegate_sue_check',           'to sue and defend lawsuits,'),
        ('delegate_govt_benefits_check', 'to apply for government benefits,'),
        ('delegate_property_check',      'to manage property or to make any gift or disposition of property,'),
        ('delegate_residence_check',     'to determine the Ward’s residence,'),
        ('delegate_medical_check',       'to consent to medical and mental health treatment,'),
        ('delegate_social_check',        'to make decisions about the Ward’s social environment or other social aspects of the Ward’s life.'),
    ]

    # Person-only: is_limited && is_scope_person_only
    _add_para(doc, '{#is_limited}{#is_scope_person_only}', space_after=0)
    for check_name, label in person_only:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')
    _add_para(doc, '{/is_scope_person_only}{/is_limited}', space_after=0)

    # Property-only: is_limited && is_scope_property_only
    _add_para(doc, '{#is_limited}{#is_scope_property_only}', space_after=0)
    for check_name, label in property_only:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')
    _add_para(doc, '{/is_scope_property_only}{/is_limited}', space_after=0)

    # Both: is_limited && scope_both
    _add_para(doc, '{#is_limited}{#scope_both}', space_after=0)
    for check_name, label in both_scope:
        p = _add_para(doc, '', left_indent=rights_indent, space_after=0)
        _add_run(p, '(')
        _add_run(p, '{' + check_name + '}')
        _add_run(p, f') {label}')
    _add_para(doc, '{/scope_both}{/is_limited}', space_after=0)


# ---------------------------------------------------------------------------
# G3-OATH  Oath of Guardian + Designation of Resident Agent and Acceptance
# ---------------------------------------------------------------------------
#
# Replaces FLSSI G-3.076 (regular guardian) and G-3.078 (emergency temporary
# guardian). Branches on matter flag is_emergency_temporary for title +
# item 1 phrasing.
# ---------------------------------------------------------------------------

def build_g3_oath():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, '{#is_minor}A minor{/is_minor}{^is_minor}An alleged incapacitated person{/is_minor}')

    _add_para(doc,
        'OATH OF {#is_emergency_temporary}EMERGENCY TEMPORARY GUARDIAN{/is_emergency_temporary}{^is_emergency_temporary}GUARDIAN{/is_emergency_temporary},',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'DESIGNATION OF RESIDENT AGENT AND ACCEPTANCE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'STATE OF FLORIDA', space_after=0)
    _add_para(doc, 'COUNTY OF {oath_county}', space_after=18)

    _add_para(doc, 'I, {proposed_guardian_name} (Affiant), state under oath that:',
              space_after=12)

    _pleading_para(doc,
        'I will faithfully perform the duties of {#is_emergency_temporary}emergency temporary guardian{/is_emergency_temporary}{^is_emergency_temporary}guardian{/is_emergency_temporary} of {aip_name} (the {ward_term}) according to law.')

    _pleading_para(doc,
        'My place of residence is {proposed_guardian_residence} and my post office address is {proposed_guardian_address}.')

    _pleading_para(doc,
        'I hereby designate {resident_agent_name}, who {#agent_is_fl_bar}is{/agent_is_fl_bar}{^agent_is_fl_bar}is not{/agent_is_fl_bar} a member of The Florida Bar, a resident of {agent_county} County, Florida, whose office address is {resident_agent_address} as my agent for the service of process or notice in any action against me, either in my representative capacity, or personally, if the personal action accrued in the performance of my duties as such guardian.')

    indent = Inches(0.5)
    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=indent, space_before=18, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{proposed_guardian_name}, Affiant', space_after=18)

    _add_para(doc,
        'Sworn to (or affirmed) and subscribed before me by means of ☐ online notarization or ☐ physical presence this _____ day of __________, 20____, by {proposed_guardian_name}, who is personally known to me or produced ______________________________ as identification.',
        space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Notary Public', space_after=0)
    _add_para(doc, 'State of Florida', space_after=24)

    # ---- Designation of Resident Agent + Acceptance ----
    _add_para(doc, 'ACCEPTANCE OF DESIGNATION AS RESIDENT AGENT',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=12, space_after=12)
    _add_para(doc,
        'I, {resident_agent_name}, certify that I am a permanent resident of {agent_county} County, Florida, and my office address is {resident_agent_address}. I hereby accept the foregoing designation as Resident Agent.',
        space_after=18)
    _add_para(doc, 'Signed on this _____ day of __________, 20___.', space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{resident_agent_name}, Resident Agent', space_after=18)

    _add_broward_ai_certification(doc, 'Oath of Guardian')
    _add_miami_dade_ai_certification(doc, 'Oath of Guardian')

    out_path = os.path.join(TEMPLATE_DIR, 'G3-OATH.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-ORDER  Order Appointing Guardian
# ---------------------------------------------------------------------------
#
# Replaces 13 FLSSI orders (G-3.060 through G-3.075). Branches via matter
# flags (minor / plenary / limited, scope_*) plus per-form has_advance_directive.
# Per the project's hard rule: judge-signed orders carry NO AI cert.
# ---------------------------------------------------------------------------

def build_g3_order():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, '{#is_minor}A minor{/is_minor}{^is_minor}An alleged incapacitated person{/is_minor}')

    _add_para(doc, 'ORDER APPOINTING {guardian_kind_caps}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{order_scope_line}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{order_subtitle}',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # Findings preamble — three branches: minor, plenary (with/without AD),
    # limited (without AD only).
    _add_para(doc,
        '{#is_minor}On the petition of {petitioner_name} for the appointment of a guardian {scope_phrase} of {aip_name}, a minor, the Court finding that it is necessary for a guardian {scope_phrase} to be appointed, it is{/is_minor}'
        '{#is_plenary}On the petition of {petitioner_name} for the appointment of a plenary guardian {scope_phrase} of {aip_name} (the Ward), the Court finds that the Ward is totally incapacitated as adjudicated by order of this Court entered {incapacity_order_date}; '
        '{^has_advance_directive}and that it is necessary for a plenary guardian to be appointed {scope_phrase} of the Ward. Further, the court finds no evidence that the Ward, prior to incapacity, executed any valid advance directive pursuant to Florida Statutes Chapter 765 and that there are no alternatives to guardianship that are sufficient to serve the needs of the Ward; and that appointment of a plenary guardian is the least restrictive alternative. It is, therefore,{/has_advance_directive}'
        '{#has_advance_directive}that there are no alternatives to guardianship that are sufficient to serve the needs of the Ward; and that appointment of a plenary guardian is the least restrictive alternative. It is, therefore,{/has_advance_directive}{/is_plenary}'
        '{#is_limited}On the petition of {petitioner_name} for the appointment of a guardian {scope_phrase} of {aip_name} (the Ward), the Court makes the following findings:{/is_limited}',
        space_after=12)

    # Limited-only findings list (numbered findings preamble).
    _add_para(doc, '{#is_limited}', space_after=0)

    _pleading_para(doc,
        '{#is_limited}The Ward was adjudicated incapacitated by Order of this Court entered on {incapacity_order_date}, and the Court, having considered alternatives to guardianship, found that no alternatives to guardianship sufficiently address the needs of the Ward, and that the restrictions imposed upon the Ward’s rights and liberties are consistent with the Ward’s welfare and safety, and are the least restrictive appropriate alternatives, reserving to the Ward the right to make decisions in all matters commensurate with the Ward’s ability to do so.{/is_limited}')

    _pleading_para(doc,
        '{#is_limited}The Order Determining Limited Incapacity established the incapacity of the Ward to exercise delegable rights as set forth in the petition.{/is_limited}')

    _pleading_para(doc,
        '{#is_limited}The nature of the guardianship is limited and it is necessary to appoint a limited guardian {scope_phrase} of the Ward.{/is_limited}')

    _pleading_para(doc,
        '{#is_limited}{^has_advance_directive}The Court finds no evidence that the Ward, prior to incapacity, executed any valid advance directive pursuant to Florida Statutes Chapter 765.{/has_advance_directive}{#has_advance_directive}The Ward, prior to incapacity, executed an advance directive pursuant to Florida Statutes Chapter 765 — see further provisions below.{/has_advance_directive}{/is_limited}')

    _add_para(doc, '{#is_limited}The Court having jurisdiction and being fully advised, it is{/is_limited}',
              space_after=12)

    _add_para(doc, '{/is_limited}', space_after=0)

    _add_para(doc, 'ADJUDGED as follows:', bold=True, space_after=12)

    # 1. Appointment.
    _pleading_para(doc,
        '{proposed_guardian_name} is qualified to serve and is hereby appointed {guardian_kind_lower} {scope_phrase} of {aip_name}{^is_minor} (the Ward){/is_minor}.')

    # 2. Bond + letters.
    _pleading_para(doc,
        'Upon taking the prescribed oath, filing a designation of resident agent and acceptance, and posting a bond in the amount of ${bond_amount} payable to the Governor of the State of Florida and to all successors in office, conditioned on the faithful performance of all duties by the guardian, letters of guardianship shall be issued.')

    # 3. Restricted account (property only).
    _pleading_para(doc,
        '{#includes_property}The guardian must place the following property of the Ward: {restricted_property}, in a restricted account in a financial institution designated pursuant to Florida Statutes section 69.031.{/includes_property}')

    # 4. Advance-directive paragraph (plenary + AD only).
    _pleading_para(doc,
        '{#is_plenary}{#has_advance_directive}The court finds that the Ward, prior to incapacity, appointed a health care surrogate pursuant to Florida Statutes Chapter 765. Such advance directive shall {advance_directive_disposition}.{/has_advance_directive}{/is_plenary}')

    # 5. Adult-only: rights retention.
    _pleading_para(doc,
        '{^is_minor}The Ward retains the rights specified in Florida Statutes section 744.3215(1) and the right to make decisions in all matters commensurate with the Ward’s abilities.{/is_minor}')

    _add_para(doc, '', space_after=18)
    _add_para(doc, 'ORDERED on this _____ day of __________, 20___.',
              first_indent=Inches(0.5), space_after=24)

    _add_para(doc, '_______________________________________', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'Circuit Judge', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    # NO AI certification — this is a judge-signed order. Hard rule.

    out_path = os.path.join(TEMPLATE_DIR, 'G3-ORDER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-LETTERS  Letters of Guardianship
# ---------------------------------------------------------------------------
#
# Replaces 14 FLSSI letters (G-3.100 through G-3.115). Branches via matter
# flags. Like Order, no AI cert (judge-signed).
# ---------------------------------------------------------------------------

def build_g3_letters():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name}')

    _add_guardianship_caption(doc, '{#is_minor}A minor{/is_minor}{^is_minor}An alleged incapacitated person{/is_minor}')

    _add_para(doc, 'LETTERS OF {letters_kind_caps}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{letters_scope_line}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'TO ALL WHOM IT MAY CONCERN:', bold=True, space_after=12)

    _add_para(doc,
        'WHEREAS {proposed_guardian_name} has been duly appointed by this Court as {guardian_kind_lower} {scope_phrase} of {aip_name}'
        '{#is_minor}, a minor,{/is_minor}'
        '{^is_minor} (the Ward){/is_minor}'
        ' and has performed all acts prerequisite to the issuance of these letters,',
        space_after=12)

    _add_para(doc,
        'NOW, THEREFORE, {proposed_guardian_name} is hereby authorized to exercise all of the powers and duties as {guardian_kind_lower} {scope_phrase} of {aip_name} that are granted by law, subject to the limitations of any order of this Court.',
        space_after=18)

    _add_para(doc, 'WITNESS my hand and the seal of this Court, on this _____ day of __________, 20___.',
              space_after=24)

    _add_para(doc, '_______________________________________', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'Clerk of the Circuit Court', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'By: _______________________________________', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'As Deputy Clerk', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    # NO AI certification — judge-/clerk-signed. Hard rule.

    out_path = os.path.join(TEMPLATE_DIR, 'G3-LETTERS.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-VOL-PETITION  Petition for Voluntary Guardianship (FLSSI G-3.035)
# ---------------------------------------------------------------------------
#
# Voluntary guardianship under §744.341 — petitioner is the (competent) Ward
# requesting that a guardian be appointed over the property. No incapacity
# proceeding required. Property scope only by statute.
# ---------------------------------------------------------------------------

def build_g3_vol_petition():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Voluntary Guardianship of {aip_name}')
    _ensure_pleading_numbering(doc)

    _add_guardianship_caption(doc, 'A voluntary ward')

    _add_para(doc, 'PETITION FOR VOLUNTARY GUARDIANSHIP',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        'Petitioner is mentally competent to understand the nature of this petition and the effect of the appointment of a guardian over Petitioner’s property.')

    _pleading_para(doc,
        'Petitioner’s residence is {petitioner_residence}, Petitioner’s post office address is {petitioner_address}, and Petitioner’s date of birth is {aip_dob}.')

    _pleading_para(doc,
        'Petitioner desires the appointment of a guardian of Petitioner’s property because of {voluntary_reason}.')

    _pleading_para(doc,
        'Petitioner has annexed to this petition a certificate of a licensed Florida physician (using FLSSI form G-3.120, or substantially similar) attesting to Petitioner’s competence to understand the nature of this petition.')

    _pleading_para(doc,
        'The proposed guardian, {proposed_guardian_name}, whose address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as guardian of the property of the Ward.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian} The relationship of the proposed guardian to Petitioner is {proposed_guardian_relationship}. The proposed guardian should be appointed because {appointment_reason}.')

    _pleading_para(doc,
        'The nature and value of the property subject to the voluntary guardianship are as follows:',
        keep_with_next=True)
    _property_items_table(doc)

    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that {proposed_guardian_name} be appointed guardian of the property of Petitioner.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Voluntary Guardianship')
    _add_miami_dade_ai_certification(doc, 'Petition for Voluntary Guardianship')

    _add_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'G3-VOL-PETITION.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-120  Physician's Certificate for Voluntary Guardianship
# ---------------------------------------------------------------------------
#
# Required attachment to a Voluntary Petition. FLSSI G-3.120.
# ---------------------------------------------------------------------------

def build_g3_120():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Voluntary Guardianship of {aip_name}')

    _add_guardianship_caption(doc, 'A voluntary ward')

    _add_para(doc, 'PHYSICIAN’S CERTIFICATE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '(Voluntary Guardianship)',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    _add_para(doc,
        'I, {physician_name}, a Florida-licensed physician (license number {physician_license_no}), state under oath that:',
        space_after=12)

    _pleading_para(doc,
        'I have examined {petitioner_name} on {physician_exam_date}.')

    _pleading_para(doc,
        'In my opinion, {petitioner_name} is mentally competent to understand the nature of the Petition for Voluntary Guardianship and the effect of the appointment of a guardian over the petitioner’s property.')

    indent = Inches(0.5)
    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=indent, space_before=18, space_after=24)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{physician_name}, M.D.', space_after=0)
    _add_para(doc, 'Florida License No. {physician_license_no}', space_after=0)
    _add_para(doc, '{physician_address}', space_after=18)

    _add_broward_ai_certification(doc, 'Physician’s Certificate (Voluntary Guardianship)')
    _add_miami_dade_ai_certification(doc, 'Physician’s Certificate (Voluntary Guardianship)')

    out_path = os.path.join(TEMPLATE_DIR, 'G3-120.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-EMERGENCY-ORDER  Order Appointing Emergency Temporary Guardian
# ---------------------------------------------------------------------------
#
# Replaces FLSSI G-3.060 (no advance directive) + G-3.060a1 (advance
# directive). Branches on has_advance_directive. Judge-signed -> no AI cert.
# ---------------------------------------------------------------------------

def build_g3_emergency_order():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name} — Emergency')

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    _add_para(doc, 'ORDER APPOINTING EMERGENCY',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'TEMPORARY GUARDIAN',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '{#has_advance_directive}(advance directive){/has_advance_directive}',
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc,
        'On the petition of {petitioner_name} for appointment of an emergency '
        'temporary guardian for {aip_name}, an alleged incapacitated person, '
        'who is represented by counsel in these proceedings, and it appearing '
        'to the court that there is an imminent danger that the physical or '
        'mental health or safety of the alleged incapacitated person will be '
        'seriously impaired or that the property of that person is in danger '
        'of being wasted, misappropriated or lost unless immediate action is '
        'taken; and the court having jurisdiction and being fully advised; it is',
        space_after=12)

    _add_para(doc, 'ADJUDGED as follows:', bold=True, space_after=12)

    _ensure_pleading_numbering(doc)

    _pleading_para(doc,
        '{proposed_guardian_name} is qualified to serve and is hereby '
        'appointed as emergency temporary guardian {scope_phrase} of '
        '{aip_name} (the Ward).')

    _pleading_para(doc,
        'Upon taking the prescribed oath, filing a designation of resident '
        'agent and acceptance, and posting bond in the amount of '
        '${bond_amount} payable to the Governor of the State of Florida and '
        'all successors in office, conditioned on the faithful performance '
        'of all duties by the guardian, letters of emergency temporary '
        'guardianship shall be issued to the emergency temporary guardian '
        'granting the following powers and duties: {emergency_powers_duties}.')

    _pleading_para(doc,
        '{^has_advance_directive}The Court is not aware whether the Ward has '
        'executed any valid advance directive pursuant to Chapter 765, Florida '
        'Statutes. If any such advance directive exists, the guardian shall '
        'exercise no authority to make health care decisions until further '
        'order of this Court.{/has_advance_directive}'
        '{#has_advance_directive}The Court is aware that the Ward has '
        'executed an advance directive pursuant to Florida Statutes Chapter '
        '765, but is unaware if it is valid and effective. The guardian shall '
        'not exercise authority to make health care decisions until further '
        'order of this Court.{/has_advance_directive}')

    _pleading_para(doc,
        'Unless further extended by order of this Court, the authority of '
        'the emergency temporary guardian will expire ninety (90) days after '
        'the date of this order, or when a guardian is appointed pursuant to '
        'Florida Statutes section 744.2005, whichever occurs first.')

    _add_para(doc, '', space_after=18)
    _add_para(doc, 'ORDERED on this _____ day of __________, 20___.',
              first_indent=Inches(0.5), space_after=24)

    _add_para(doc, '_______________________________________', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'Circuit Judge', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    # NO AI certification — judge-signed.

    out_path = os.path.join(TEMPLATE_DIR, 'G3-EMERGENCY-ORDER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# G3-EMERGENCY-LETTERS  Letters of Emergency Temporary Guardianship
# ---------------------------------------------------------------------------
#
# Replaces FLSSI G-3.110. Judge-/clerk-signed -> no AI cert.
# ---------------------------------------------------------------------------

def build_g3_emergency_letters():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Guardianship of {aip_name} — ETG Letters')

    _add_guardianship_caption(doc, 'An alleged incapacitated person')

    _add_para(doc, 'LETTERS OF EMERGENCY',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'TEMPORARY GUARDIANSHIP',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'TO ALL WHOM IT MAY CONCERN:', bold=True, space_after=12)

    _add_para(doc,
        'WHEREAS, {proposed_guardian_name} has been appointed emergency '
        'temporary guardian {scope_phrase} of {aip_name} (the Ward), and has '
        'taken the prescribed oath and performed all other acts prerequisite '
        'to issuance of letters of emergency temporary guardianship of the '
        'Ward,',
        space_after=12)

    _add_para(doc,
        'NOW THEREFORE, I, the undersigned circuit judge, declare '
        '{proposed_guardian_name} duly qualified under the laws of the State '
        'of Florida to act as emergency temporary guardian {scope_phrase} of '
        '{aip_name}, with full power to exercise the following powers and '
        'duties: {emergency_powers_duties}.',
        space_after=12)

    _add_para(doc,
        'The guardian '
        '{#emergency_health_care_authority}shall{/emergency_health_care_authority}'
        '{^emergency_health_care_authority}shall not{/emergency_health_care_authority} '
        'have authority to make health care decisions until further order of '
        'this Court.',
        space_after=12)

    _add_para(doc,
        'The authority of the emergency temporary guardian expires ninety '
        '(90) days after the date hereof, unless earlier terminated by the '
        'appointment of a guardian or extended by order of this court.',
        space_after=24)

    _add_para(doc, 'WITNESS my hand and the seal of this Court, on this _____ day of __________, 20___.',
              space_after=24)

    _add_para(doc, '_______________________________________', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)
    _add_para(doc, 'Circuit Judge', align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=0)

    # NO AI certification — judge-signed.

    out_path = os.path.join(TEMPLATE_DIR, 'G3-EMERGENCY-LETTERS.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    build_g2_010()
    build_g2_140()
    build_g3_010()
    build_g3_025()
    build_g3_026()
    build_g3_petition()
    build_g3_emergency()
    build_g3_oath()
    build_g3_order()
    build_g3_letters()
    build_g3_vol_petition()
    build_g3_120()
    # Phase 11 (F): emergency Order + Letters
    build_g3_emergency_order()
    build_g3_emergency_letters()
