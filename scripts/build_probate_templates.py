#!/usr/bin/env python3
"""
Build clean single-column probate templates for Ginsberg Shulman.

Four "smart" templates cover every variant of formal-administration opening
via docxtemplater conditionals — FLSSI form numbers are meaningless to the
firm, so this collapses 23 FLSSI IDs into 4 maintainable templates:

  P3-PETITION  Petition for Administration
    Axes: testate/intestate, single/multiple petitioners, domiciliary/ancillary
  P3-OATH      Oath of Personal Representative + Designation of Resident Agent
    Axes: single/multiple PRs
  P3-ORDER     Order Admitting Will and Appointing PR / Order Appointing PR
    Axes: testate/intestate, single/multiple PRs, domiciliary/ancillary
  P3-LETTERS   Letters of Administration / Letters of Ancillary Administration
    Axes: single/multiple PRs, domiciliary/ancillary

Shared conventions (all 4 templates):
  - Single-column caption via borderless table
  - Real Word numbering (numPr/numId=1) per docx-numbering skill
  - Grammar pre-computed in prepareTemplateData() — templates use
    {petitioner_label}/{petitioner_verb_alleges}/{pr_names}/etc.
  - Broward AI certification above signature block (county-conditional)

Shared helpers (_add_para, _pleading_para, _inject_numbering_part, etc.)
are imported from build_guardianship_templates.
"""

import os

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from build_guardianship_templates import (
    FONT, BODY_SIZE, SMALL_SIZE, TEMPLATE_DIR,
    _add_para, _add_run, _set_run,
    _borderless_table, _table_with_borders, _clear_cell, _cell_para,
    _apply_page_setup, _apply_running_header,
    _ensure_pleading_numbering, _pleading_para,
    _inject_numbering_part,
    _add_broward_ai_certification,
    _add_miami_dade_ai_certification,
)


# ---------------------------------------------------------------------------
# Probate-specific helpers
# ---------------------------------------------------------------------------

def _add_probate_caption(doc, *, decedent_tag='{decedent_full_name}'):
    """Probate caption:
      IN THE CIRCUIT COURT FOR {county_caption} COUNTY, FLORIDA (centered, bold)
      4-row borderless table (decedent on the left, File No./Division on the right),
      every line bold.

    {county_caption} is set by prepareTemplateData() in app.js to an uppercased
    version of the matter county, so e.g. "Broward" renders as "BROWARD".
    """
    _add_para(doc, 'IN THE CIRCUIT COURT FOR {county_caption} COUNTY, FLORIDA',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    caption = _borderless_table(doc, rows=4, cols=2, col_widths_in=[3.5, 3.0])
    _clear_cell(caption.cell(0, 0)); _cell_para(caption.cell(0, 0), 'IN RE: ESTATE OF', bold=True, space_after=0)
    _clear_cell(caption.cell(0, 1)); _cell_para(caption.cell(0, 1), 'PROBATE DIVISION', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 0)); _cell_para(caption.cell(1, 0), f'{decedent_tag},', bold=True, space_after=0)
    _clear_cell(caption.cell(1, 1)); _cell_para(caption.cell(1, 1), '', bold=True, space_after=0)
    _clear_cell(caption.cell(2, 0)); _cell_para(caption.cell(2, 0), 'Deceased.', bold=True, space_after=0)
    _clear_cell(caption.cell(2, 1)); _cell_para(caption.cell(2, 1), 'File No. {file_no}', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 0)); _cell_para(caption.cell(3, 0), '', bold=True, space_after=0)
    _clear_cell(caption.cell(3, 1)); _cell_para(caption.cell(3, 1), 'Division {division}', bold=True, space_after=0)
    _add_para(doc, '', space_after=18)


def _estate_assets_table(doc):
    """Two-column bordered table for probate estate assets with repeating row."""
    tbl = _table_with_borders(doc, rows=2, cols=2, col_widths_in=[4.5, 2.0])
    for cell in tbl.rows[0].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(0, 0), 'NATURE OF ASSET', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 1), 'APPROXIMATE VALUE', bold=True, space_after=0)
    for cell in tbl.rows[1].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(1, 0), '{#estate_assets}{asset_description}', space_after=0)
    _cell_para(tbl.cell(1, 1), '{asset_value_formatted}{/estate_assets}', space_after=0)
    return tbl


def _distribution_table(doc):
    """Three-column bordered table for summary-admin distribution plan.
    Same row pattern as beneficiaries: header + one looping row
    (`{#summary_admin_distributees}...{/summary_admin_distributees}`)."""
    tbl = _table_with_borders(doc, rows=2, cols=3, col_widths_in=[2.0, 2.8, 1.7])
    for cell in tbl.rows[0].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(0, 0), 'NAME', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 1), 'ADDRESS', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 2), 'ASSET, SHARE OR AMOUNT', bold=True, space_after=0)
    for cell in tbl.rows[1].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(1, 0), '{#summary_admin_distributees}{dist_name}', space_after=0)
    _cell_para(tbl.cell(1, 1), '{dist_address}', space_after=0)
    _cell_para(tbl.cell(1, 2), '{dist_share}{/summary_admin_distributees}', space_after=0)
    return tbl


def _beneficiaries_table(doc):
    """Four-column bordered table for probate beneficiaries with a
    docxtemplater paragraph-loop row."""
    tbl = _table_with_borders(doc, rows=2, cols=4, col_widths_in=[1.8, 2.4, 1.4, 0.9])
    for cell in tbl.rows[0].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(0, 0), 'NAME', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 1), 'ADDRESS', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 2), 'RELATIONSHIP', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 3), 'BIRTH YR.', bold=True, space_after=0)
    for cell in tbl.rows[1].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(1, 0), '{#beneficiaries}{ben_name}', space_after=0)
    _cell_para(tbl.cell(1, 1), '{ben_address}', space_after=0)
    _cell_para(tbl.cell(1, 2), '{ben_relationship}', space_after=0)
    _cell_para(tbl.cell(1, 3), '{ben_year_of_birth}{/beneficiaries}', space_after=0)
    return tbl


def _add_probate_signature_block(doc):
    """Signature block: one signer line per petitioner (via loop) + attorney.

    Signing date left blank — filler at signing time (won't be known at draft).
    """
    indent = Inches(0.5)
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)

    # Per-petitioner signature line (loops over petitioners array; renders a
    # single sig line when there's only one petitioner).
    _add_para(doc, '{#petitioners}', space_after=0)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{pet_name}, Petitioner', space_after=24)
    _add_para(doc, '{/petitioners}', space_after=0)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)


def _add_order_signature_block(doc):
    """Signature block for court orders: date stamp + judge signature."""
    _add_para(doc, '', space_after=18)
    _add_para(doc, 'DONE AND ORDERED in {county} County, Florida.', space_after=36)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'CIRCUIT JUDGE', space_after=24)
    _add_para(doc, 'Copies furnished to:', space_after=0)
    _add_para(doc, '{attorney_name}, {attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=0)


# ---------------------------------------------------------------------------
# P3-PETITION  Petition for Administration (smart / conditional)
# ---------------------------------------------------------------------------

def build_p3_petition():
    """One template covering every opening-petition variant via conditionals:
      {is_testate}        — testate vs intestate
      {is_ancillary}      — domiciliary vs ancillary
      {multiple_petitioners} / {multiple_prs} — grammar + loop branches
    Grammar strings ({petitioner_label}, {petitioner_verb_alleges}, etc.)
    are precomputed in prepareTemplateData().
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    # ---- Title (conditional on is_ancillary) ----
    _add_para(doc, 'PETITION FOR{#is_ancillary} ANCILLARY{/is_ancillary} ADMINISTRATION',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Intro — unnumbered. Grammar pre-computed.
    _add_para(doc, '{petitioner_label}, {petitioner_names}, {petitioner_verb_alleges}:',
              space_after=12)

    # 1. Interest + address (loops petitioners; grammar switches on multiple_petitioners)
    _pleading_para(doc,
        '{^multiple_petitioners}{#petitioners}Petitioner has an interest in the above '
        'estate as {pet_interest}. Petitioner\u2019s address is {pet_address}, and the '
        'name and office address of petitioner\u2019s attorney are set forth at the end of '
        'this petition.{/petitioners}{/multiple_petitioners}'
        '{#multiple_petitioners}Each petitioner has an interest in the above estate, as '
        'set forth below: {#petitioners}{pet_name} is {pet_interest}, whose address is '
        '{pet_address}. {/petitioners}The name and office address of petitioners\u2019 '
        'attorney are set forth at the end of this petition.{/multiple_petitioners}')

    # 2. Decedent (ancillary vs domiciliary + testate vs intestate)
    _pleading_para(doc,
        'Decedent, {decedent_full_name}, whose last known address was {decedent_address}, '
        'and the last four digits of whose social security number are {decedent_ssn_last4}, '
        'died on {decedent_death_date} at {decedent_death_place}. On the date of death, '
        'decedent was domiciled in '
        '{^is_ancillary}{decedent_domicile} County, Florida{/is_ancillary}'
        '{#is_ancillary}{decedent_domicile_state} and left property in {county} County, '
        'Florida{/is_ancillary}'
        '{^is_testate}, and died intestate{/is_testate}.')

    # 3. Beneficiaries
    _pleading_para(doc,
        'So far as is known, the names of the beneficiaries of this estate and of the '
        'decedent\u2019s surviving spouse, if any, their addresses and relationships to '
        'decedent, and the years of birth of any who are minors, are:',
        keep_with_next=True)
    _beneficiaries_table(doc)

    # 4. Venue
    _pleading_para(doc,
        'Venue of this proceeding is in this county because {venue_reason}.')

    # 5. PR qualifications (single PR vs multiple PRs; FL-resident conditional on each)
    _pleading_para(doc,
        '{^multiple_prs}{pr_names}, whose address is {pr_address}, is qualified to serve '
        'as personal representative of the decedent\u2019s estate: {pr_names} has not been '
        'convicted of a felony, is mentally and physically able to perform the duties of '
        'personal representative, and is 18 years of age or older. '
        '{#pr_is_fl_resident}{pr_names} is a resident of Florida.{/pr_is_fl_resident}'
        '{^pr_is_fl_resident}{pr_names} is not a resident of Florida but is related to '
        'the decedent as {pr_relationship} and is qualified to serve as personal '
        'representative under Florida Statutes section 733.304.{/pr_is_fl_resident}'
        '{/multiple_prs}'
        '{#multiple_prs}Each proposed personal representative is qualified to serve: '
        '{#prs}{pr_name}, whose address is {pr_address}, has not been convicted of a '
        'felony, is mentally and physically able to perform the duties of personal '
        'representative, and is 18 years of age or older. '
        '{#pr_is_fl_resident}{pr_name} is a resident of Florida.{/pr_is_fl_resident}'
        '{^pr_is_fl_resident}{pr_name} is not a resident of Florida but is related to the '
        'decedent as {pr_relationship} and is qualified under Florida Statutes section '
        '733.304.{/pr_is_fl_resident} {/prs}{/multiple_prs}')

    # 6. Petitioner(s) prior conviction
    _pleading_para(doc,
        '{petitioner_label} {#petitioner_has_prior_conviction}{petitioner_verb_has}'
        '{/petitioner_has_prior_conviction}{^petitioner_has_prior_conviction}'
        '{petitioner_verb_has} not{/petitioner_has_prior_conviction} been convicted in any '
        'state or foreign jurisdiction of abuse, neglect, or exploitation of an elderly '
        'person or a disabled adult, as those terms are defined in Florida Statutes '
        'section 825.101.')

    # 7. Preference
    _pleading_para(doc,
        '{^higher_preference_exists}No person has equal or higher preference to be '
        'appointed personal representative.{/higher_preference_exists}'
        '{#higher_preference_exists}The following person(s) has/have equal or higher '
        'preference to be appointed personal representative: {higher_preference_names}. '
        '{#higher_preference_formal_notice}Formal notice of this petition will be served '
        'on such person(s).{/higher_preference_formal_notice}'
        '{^higher_preference_formal_notice}Formal notice of this petition will not be '
        'served on such person(s).{/higher_preference_formal_notice}'
        '{/higher_preference_exists}')

    # 8. Assets
    _pleading_para(doc,
        'The nature and approximate value of the assets in this estate are:',
        keep_with_next=True)
    _estate_assets_table(doc)

    # 9. Estate tax return
    _pleading_para(doc,
        'This estate {#estate_tax_return_required}will{/estate_tax_return_required}'
        '{^estate_tax_return_required}will not{/estate_tax_return_required} be required '
        'to file a federal estate tax return.')

    # 10. Domiciliary/principal proceedings — ancillary ALWAYS has pending domiciliary
    _pleading_para(doc,
        '{#is_ancillary}Domiciliary proceedings are pending in {domiciliary_court_name}, '
        'the address of which is {domiciliary_court_address}, and letters have been '
        'issued to {domiciliary_representative}, whose address is '
        '{domiciliary_representative_address}.{/is_ancillary}'
        '{^is_ancillary}Domiciliary or principal proceedings '
        '{#domiciliary_proceedings_pending}are known to be pending in '
        '{domiciliary_court_name}, the address of which is {domiciliary_court_address}. '
        'Letters have been issued to {domiciliary_representative}, whose address is '
        '{domiciliary_representative_address}.{/domiciliary_proceedings_pending}'
        '{^domiciliary_proceedings_pending}are not known to be pending in another state '
        'or country.{/domiciliary_proceedings_pending}{/is_ancillary}')

    # 11. Will disposition (testate) / Unaware of unrevoked wills (intestate) — ONE paragraph.
    # Will status (original in court, accompanies petition, authenticated copy, etc.) is a
    # filing-time fact — omitted here. Clerk-review language defaults to neutral "accompanies
    # this petition"; revise at filing if the original was previously deposited.
    _pleading_para(doc,
        '{#is_testate}The decedent\u2019s last will dated {will_date}'
        '{#has_codicil}, and codicil(s) dated {codicil_dates}{/has_codicil}, '
        'accompanies this petition. Petitioner is unaware of any unrevoked will or codicil '
        'of decedent other than as set forth above.{/is_testate}'
        '{^is_testate}After the exercise of reasonable diligence, petitioner is unaware '
        'of any unrevoked wills or codicils of decedent.{/is_testate}')

    # Closing — unnumbered.
    indent = Inches(0.5)
    _add_para(doc,
        'WHEREFORE, {petitioner_label} respectfully '
        '{^multiple_petitioners}requests{/multiple_petitioners}'
        '{#multiple_petitioners}request{/multiple_petitioners} that '
        '{#is_testate}the decedent\u2019s will be admitted to probate and that {/is_testate}'
        '{pr_names} be appointed {pr_label} of the estate of the decedent'
        '{#is_ancillary} and that letters of ancillary administration be issued'
        '{/is_ancillary}.',
        first_indent=indent, space_before=12, space_after=12)
    _add_para(doc,
        'Under penalties of perjury, {petitioner_label} declare{^multiple_petitioners}s'
        '{/multiple_petitioners} that {petitioner_label} {petitioner_verb_has} read the '
        'foregoing, and the facts alleged are true, to the best of '
        '{petitioner_poss} knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Administration')
    _add_miami_dade_ai_certification(doc, 'Petition for Administration')

    _add_probate_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-PETITION.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-OATH  Oath of Personal Representative + Designation of Resident Agent
# ---------------------------------------------------------------------------

def build_p3_oath():
    """Combined oath + designation of resident agent. Single document per
    PR; when there are multiple PRs, each one signs their own block below.
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'OATH OF {pr_label_caps}',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    _add_para(doc, 'AND DESIGNATION OF RESIDENT AGENT AND ACCEPTANCE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Oath section — one body paragraph per PR via loop
    _add_para(doc, 'OATH', bold=True, space_after=6)
    _add_para(doc,
        '{#prs}I, {pr_name}, swear or affirm that I will faithfully administer the '
        'estate of the above-named decedent according to law.{/prs}',
        space_after=24)

    # Per-PR oath signature block
    _add_para(doc, '{#prs}', space_after=0)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{pr_name}', space_after=18)
    _add_para(doc, 'Sworn to (or affirmed) and subscribed before me by means of '
                   '\u2610 online notarization or \u2610 physical presence this '
                   '_____ day of ______________________, 20____, by {pr_name}, '
                   'who is personally known to me or produced '
                   '______________________________ as identification.', space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Notary Public', space_after=0)
    _add_para(doc, 'State of Florida', space_after=24)
    _add_para(doc, '{/prs}', space_after=0)

    # Designation of Resident Agent
    _add_para(doc, 'DESIGNATION OF RESIDENT AGENT AND ACCEPTANCE',
              bold=True, space_before=12, space_after=12)
    _add_para(doc,
        '{pr_label_title} hereby designate{^multiple_prs}s{/multiple_prs} {resident_agent_name}, '
        'whose address is {resident_agent_address}, as resident agent '
        'to accept service of process within the State of Florida in any '
        'action or proceeding affecting the estate of the above-named decedent.',
        space_after=18)

    # PR signature(s) for the designation
    _add_para(doc, '{#prs}', space_after=0)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{pr_name}, {pr_label_title}', space_after=18)
    _add_para(doc, '{/prs}', space_after=6)

    # Resident agent acceptance
    _add_para(doc, 'ACCEPTANCE', bold=True, space_before=12, space_after=12)
    _add_para(doc,
        'I, {resident_agent_name}, having a business address within the State of Florida, '
        'hereby accept the foregoing designation as resident agent.',
        space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{resident_agent_name}', space_after=0)
    _add_para(doc, 'Resident Agent', space_after=18)

    _add_broward_ai_certification(doc, 'Oath of Personal Representative')
    _add_miami_dade_ai_certification(doc, 'Oath of Personal Representative')

    out_path = os.path.join(TEMPLATE_DIR, 'P3-OATH.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-ORDER  Order Admitting Will & Appointing PR / Order Appointing PR
# ---------------------------------------------------------------------------

def build_p3_order():
    """Proposed order granting the petition. Handles:
      - testate: admits will, appoints PR
      - intestate: appoints PR only
      - single/multiple PRs
      - domiciliary (Letters of Administration) vs ancillary (Letters of
        Ancillary Administration)
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    # Title varies on testate + ancillary
    _add_para(doc,
        'ORDER '
        '{#is_testate}ADMITTING WILL TO PROBATE AND {/is_testate}'
        'APPOINTING {pr_label_caps}'
        '{#is_ancillary} OF ANCILLARY ADMINISTRATION{/is_ancillary}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Findings preamble
    _add_para(doc,
        'On the petition of {petitioner_names} for administration of the estate of '
        '{decedent_full_name}, deceased, the Court finds that the decedent died on '
        '{decedent_death_date}; that the decedent was domiciled in '
        '{^is_ancillary}{decedent_domicile} County, Florida{/is_ancillary}'
        '{#is_ancillary}{decedent_domicile_state}{/is_ancillary} at the time of death; '
        'that {pr_names} {pr_verb_is} entitled to preference in appointment and '
        '{pr_verb_is} qualified to serve as {pr_label}; and that the petition '
        'satisfies the requirements of Florida Probate Code. It is',
        space_after=18)

    _add_para(doc, 'ORDERED and ADJUDGED:', bold=True, space_after=12)

    # Testate-only: admit will
    _pleading_para(doc,
        '{#is_testate}The decedent\u2019s last will dated {will_date}'
        '{#has_codicil}, and codicil(s) dated {codicil_dates}{/has_codicil} '
        '{#will_is_self_proved}{pr_verb_is} self-proved and {/will_is_self_proved}'
        '{pr_verb_is} admitted to probate according to law.{/is_testate}'
        '{^is_testate}The decedent died intestate. Administration of the '
        'decedent\u2019s estate is granted.{/is_testate}')

    # Appoint PR(s). Bond omitted from the questionnaire — if the court wants bond,
    # it will enter a separate order requiring it.
    _pleading_para(doc,
        '{pr_names} {pr_verb_is} appointed as {pr_label} of the estate of the decedent, '
        'to serve without bond.')

    # Letters
    _pleading_para(doc,
        'Upon {pr_pronoun_his_her} qualification by filing of {pr_pronoun_his_her} oath, '
        'letters of {#is_ancillary}ancillary {/is_ancillary}administration shall issue to '
        '{pr_names}.')

    _add_order_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-ORDER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-LETTERS  Letters of Administration / Letters of Ancillary Administration
# ---------------------------------------------------------------------------

def build_p3_letters():
    """Letters issued by the clerk once the order is entered. Single template
    covers domiciliary + ancillary, single + multiple PRs."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc,
        'LETTERS OF {#is_ancillary}ANCILLARY {/is_ancillary}ADMINISTRATION',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'TO ALL WHOM IT MAY CONCERN:', bold=True, space_after=12)

    _add_para(doc,
        'WHEREAS, {decedent_full_name}, a resident of '
        '{^is_ancillary}{decedent_domicile} County, Florida{/is_ancillary}'
        '{#is_ancillary}{decedent_domicile_state}{/is_ancillary}, died on '
        '{decedent_death_date}, '
        '{#is_testate}owning assets in this state, and a duly executed last will'
        '{#has_codicil} and codicil(s) dated {codicil_dates}{/has_codicil} of the decedent has been '
        'admitted to probate in this Court; and{/is_testate}'
        '{^is_testate}leaving assets in this state and having died intestate; and'
        '{/is_testate}',
        space_after=12)

    _add_para(doc,
        'WHEREAS, {pr_names} {pr_verb_is} qualified under the laws of the State of '
        'Florida to act as {pr_label}'
        '{#is_ancillary} of ancillary administration{/is_ancillary} of the estate of '
        'the decedent, and has/have taken the prescribed oath;',
        space_after=12)

    _add_para(doc,
        'NOW, THEREFORE, I, the undersigned Circuit Judge, do declare {pr_names} '
        'duly qualified under the laws of the State of Florida to act as {pr_label}'
        '{#is_ancillary} of ancillary administration{/is_ancillary} of the estate of '
        '{decedent_full_name}, deceased, with full power to administer the estate '
        'according to law; to ask, demand, sue for, recover and receive the property '
        'of the decedent; to pay the debts of the decedent as far as the assets of '
        'the estate will permit and the law directs; and to make distribution of the '
        'estate according to law.',
        space_after=24)

    _add_para(doc, 'WITNESS my hand and the seal of this Court.',
              space_after=24)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'CIRCUIT JUDGE', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-LETTERS.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-0900  Notice of Designation of Email Addresses for Service
# ---------------------------------------------------------------------------

def build_p1_0900():
    """Rule 2.516(b)(1)(A) notice of primary/secondary service emails.
    Attorney-signed only (no petitioner signature). Signing date left blank
    (filled at signing).
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE OF DESIGNATION OF EMAIL ADDRESSES',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'FOR SERVICE OF DOCUMENTS',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'Pursuant to Florida Rule of General Practice and Judicial Administration '
        '2.516(b)(1)(A), the undersigned counsel gives notice of the following '
        'primary and secondary e-mail addresses for service in this matter:',
        first_indent=indent, space_after=12)

    _add_para(doc, 'Primary Email Address: {attorney_email}',
              first_indent=indent, space_after=6)
    _add_para(doc,
        'Secondary Email Address: '
        '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}'
        '{^attorney_email_secondary}N/A{/attorney_email_secondary}',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Notice of Designation of Email Addresses')
    _add_miami_dade_ai_certification(doc, 'Notice of Designation of Email Addresses')

    # Attorney-only signature block (no petitioner sig on a service notice).
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0900.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# Shared helpers for P1 service / consent forms (Batch 1)
# ---------------------------------------------------------------------------

def _add_attorney_signature_block(doc):
    """Attorney-only signature block for P1 service / consent / notice forms.
    Signing date left blank (handwritten at signing per Phase 7b convention).
    """
    indent = Inches(0.5)
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)


# ---------------------------------------------------------------------------
# P1-0400  Request for Notice and Copies
# ---------------------------------------------------------------------------

def build_p1_0400():
    """Florida Probate Rule 5.060 — request by an interested person for notice
    and copies of subsequent pleadings/documents in the estate.
    Signed by requestor + attorney; cert of service to PR's attorney.
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'REQUEST FOR NOTICE AND COPIES',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, '{requestor_name}, alleges:', space_after=12)

    _pleading_para(doc, 'I have an interest in the above estate as {requestor_interest}.')
    _pleading_para(doc, 'My residence is {requestor_residence}.')
    _pleading_para(doc, 'My post office address is {requestor_mailing_address}.')
    _pleading_para(doc, 'The name and address of my attorney, if any, are set forth below.')
    _pleading_para(doc, 'A copy of this request is being served on the attorney for the personal representative.')

    indent = Inches(0.5)
    _add_para(doc,
        'I request that, as long as I am an interested person, notice of all '
        'further proceedings in this estate and copies of subsequent pleadings '
        'and documents be sent to:',
        first_indent=indent, space_after=6)
    _add_para(doc, '{service_address}', first_indent=indent, space_after=12)

    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, '
        'and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Request for Notice and Copies')
    _add_miami_dade_ai_certification(doc, 'Request for Notice and Copies')

    # Requestor signature, then attorney signature.
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{requestor_name}, Requesting Party', space_after=24)

    _add_attorney_signature_block(doc)

    # Certificate of Service
    _add_para(doc, '', space_after=12)
    _add_para(doc, 'CERTIFICATE OF SERVICE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    _add_para(doc,
        'I CERTIFY that a copy hereof has been furnished to {cos_recipients} '
        'by {cos_method} on _____ day of ______________________, 20____.',
        first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0400.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# Shared: adversary caption supplement
# ---------------------------------------------------------------------------

def _add_adversary_caption_supplement(doc):
    """Adds the adversary-proceeding caption block ABOVE _add_probate_caption().

    Wraps each line in {#is_adversary}...{/is_adversary} marker paragraphs so
    the entire block disappears in non-adversary mode (no trailing whitespace).
    """
    _add_para(doc, '{#is_adversary}', space_after=0)
    _add_para(doc, 'Adversary Proceeding No. {adversary_proceeding_no}',
              align=WD_ALIGN_PARAGRAPH.RIGHT, space_after=12)
    _add_para(doc, '{petitioner_name},', space_after=0)
    indent = Inches(0.5)
    _add_para(doc, 'Petitioner,', first_indent=indent, space_after=0)
    _add_para(doc, 'vs.', space_after=0)
    _add_para(doc, '{respondent_name},', space_after=0)
    _add_para(doc, 'Respondent.', first_indent=indent, space_after=18)
    _add_para(doc, '{/is_adversary}', space_after=0)


# ---------------------------------------------------------------------------
# P1-FORMAL-NOTICE  Smart Formal Notice (consolidates P1-0500/0501)
#   Axis: is_adversary
# ---------------------------------------------------------------------------

def build_p1_formal_notice():
    """Florida Probate Rule 5.040(a) formal notice. Served by attorney with
    a copy of the underlying petition; recipient has 20 days to respond.

    Adversary variant adds:
      - "Adversary Proceeding No. ___" line above caption
      - Petitioner / vs. / Respondent block
      - "(adversary)" subtitle on title
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)
    _add_adversary_caption_supplement(doc)

    _add_para(doc, 'FORMAL NOTICE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{#is_adversary}(adversary){/is_adversary}',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'TO:', space_after=6)
    _add_para(doc, '{notice_to}', first_indent=indent, space_after=18)

    _add_para(doc,
        'YOU ARE NOTIFIED that a {petition_being_noticed} has been filed in '
        'this court, a copy of which accompanies this notice. You are required '
        'to serve written defenses on the undersigned within 20 days after '
        'service of this notice, exclusive of the day of service, and to file '
        'the original of the written defenses with the clerk of the above court '
        'either before service or immediately thereafter. Failure to serve and '
        'file written defenses as required may result in a judgment or order '
        'for the relief demanded in the pleading or motion, without further '
        'notice.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Formal Notice')
    _add_miami_dade_ai_certification(doc, 'Formal Notice')

    _add_attorney_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-FORMAL-NOTICE.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-PROOF-OF-SERVICE-FN  Smart Proof of Service of Formal Notice
#   Consolidates FLSSI P1-0507/0510/0511/0512/0513.
#   Axes: is_adversary × service_type (certified / first_class / in_manner_of)
# ---------------------------------------------------------------------------

def build_p1_proof_of_service_fn():
    """Sworn proof of service.
      service_type=formal_notice_certified  → P1-0510 / P1-0511
      service_type=formal_notice_first_class → P1-0507 (no adversary version
                                                in FLSSI; we support is_adversary
                                                for completeness)
      service_type=in_the_manner_of          → P1-0512 / P1-0513
    """
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)
    _add_adversary_caption_supplement(doc)

    # Title — three branches by service_type
    _add_para(doc,
        '{#service_type_certified}PROOF OF SERVICE OF FORMAL NOTICE{/service_type_certified}'
        '{#service_type_first_class}PROOF OF SERVICE OF FORMAL NOTICE BY FIRST CLASS MAIL{/service_type_first_class}'
        '{#service_type_in_manner_of}PROOF OF SERVICE IN THE MANNER OF FORMAL NOTICE{/service_type_in_manner_of}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{#is_adversary}(adversary){/is_adversary}',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)

    # Body — branches on service_type
    _add_para(doc,
        'Under penalties of perjury, I swear or affirm that on {service_date}, '
        'a copy of {pleading_title} '
        '{#service_type_in_manner_of}filed in the above proceeding was mailed by '
        'United States registered or certified mail, return receipt requested, '
        'postage prepaid, or was delivered in a manner permitted by Florida '
        'Probate Rule 5.040(a), to:'
        '{/service_type_in_manner_of}'
        '{#service_type_certified}and a copy of the formal notice thereof filed '
        'in the above proceeding were mailed by United States registered or '
        'certified mail, return receipt requested, postage prepaid, or were '
        'delivered in a manner permitted by Florida Probate Rule 5.040(a), to:'
        '{/service_type_certified}'
        '{#service_type_first_class}and a copy of the formal notice thereof '
        'filed in the above proceeding were mailed by United States First Class '
        'Mail, to:'
        '{/service_type_first_class}',
        first_indent=indent, space_after=12)

    _add_para(doc, '{service_recipients}', first_indent=indent, space_after=18)

    # First-class only: Florida Probate Rule 5.040(a)(3)(D) justification
    _add_para(doc, '{#service_type_first_class}', space_after=0)
    _add_para(doc,
        'Service by First Class Mail is appropriate under Florida Probate Rule '
        '5.040(a)(3)(D), because:',
        first_indent=indent, space_after=6)
    _add_para(doc,
        '(i) registered or certified mail service to the addressee requiring a '
        'signed receipt is unavailable and delivery by commercial delivery '
        'service requiring a signed receipt was also unavailable;',
        first_indent=indent, space_after=6)
    _add_para(doc,
        '(ii) delivery pursuant to subdivision (a)(3)(A) was attempted and was '
        'refused by the addressee; or',
        first_indent=indent, space_after=6)
    _add_para(doc,
        '(iii) delivery pursuant to subdivision (a)(3)(A) was attempted and was '
        'unclaimed after notice to the addressee by the delivering entity.',
        first_indent=indent, space_after=18)
    _add_para(doc, '{/service_type_first_class}', space_after=0)

    # Certified / in-the-manner-of: receipts attached language
    _add_para(doc,
        '{^service_type_first_class}Signed receipts or other evidence that '
        "service was made on each addressee or the addressee's agent are "
        'attached.{/service_type_first_class}',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Proof of Service of Formal Notice')
    _add_miami_dade_ai_certification(doc, 'Proof of Service of Formal Notice')

    _add_attorney_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-PROOF-OF-SERVICE-FN.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-0530  Notice of Hearing
# ---------------------------------------------------------------------------

def build_p1_0530():
    """Notice of Hearing per Florida Probate Rule 5.060 / 2.516.
    Signed by attorney; includes ADA accommodation language."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE OF HEARING',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'TO:', space_after=6)
    _add_para(doc, '{notice_to}', first_indent=indent, space_after=18)

    _add_para(doc,
        'YOU ARE HEREBY NOTIFIED that the undersigned will call up for hearing '
        "before the Honorable {judge_name}, judge of the above court, in the "
        "judge's chambers in the {county} County Courthouse, the address of "
        'which is {courthouse_address}, Florida, on {hearing_date}, at '
        '{hearing_time} o’clock {hearing_time_ampm}.M., or as soon thereafter '
        'as same may be heard, the {hearing_subject}.',
        first_indent=indent, space_after=12)

    _add_para(doc,
        'Time set aside by the court is {time_set_aside}.',
        first_indent=indent, space_after=18)

    _add_para(doc, 'PLEASE GOVERN YOURSELVES ACCORDINGLY.',
              first_indent=indent, space_after=18)

    _add_para(doc,
        'I CERTIFY that a copy hereof has been furnished to the above '
        'addressees by {cos_method} on _____ day of ______________________, '
        '20____.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Notice of Hearing')
    _add_miami_dade_ai_certification(doc, 'Notice of Hearing')

    _add_attorney_signature_block(doc)

    # ADA accommodation notice (FLSSI standard footer).
    _add_para(doc, '', space_after=12)
    _add_para(doc,
        'If you are a person with a disability who needs any accommodation in '
        'order to participate in this proceeding, you are entitled, at no cost '
        'to you, to the provision of certain assistance. Please contact '
        '{ada_contact} at least 7 days before your scheduled court appearance, '
        'or immediately upon receiving this notification if the time before '
        'the scheduled appearance is less than 7 days; if you are hearing or '
        'voice impaired, call 711.',
        space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0530.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-CAVEAT  Smart caveat: consolidates FLSSI P1-0301/0305/0311/0315
#   Axes: caveator_type (creditor / interested_person)
#         caveator_is_nonresident (drives FL-attorney representation paragraph
#                                   + attorney signature block)
#   Pro se variants (P1-0300, P1-0310 with designated agent) are intentionally
#   out of scope — the firm always represents the caveator.
# ---------------------------------------------------------------------------

def build_p1_caveat():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    # Title — branches on caveator_type
    _add_para(doc,
        '{#caveator_is_creditor}CAVEAT BY CREDITOR{/caveator_is_creditor}'
        '{#caveator_is_ip}CAVEAT BY INTERESTED PERSON{/caveator_is_ip}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    # Paragraph 1: caveator's interest + decedent ID block
    _pleading_para(doc,
        'The interest of the caveator is '
        '{#caveator_is_creditor}that of a creditor of {decedent_full_name}, '
        'deceased{/caveator_is_creditor}'
        '{#caveator_is_ip}{caveator_interest}, of {decedent_full_name}, '
        'deceased{/caveator_is_ip}, whose last known residence address is '
        '{decedent_last_residence}, the last four digits of whose social '
        'security number, if known, are {decedent_ssn_last4}, whose year of '
        'birth, if known, is {decedent_year_of_birth}, and who died on or '
        'about {decedent_death_date}.')

    # Paragraph 2: caveator's name + addresses
    _pleading_para(doc,
        "Caveator's name, mailing address, and residence address are: "
        '{caveator_name}; {caveator_mailing_address}; '
        '{caveator_residence_address}.')

    # Paragraph 3 (nonresident only): FL attorney representation
    _pleading_para(doc,
        '{#caveator_is_nonresident}Caveator, a non-resident of the State of '
        'Florida, is represented by an attorney admitted to practice in '
        'Florida.{/caveator_is_nonresident}'
        '{^caveator_is_nonresident}Caveator is a resident of the State of '
        'Florida.{/caveator_is_nonresident}')

    indent = Inches(0.5)

    # Request paragraph — branches on type AND residency
    _add_para(doc,
        '{#caveator_is_creditor}Caveator requests that the clerk notify the '
        "{#caveator_is_nonresident}caveator's attorney{/caveator_is_nonresident}"
        '{^caveator_is_nonresident}caveator{/caveator_is_nonresident} in '
        'writing of the date of issuance of letters of administration and of '
        'the names and addresses of the personal representative and the '
        "personal representative's attorney, and that caveator be given such "
        'additional notice as the Florida Probate Rules require.'
        '{/caveator_is_creditor}'
        '{#caveator_is_ip}Caveator requests that the court not admit a will '
        'of the decedent to probate or appoint a personal representative '
        'without formal notice '
        '{#caveator_is_nonresident}on caveator or his or her designated '
        'agent{/caveator_is_nonresident}'
        '{^caveator_is_nonresident}on caveator{/caveator_is_nonresident}, '
        'and that caveator be given such additional notice as the Florida '
        'Probate Rules require.{/caveator_is_ip}',
        first_indent=indent, space_after=18)

    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the '
        'foregoing, and the facts alleged are true, to the best of my '
        'knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Caveat')
    _add_miami_dade_ai_certification(doc, 'Caveat')

    # Caveator signature (always)
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{caveator_name}, Caveator', space_after=24)

    # Attorney signature block — only when nonresident (firm represents
    # caveator as Florida attorney).
    _add_para(doc, '{#caveator_is_nonresident}', space_after=0)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for Caveator', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)
    _add_para(doc, '{/caveator_is_nonresident}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-CAVEAT.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P2-PETITION  Smart Petition for Summary Administration
#   Consolidates 8 FLSSI variants (P2-0204/0205/0214/0215/0219/0220/0224/0225).
#   Axes: is_testate × is_ancillary × multiple_petitioners
#   Self-proved status is a property of the will/order, not the petition.
# ---------------------------------------------------------------------------

def build_p2_petition():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'PETITION FOR SUMMARY ADMINISTRATION',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, '{petitioner_label}, {petitioner_names}, {petitioner_verb_alleges}:',
              space_after=12)

    # 1. Interest of petitioner(s)
    _pleading_para(doc,
        '{^multiple_petitioners}{#petitioners}Petitioner has an interest in the above '
        'estate as {pet_interest}. Petitioner’s address is set forth in paragraph 3 '
        'and the name and office address of petitioner’s attorney are set forth at '
        'the end of this petition.{/petitioners}{/multiple_petitioners}'
        '{#multiple_petitioners}Each petitioner has an interest in the above estate as '
        'set forth below: {#petitioners}{pet_name} is {pet_interest}. {/petitioners}'
        'Their addresses are set forth in paragraph 3 and the name and office address of '
        'their attorney are set forth at the end of this petition.{/multiple_petitioners}')

    # 2. Decedent
    _pleading_para(doc,
        'Decedent, {decedent_full_name}, whose last known address was {decedent_address}, '
        'and the last four digits of whose social security number are {decedent_ssn_last4}, '
        'died on {decedent_death_date}, at {decedent_death_place}, and on the date of death, '
        'decedent was domiciled in '
        '{^is_ancillary}{decedent_domicile} County, Florida{/is_ancillary}'
        '{#is_ancillary}{decedent_domicile_state}{/is_ancillary}.')

    # 3. Beneficiaries
    _pleading_para(doc,
        'So far as is known, the names of the beneficiaries of this estate and of '
        'decedent’s surviving spouse, if any, their addresses and relationships to '
        'decedent, and the years of birth of any who are minors, are:',
        keep_with_next=True)
    _beneficiaries_table(doc)

    # 4. Venue
    _pleading_para(doc,
        'Venue of this proceeding is in this county because {venue_reason}.')

    # 5. Will (testate only) — branches on FL res vs ancillary
    _pleading_para(doc,
        '{#is_testate}'
        '{^is_ancillary}The original of the decedent’s last will, dated {will_date}'
        '{#has_codicil}, and codicil(s) dated {codicil_dates}{/has_codicil}, '
        'is/are in the possession of the above court or accompany/accompanies this petition.'
        '{/is_ancillary}'
        '{#is_ancillary}An authenticated copy of the decedent’s last will, dated '
        '{will_date}{#has_codicil}, and codicil(s) dated {codicil_dates}{/has_codicil}, '
        'and an authenticated copy of so much of the domiciliary proceedings as is required '
        'by Florida Probate Rule 5.470 accompany this petition.{/is_ancillary}'
        '{/is_testate}'
        '{^is_testate}After the exercise of reasonable diligence, {petitioner_label} '
        '{petitioner_verb_is} unaware of any unrevoked wills or codicils of decedent.'
        '{/is_testate}')

    # 6. No other unrevoked wills (testate only)
    _pleading_para(doc,
        '{#is_testate}{petitioner_label} {petitioner_verb_is} unaware of any unrevoked '
        'will or codicil of decedent other than as set forth in paragraph 5.{/is_testate}')

    # 7. Qualifies for summary admin — render only the checked grounds
    _pleading_para(doc,
        '{petitioner_label} {petitioner_verb_is} entitled to summary administration because:'
        '{#qualifies_no_admin_required} Decedent’s will does not direct administration '
        'as required by Florida Statutes Chapter 733.{/qualifies_no_admin_required}'
        '{#qualifies_under_75k} To the best knowledge of the {#multiple_petitioners}'
        'petitioners{/multiple_petitioners}{^multiple_petitioners}petitioner'
        '{/multiple_petitioners}, the value of the entire estate subject to administration '
        'in this state, less the value of property exempt from the claims of creditors, '
        'does not exceed $75,000.{/qualifies_under_75k}'
        '{#qualifies_2_year_rule} The decedent has been dead for more than two years.'
        '{/qualifies_2_year_rule}')

    # 8. Domiciliary / principal probate proceedings
    _pleading_para(doc,
        '{#is_ancillary}Domiciliary proceedings are pending in {domiciliary_court_name}, '
        'the address of which is {domiciliary_court_address}, and letters have been issued '
        'to {domiciliary_representative}, whose address is '
        '{domiciliary_representative_address}.{/is_ancillary}'
        '{^is_ancillary}Domiciliary or principal probate proceedings '
        '{#domiciliary_proceedings_pending}are known to be pending in '
        '{domiciliary_court_name}, the address of which is {domiciliary_court_address}. '
        'Letters have been issued to {domiciliary_representative}, whose address is '
        '{domiciliary_representative_address}.{/domiciliary_proceedings_pending}'
        '{^domiciliary_proceedings_pending}are not known to be pending in another state '
        'or country.{/domiciliary_proceedings_pending}{/is_ancillary}')

    # 9. Assets
    _pleading_para(doc,
        'The following is a complete list of the assets in this estate and their estimated '
        'values, together with those assets claimed to be exempt:',
        keep_with_next=True)
    _estate_assets_table(doc)

    # 10. Creditors status — single-select, render only the chosen branch
    _pleading_para(doc,
        'With respect to claims of creditors:'
        '{#creditors_all_barred} All claims of creditors are barred.{/creditors_all_barred}'
        '{#creditors_no_debt} {petitioner_label} {petitioner_verb_has} made diligent search '
        'and reasonable inquiry for any known or reasonably ascertainable creditors and the '
        'estate is not indebted.{/creditors_no_debt}'
        '{#creditors_has_debt} {petitioner_label} {petitioner_verb_has} made diligent search '
        'and reasonable inquiry for any known or reasonably ascertainable creditors and the '
        'estate is indebted; provision for the payment of debts and the information required '
        'by Florida Statutes section 735.206 and Florida Probate Rule 5.530 are set forth on '
        'the attached schedule.{/creditors_has_debt}')

    # 11. Notice to creditors plan
    _pleading_para(doc,
        'All creditors ascertained to have claims and which have not joined in this '
        'petition or consented to entry of the order requested will be served by formal '
        'notice with a copy of this petition. {petitioner_label} '
        '{#multiple_petitioners}acknowledge{/multiple_petitioners}'
        '{^multiple_petitioners}acknowledges{/multiple_petitioners} that any known or '
        'reasonably ascertainable creditor who did not receive timely notice of this '
        'petition and for whom provision for payment was not made may enforce a timely '
        'claim and, if the creditor prevails, shall be awarded reasonable attorney’s '
        'fees as an element of costs against those who joined in the petition.')

    # 12. Distribution
    _pleading_para(doc,
        'It is proposed that all assets of the decedent, including exempt property, be '
        'distributed to the following:',
        keep_with_next=True)
    _distribution_table(doc)

    # Closing — unnumbered
    indent = Inches(0.5)
    _add_para(doc,
        '{petitioner_label} '
        '{#multiple_petitioners}waive{/multiple_petitioners}'
        '{^multiple_petitioners}waives{/multiple_petitioners} notice of hearing on this '
        'petition and '
        '{#multiple_petitioners}request{/multiple_petitioners}'
        '{^multiple_petitioners}requests{/multiple_petitioners} that '
        '{#is_testate}the decedent’s last will and codicil(s), if applicable, be '
        'admitted to probate and {/is_testate}'
        'an order of summary administration be entered directing distribution of the '
        'assets in the estate in accordance with the schedule set forth above.',
        first_indent=indent, space_before=12, space_after=12)

    _add_para(doc,
        'Under penalties of perjury, {petitioner_label} declare'
        '{^multiple_petitioners}s{/multiple_petitioners} that {petitioner_label} '
        '{petitioner_verb_has} read the foregoing, and the facts alleged are true, to the '
        'best of {petitioner_poss} knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Summary Administration')
    _add_miami_dade_ai_certification(doc, 'Petition for Summary Administration')

    _add_probate_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P2-PETITION.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P2-ORDER  Smart Order of Summary Administration
#   Testate path: combined Order Admitting Will to Probate AND of Summary
#                 Administration (per firm convention — David always uses
#                 the combined order for testate).
#   Intestate path: just Order of Summary Administration.
#
#   Replaces P2-0300/0310/0320/0322/0325 + P2-0500 (and the firm's intent
#   for testate-nonresident combined, which FLSSI does not provide directly).
#
#   Axes: is_testate × is_ancillary × is_self_proved (testate only)
#         × is_auth_copy_of_will (testate ancillary only)
# ---------------------------------------------------------------------------

def build_p2_order():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    # Title — branches on testate × ancillary × auth-copy
    _add_para(doc,
        '{#is_testate}ORDER ADMITTING WILL'
        '{#is_ancillary} OF NONRESIDENT{/is_ancillary} TO PROBATE'
        '{/is_testate}'
        '{#is_testate} AND OF SUMMARY ADMINISTRATION{/is_testate}'
        '{^is_testate}ORDER OF SUMMARY ADMINISTRATION{/is_testate}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)

    _add_para(doc,
        '{#is_testate}{#is_self_proved}(self-proved){/is_self_proved}'
        '{#is_auth_copy_of_will}(authenticated copy of will){/is_auth_copy_of_will}'
        '{/is_testate}'
        '{^is_testate}{#is_ancillary}(nonresident decedent){/is_ancillary}{/is_testate}',
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)

    # Recital — branches heavily on testate / self-proved
    _add_para(doc,
        'On the petition of {petitioner_names} for summary administration of the estate of '
        '{decedent_full_name}, deceased, the court finding that the decedent died on '
        '{decedent_death_date}; that all interested persons have been served proper notice '
        'of the petition and hearing or have waived notice thereof; that the material '
        'allegations of the petition are true; '
        '{#is_testate}that the writing presented to this court as the last will of '
        '{decedent_full_name}, deceased, '
        '{#is_self_proved}having been executed in conformity with law and made self-proved '
        'at the time of its execution by the acknowledgement of the decedent and the '
        'affidavits of the witnesses, made before an officer authorized to administer '
        'oaths and evidenced by the officer’s certificate attached to or following '
        'the will in the form required by law, and{/is_self_proved}'
        '{^is_self_proved}having been established by the oath of {will_witnesses}, a '
        'subscribing and attesting witness, as being the last will of the decedent, and'
        '{/is_self_proved} no objection having been made to its probate; '
        '{/is_testate}'
        'and that the decedent’s estate qualifies for summary administration and an '
        'Order '
        '{#is_testate}Admitting Will to Probate and of Summary Administration{/is_testate}'
        '{^is_testate}of Summary Administration{/is_testate} should be entered, it is',
        first_indent=indent, space_after=12)

    _add_para(doc, 'ADJUDGED that:', bold=True, space_after=12)

    # Testate-only: numbered ¶1 admits the will to probate.
    # Wrapped in {#is_testate}...{/is_testate} marker paragraphs so the entire
    # numbered paragraph is dropped (and renumbered) in intestate mode.
    _add_para(doc, '{#is_testate}', space_after=0)
    _pleading_para(doc,
        'The will dated {will_date}'
        '{#has_codicil}, together with codicil(s) dated {codicil_dates},{/has_codicil}'
        ' attested by {will_witnesses} as subscribing and attesting witnesses, is admitted '
        'to probate according to law as the last will of the decedent.')
    _add_para(doc, '{/is_testate}', space_after=0)

    # Always: distribution preamble (numbered ¶2 in testate, ¶1 in intestate)
    _pleading_para(doc,
        'There be immediate distribution of the assets of the decedent as follows:',
        keep_with_next=True)
    _distribution_table(doc)

    _pleading_para(doc,
        'Those to whom specified assets of the decedent’s estate are distributed by '
        'this order have the right to receive and collect those assets and to maintain '
        'actions to enforce their rights.')

    _pleading_para(doc,
        'Debtors of the decedent, those holding property of the decedent, and those with '
        'whom securities or other property of decedent are registered, are authorized and '
        'directed to comply with this order by paying, delivering, or transferring to the '
        'beneficiaries specified above the parts of the decedent’s estate distributed '
        'to them by this order, and the persons so paying, delivering, or transferring '
        'shall not be accountable to anyone else for the property.')

    _add_para(doc, '', space_after=12)
    _add_para(doc, 'ORDERED on _____ day of ______________________, 20____.',
              space_after=36)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Circuit Judge', space_after=18)

    out_path = os.path.join(TEMPLATE_DIR, 'P2-ORDER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P2-0355  Notice to Creditors (summary administration)
# ---------------------------------------------------------------------------

def build_p2_0355():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE TO CREDITORS',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '(summary administration)',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc, 'TO ALL PERSONS HAVING CLAIMS OR DEMANDS AGAINST',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'THE ABOVE ESTATE:',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'You are hereby notified that an Order of Summary Administration has been entered '
        'in the estate of {decedent_full_name}, deceased, File Number {file_no}, by the '
        'Circuit Court for {county} County, Florida, Probate Division, the address of '
        'which is {court_address}; that the decedent’s date of death was '
        '{decedent_death_date}; that the total value of the estate is {total_estate_value} '
        'and that the names and addresses of those to whom it has been assigned by such '
        'order are:',
        first_indent=indent, space_after=12)

    _distribution_table(doc)

    _add_para(doc, '', space_after=12)
    _add_para(doc, 'ALL INTERESTED PERSONS ARE NOTIFIED THAT:',
              bold=True, space_after=12)

    _add_para(doc,
        'All creditors of the estate of the decedent and persons having claims or demands '
        'against the estate of the decedent other than those for whom provision for full '
        'payment was made in the Order of Summary Administration must file their claims '
        'with this court WITHIN THE TIME PERIODS SET FORTH IN FLORIDA STATUTES SECTION '
        '733.702. ALL CLAIMS AND DEMANDS NOT SO FILED WILL BE FOREVER BARRED. '
        'NOTWITHSTANDING ANY OTHER APPLICABLE TIME PERIOD, ANY CLAIM FILED TWO (2) YEARS '
        'OR MORE AFTER THE DECEDENT’S DATE OF DEATH IS BARRED.',
        first_indent=indent, space_after=18)

    _add_para(doc,
        'The date of first publication of this Notice is _____ day of '
        '______________________, 20____.',
        first_indent=indent, space_after=24)

    _add_broward_ai_certification(doc, 'Notice to Creditors (Summary Administration)')
    _add_miami_dade_ai_certification(doc, 'Notice to Creditors (Summary Administration)')

    # Signature: attorney + petitioner (Person Giving Notice).
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{petitioner_name}, Person Giving Notice', space_after=24)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P2-0355.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# Shared notary block (acknowledged before notary; Phase 7b style — handwritten
# date and ID at notarization). Used by P1-0620 and P3-CURATOR-OATH.
# ---------------------------------------------------------------------------

def _add_notary_block(doc, *, signer_tag='{joining_party_name}'):
    indent = Inches(0.5)
    _add_para(doc,
        'Sworn to (or affirmed) and subscribed before me by means of '
        '☐ online notarization or ☐ physical presence this _____ day '
        f'of ______________________, 20____, by {signer_tag}, who is personally '
        'known to me or produced ______________________________ as identification.',
        first_indent=indent, space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Notary Public', space_after=0)
    _add_para(doc, 'State of Florida', space_after=0)
    _add_para(doc, 'My Commission Expires: ______________________', space_after=0)


# ---------------------------------------------------------------------------
# P1-0100  Petition to Open Safe Deposit Box
# ---------------------------------------------------------------------------

def build_p1_0100():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'PETITION TO OPEN SAFE DEPOSIT BOX',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        "Petitioner has an interest in the above estate as {petitioner_interest}. "
        "Petitioner's address is {petitioner_address}, and the name and address of "
        "petitioner's attorney are set forth at the end of this petition.")

    _pleading_para(doc,
        'Decedent, {decedent_full_name}, whose last known address was '
        '{decedent_address}, and, if known, whose age was {decedent_age_at_death}, '
        'died on {decedent_death_date}, at {decedent_death_place}, and on the date '
        'of death decedent was domiciled in {decedent_domicile} County, Florida.')

    _pleading_para(doc,
        'The decedent was the lessee or co-lessee of a safe deposit box (No. '
        '{sdb_number}) leased to the decedent by {sdb_lessor_name}, the address '
        'of which is {sdb_lessor_address}. The name(s) of any co-lessee(s), if '
        'known, is/are: {sdb_colessees}.')

    _pleading_para(doc,
        'Petitioner is informed and believes that the decedent may have left in '
        'the safe deposit box: (a) a will or codicil of the decedent or a writing '
        'described in Florida Statutes section 732.515 purporting to identify '
        'devises of tangible property; (b) a deed to a burial plot; (c) a writing '
        'giving burial instructions; and/or (d) insurance policies on the life of '
        'the decedent.')

    indent = Inches(0.5)
    _add_para(doc,
        'Petitioner requests that an order be entered authorizing petitioner, in '
        'the presence of an officer of the lessor, to open and examine the contents '
        'of the safe deposit box leased or co-leased by the decedent and directing '
        'the lessor to deliver: (a) to the court having probate jurisdiction in the '
        'county where the lessor is located, any writing purporting to be a will or '
        'codicil of the decedent and any writing described in Florida Statutes '
        'section 732.515 purporting to identify devises of tangible property; '
        '(b) to petitioner, any writing purporting to be a deed to a burial plot or '
        'to give burial instructions; and (c) to the beneficiary named therein, any '
        'document purporting to be an insurance policy on the life of the decedent; '
        'and directing the lessor to make a complete copy of any document removed '
        'and delivered and to place that copy, together with a memorandum of '
        'delivery identifying the officer, the person to whom the document was '
        'delivered and the date of delivery, in the safe deposit box leased or '
        'co-leased by the decedent.',
        first_indent=indent, space_after=18)

    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, '
        'and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition to Open Safe Deposit Box')
    _add_miami_dade_ai_certification(doc, 'Petition to Open Safe Deposit Box')

    _add_probate_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0100.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-0620  Joinder, Waiver and Consent (with notary block per local rules)
# ---------------------------------------------------------------------------

def build_p1_0620():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'JOINDER, WAIVER AND CONSENT',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'The undersigned, whose name is {joining_party_name}, and who has an '
        'interest in this estate as {joining_party_interest}, acknowledges receipt '
        'of a copy of the {petition_title} heretofore filed in this proceeding, '
        'joins in the petition, waives hearing and notice of hearing thereon, and '
        'consents to the entry of an order granting the relief requested in the '
        'petition.',
        first_indent=indent, space_after=24)

    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{joining_party_name}', space_after=24)

    _add_notary_block(doc, signer_tag='{joining_party_name}')

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0620.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-NOTICE-CONFIDENTIAL  Smart Notice of Confidential Info
#   Consolidates FLSSI P1-0640 (non-contemporaneous) + P1-0641 (contemporaneous).
#   Axis: is_contemporaneous
# ---------------------------------------------------------------------------

def build_p1_notice_confidential():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE OF CONFIDENTIAL INFORMATION',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'WITHIN COURT FILING',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '{#is_contemporaneous}(contemporaneous filing){/is_contemporaneous}'
        '{^is_contemporaneous}(non-contemporaneous filing){/is_contemporaneous}',
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'Pursuant to Florida Rule of General Practice and Judicial Administration '
        '2.420(d)(2), I hereby certify:',
        first_indent=indent, space_after=12)

    # Body — branches on is_contemporaneous
    _add_para(doc,
        '{#is_contemporaneous}I am filing herewith a document containing '
        'confidential information as described in Rule 2.420(d)(1)(B), and '
        'that:{/is_contemporaneous}'
        '{^is_contemporaneous}A document was previously filed in this case that '
        'contains confidential information as described in Rule 2.420(d)(1)(B), '
        'but a Notice of Confidential Information within Court Filing was not '
        'filed with the document and the confidential information was not '
        'maintained as confidential by the clerk of the court. I hereby notify '
        'the clerk that this confidential information is located as '
        'follows:{/is_contemporaneous}',
        first_indent=indent, space_after=12)

    # Document title (always)
    _add_para(doc,
        '(a) Title/type of document: {conf_doc_title}',
        first_indent=indent, space_after=6)

    # Non-contemporaneous: extra fields (filing date, doc date, docket entry)
    _add_para(doc,
        '{^is_contemporaneous}(b) Date of filing: {conf_doc_filing_date}'
        '{/is_contemporaneous}',
        first_indent=indent, space_after=6)
    _add_para(doc,
        '{^is_contemporaneous}(c) Date of document: {conf_doc_date}'
        '{/is_contemporaneous}',
        first_indent=indent, space_after=6)
    _add_para(doc,
        '{^is_contemporaneous}(d) Docket entry number: {conf_docket_entry}'
        '{/is_contemporaneous}',
        first_indent=indent, space_after=6)

    # Confidential location (always — relabel index based on contemporaneous)
    _add_para(doc,
        '{#is_contemporaneous}(b){/is_contemporaneous}'
        '{^is_contemporaneous}(e){/is_contemporaneous} '
        '{#conf_entire_document}Entire document is confidential.'
        '{/conf_entire_document}'
        '{^conf_entire_document}Precise location of confidential information '
        'in document: {conf_precise_location}{/conf_entire_document}',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Notice of Confidential Information')
    _add_miami_dade_ai_certification(doc, 'Notice of Confidential Information')

    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=24)

    # Certificate of Service
    _add_para(doc, 'CERTIFICATE OF SERVICE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    _add_para(doc,
        'I HEREBY CERTIFY that a copy of the foregoing was furnished by '
        '{cos_method} to {cos_recipients} on this _____ day of '
        '______________________, 20____.',
        first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=18)
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-NOTICE-CONFIDENTIAL.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-CURATOR-PETITION  Petition to Appoint Curator
# ---------------------------------------------------------------------------

def build_p3_curator_petition():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'PETITION TO APPOINT CURATOR',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    _pleading_para(doc,
        'Petitioner has an interest in the above estate as {petitioner_interest}. '
        "Petitioner's address is {petitioner_address}, and the name and office "
        "address of petitioner's attorney are set forth at the end of this petition.")

    _pleading_para(doc,
        'Decedent, {decedent_full_name}, whose last known address was '
        '{decedent_address}, and, if known, whose age was {decedent_age_at_death} '
        'and the last four digits of whose social security number are '
        '{decedent_ssn_last4}, died on {decedent_death_date} at '
        '{decedent_death_place}. On the date of death, decedent was domiciled in '
        '{decedent_domicile} County, Florida.')

    _pleading_para(doc,
        'So far as is known, the names of persons apparently entitled to letters '
        'of administration, the beneficiaries of this estate and the decedent’s '
        'surviving spouse, if any, their addresses and relationships to decedent, '
        'and the years of birth of any who are minors, are:',
        keep_with_next=True)
    _beneficiaries_table(doc)

    _pleading_para(doc,
        'Venue of this proceeding is in this county because {venue_reason}.')

    _pleading_para(doc,
        'The nature and approximate value of the assets of the estate are:',
        keep_with_next=True)
    _estate_assets_table(doc)

    _pleading_para(doc,
        'Petitioner proposes that {proposed_curator_name}, whose address is '
        '{proposed_curator_address}, and who is qualified under the laws of the '
        "State of Florida to serve as personal representative of the decedent's "
        'estate, be appointed as curator of the estate.')

    _pleading_para(doc,
        'The court should appoint a curator because {curator_appointment_reason}.')

    indent = Inches(0.5)
    _add_para(doc,
        'WHEREFORE, Petitioner respectfully requests that {proposed_curator_name} '
        'be appointed curator of the estate of the decedent.',
        first_indent=indent, space_before=12, space_after=12)

    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, '
        'and the facts alleged are true, to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition to Appoint Curator')
    _add_miami_dade_ai_certification(doc, 'Petition to Appoint Curator')

    _add_probate_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-CURATOR-PETITION.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-CURATOR-ORDER  Smart Order Appointing Curator (bond y/n)
#   Consolidates FLSSI P3-0065 (no bond) + P3-0070 (bond required).
# ---------------------------------------------------------------------------

def build_p3_curator_order():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'ORDER APPOINTING CURATOR',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '{#bond_required}(bond required){/bond_required}'
        '{^bond_required}(no bond){/bond_required}',
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'On the petition of {petitioner_name} to appoint a curator of the estate '
        'of {decedent_full_name}, deceased, the court finding that the decedent '
        'died on {decedent_death_date}, and that {curator_name} is entitled to '
        'appointment as curator of the estate because {curator_appointment_reason}, '
        'it is',
        first_indent=indent, space_after=12)

    _add_para(doc, 'ADJUDGED that:', bold=True, space_after=12)

    _pleading_para(doc,
        '{curator_name} is appointed curator of the estate of the decedent, and '
        'that upon taking the prescribed oath and filing designation and '
        'acceptance of resident agent'
        '{#bond_required} and posting bond in the amount of {bond_amount}'
        '{/bond_required}, letters of curatorship shall be issued.')

    _pleading_para(doc, 'The curator shall have the following powers: {curator_powers}')

    _add_para(doc, '', space_after=12)
    _add_para(doc, 'ORDERED on _____ day of ______________________, 20____.',
              space_after=36)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Circuit Judge', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-CURATOR-ORDER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-CURATOR-OATH  Oath of Curator + Designation/Acceptance of Resident Agent
# ---------------------------------------------------------------------------

def build_p3_curator_oath():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'OATH OF CURATOR AND',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'DESIGNATION AND ACCEPTANCE OF RESIDENT AGENT',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'STATE OF FLORIDA', space_after=0)
    _add_para(doc, 'COUNTY OF {notary_county}', space_after=18)

    _add_para(doc,
        'I, {curator_name} (Affiant), state under oath that:',
        space_after=12)

    _pleading_para(doc,
        'I am qualified within the provisions of sections 733.302, 733.303, '
        '733.304, and 733.501, Florida Statutes, to serve as curator of the '
        "estate of {decedent_full_name}, deceased. I have reviewed the statutes "
        "and understand the qualifications. Under penalties of perjury, I "
        'certify that the following statements are true: (a) I am 18 years of '
        'age or older; (b) I have never been convicted of a felony; (c) I have '
        'never been convicted in any state or foreign jurisdiction of abuse, '
        'neglect or exploitation of an elderly person or a disabled adult, as '
        'those terms are defined in Florida Statutes section 825.101; (d) I am '
        'mentally and physically able to perform the duties of curator; and '
        '(e) I am a resident of the State of Florida, or, if I am not, I am '
        'qualified to serve under Florida Statutes section 733.304.')

    _pleading_para(doc,
        'I will faithfully administer the estate of the decedent according to '
        'law.')

    _pleading_para(doc,
        'My place of residence is {curator_residence} and my post office '
        'address is {curator_mailing_address}.')

    _pleading_para(doc,
        'I will promptly file and serve a notice on all interested persons at '
        'any time I know that I would not be qualified for appointment, and '
        'will include the reason I would not then be qualified and the date on '
        'which the disqualifying event occurred.')

    _pleading_para(doc,
        'I will file and serve a notice within 20 days on all interested '
        'persons in the event there is a change in my residence address, '
        'street address, or mailing address.')

    _pleading_para(doc,
        'I hereby designate {resident_agent_name}, who '
        '{#resident_agent_is_fla_bar}is{/resident_agent_is_fla_bar}'
        '{^resident_agent_is_fla_bar}is not{/resident_agent_is_fla_bar} a '
        'member of The Florida Bar, whose address is {resident_agent_address}, '
        'as my agent for the service of process or notice in any action '
        'against me, either in my representative capacity, or personally, if '
        'the personal action accrued in the administration of the estate.')

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{curator_name}, Affiant', space_after=24)

    _add_notary_block(doc, signer_tag='{curator_name}')

    # Acceptance by resident agent
    _add_para(doc, '', space_after=12)
    _add_para(doc, 'ACCEPTANCE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)
    _add_para(doc,
        'I CERTIFY that my address is as indicated above. I hereby accept the '
        'foregoing designation as Resident Agent.',
        first_indent=indent, space_after=24)
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{resident_agent_name}, Resident Agent', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-CURATOR-OATH.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-CURATOR-LETTERS  Letters of Curatorship (judge-signed)
# ---------------------------------------------------------------------------

def build_p3_curator_letters():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'LETTERS OF CURATORSHIP',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'TO ALL WHOM IT MAY CONCERN', bold=True, space_after=12)

    _add_para(doc,
        'WHEREAS, {decedent_full_name}, a resident of {decedent_address}, died '
        'on {decedent_death_date}, owning assets in the State of Florida, and',
        first_indent=indent, space_after=12)

    _add_para(doc,
        'WHEREAS, {curator_name} has been appointed curator of the estate of '
        'the decedent and has performed all acts prerequisite to issuance of '
        'Letters of Curatorship in the estate,',
        first_indent=indent, space_after=12)

    _add_para(doc,
        'NOW, THEREFORE, I, the undersigned circuit judge, declare {curator_name} '
        'duly qualified to act as curator of this estate of {decedent_full_name}, '
        'deceased, and to perform the following duties or functions of a personal '
        'representative: {curator_powers}',
        first_indent=indent, space_after=24)

    _add_para(doc, 'ORDERED on _____ day of ______________________, 20____.',
              space_after=36)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Circuit Judge', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-CURATOR-LETTERS.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-OATH-WITNESS  Smart Oath of Witness to Will or Codicil
#   Consolidates FLSSI P3-0300 (will), P3-0301 (will copy), P3-0310 (codicil),
#   P3-0311 (codicil copy). Axes: is_codicil × is_copy.
# ---------------------------------------------------------------------------

def build_p3_oath_witness():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc,
        '{^is_codicil}OATH OF WITNESS TO WILL{/is_codicil}'
        '{#is_codicil}OATH OF WITNESS TO CODICIL{/is_codicil}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '{#is_copy}(copy){/is_copy}',
              align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'STATE OF FLORIDA', space_after=0)
    _add_para(doc, 'COUNTY OF {notary_county}', space_after=18)

    # Body — branches on is_copy and is_codicil
    _add_para(doc,
        'The undersigned, {witness_name}, being duly sworn says that '
        '{#is_copy}the photographic copy annexed to this oath is a true copy of '
        '{/is_copy}'
        '{^is_copy}the writing exhibited to the undersigned as {/is_copy}'
        '{^is_codicil}the last will of {decedent_full_name}, deceased{/is_codicil}'
        '{#is_codicil}the {codicil_ordinal} codicil to the last will of '
        '{decedent_full_name}, deceased{/is_codicil}, '
        '{#is_copy}that the decedent executed{/is_copy}'
        '{^is_copy}is the same writing that the decedent executed{/is_copy} '
        'and that the undersigned and {co_witness_name} subscribed as attesting '
        'witnesses on {execution_date}; that the decedent signed the '
        '{#is_copy}original of the {/is_copy}writing at the end in the presence '
        'of the attesting witnesses or acknowledged in the presence of the '
        'attesting witnesses that the decedent had previously signed the writing '
        'at the end; and that the witnesses, in the presence of the decedent and '
        'in the presence of each other, subscribed their names to the '
        '{#is_copy}original of the {/is_copy}writing as attesting witnesses.',
        first_indent=indent, space_after=24)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{witness_name}, Affiant', space_after=24)

    _add_para(doc,
        'Sworn to and subscribed before me on _____ day of '
        '______________________, 20____.',
        first_indent=indent, space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '(Circuit Judge) (Clerk) (Deputy Clerk)', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-OATH-WITNESS.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-PROOF-WILL  Smart Proof of Will or Codicil
#   Consolidates FLSSI P3-0320 (will) + P3-0330 (codicil). Axis: is_codicil.
#   Two select-driven branches: witness_unavailable_reason, affiant_relation.
# ---------------------------------------------------------------------------

def build_p3_proof_will():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc,
        '{^is_codicil}PROOF OF WILL{/is_codicil}'
        '{#is_codicil}PROOF OF CODICIL{/is_codicil}',
        align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc, 'STATE OF FLORIDA', space_after=0)
    _add_para(doc, 'COUNTY OF {notary_county}', space_after=18)

    _add_para(doc,
        'The undersigned, {affiant_name}, being duly sworn says that:',
        space_after=12)

    _pleading_para(doc,
        'The attesting witnesses to the writing dated {execution_date}, '
        'exhibited to the undersigned as '
        '{^is_codicil}the last will of {decedent_full_name}, deceased{/is_codicil}'
        '{#is_codicil}the {codicil_ordinal} codicil to the last will of '
        '{decedent_full_name}, deceased{/is_codicil}: '
        '{#witness_unavailable_cannot_be_found}cannot be found.'
        '{/witness_unavailable_cannot_be_found}'
        '{#witness_unavailable_incapacitated}have become incapacitated after the '
        'execution of the {^is_codicil}will{/is_codicil}{#is_codicil}codicil'
        '{/is_codicil}.{/witness_unavailable_incapacitated}'
        '{#witness_unavailable_unavailable}are unavailable so that their '
        'testimony cannot be taken within a reasonable time.'
        '{/witness_unavailable_unavailable}')

    _pleading_para(doc,
        'The undersigned: '
        '{#affiant_is_pr_nominated}is the personal representative nominated by '
        'the {^is_codicil}will{/is_codicil}{#is_codicil}will or codicil'
        '{/is_codicil}.{/affiant_is_pr_nominated}'
        '{#affiant_has_no_interest}has no interest in the estate under the '
        '{^is_codicil}will{/is_codicil}{#is_codicil}will or codicil'
        '{/is_codicil}.{/affiant_has_no_interest}')

    _pleading_para(doc,
        'The undersigned believes the writing exhibited to the undersigned to '
        'be the {^is_codicil}last will{/is_codicil}'
        '{#is_codicil}{codicil_ordinal} codicil to the last will{/is_codicil} '
        'of the decedent.')

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{affiant_name}, Affiant', space_after=24)

    _add_para(doc,
        'Sworn to and subscribed before me on _____ day of '
        '______________________, 20____.',
        first_indent=indent, space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '(Circuit Judge) (Clerk) (Deputy Clerk)', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-PROOF-WILL.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P1-0800  Notice of Trust (formerly P4-0800)
#   Filed by trustee (or attorney for trustee) when settlor of a Fla. Stat.
#   §733.707(3) trust dies — puts court + PR on notice that the trust exists
#   and may be liable for estate administration expenses if estate is
#   insufficient.
# ---------------------------------------------------------------------------

def build_p1_0800():
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE OF TRUST',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        '{decedent_full_name}, a resident of {decedent_domicile} County, '
        'Florida, who died on {decedent_death_date}, was the settlor of a '
        'trust entitled:',
        first_indent=indent, space_after=6)

    _add_para(doc, '{trust_name}', first_indent=indent, bold=True, space_after=6)

    _add_para(doc,
        'dated {trust_date}, which is a trust described in Florida Statutes '
        "section 733.707(3) and is liable for the expenses of the administration "
        "of the decedent's estate and enforceable claims of the decedent's "
        'creditors to the extent the decedent’s estate is insufficient to '
        'pay them, as provided in Florida Statutes section 733.607(2).',
        first_indent=indent, space_after=12)

    _add_para(doc,
        'The name and address of the trustee are set forth below.',
        first_indent=indent, space_after=12)

    _add_para(doc,
        'The clerk shall file and index this notice of trust in the same manner '
        'as a caveat, unless there exists a probate proceeding for the '
        "settlor’s estate in which case this notice of trust must be filed "
        'in the probate proceeding and the clerk shall send a copy to the '
        'personal representative.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Notice of Trust')
    _add_miami_dade_ai_certification(doc, 'Notice of Trust')

    # Trustee signature block
    _add_para(doc, 'Signed on this _____ day of ______________________, 20____.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{trustee_name}, Trustee', space_after=0)
    _add_para(doc, '{trustee_address}', space_after=24)

    # Clerk acknowledgement block (filled in by the clerk on filing)
    _add_para(doc,
        'Copy mailed to attorney for the Personal Representative on '
        '_____ day of ______________________, 20____.',
        first_indent=indent, space_after=18)
    _add_para(doc, 'CLERK OF THE CIRCUIT COURT', space_after=12)
    _add_para(doc, 'By: _______________________________________', space_after=0)
    _add_para(doc, '       Deputy Clerk', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P1-0800.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# Discharge cluster (full-waiver path) — Phase 11 (E)
# ---------------------------------------------------------------------------
# Replaces 8 FLSSI forms (P5-0510/0511, P5-0550/0551, P5-0700/0701,
# P5-0800/0810) with 4 smart templates branching on multiple_prs and
# receipt_includes_refunding. David never uses the non-full-waiver path
# (P5-0400/0401/0500/0501/0410/0411/0420 — all dropped).
# ---------------------------------------------------------------------------

def build_p5_petition_discharge_full_waiver():
    """Petition for Discharge (full waiver path) — replaces P5-0550 +
    P5-0551. Branches on multiple_prs for grammar."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'PETITION FOR DISCHARGE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '(full waiver — '
        '{^multiple_prs}single personal representative{/multiple_prs}'
        '{#multiple_prs}multiple personal representatives{/multiple_prs})',
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # Intro — petitioner == PR(s).
    _add_para(doc,
        '{petitioner_label}, {pr_names}, as {pr_label} of the above estate, '
        '{petitioner_verb_alleges}:',
        space_after=12)

    # 1. Date of death + letters issuance.
    _pleading_para(doc,
        'The decedent, {decedent_full_name}, a resident of '
        '{^is_ancillary}{decedent_domicile} County, Florida{/is_ancillary}'
        '{#is_ancillary}{decedent_domicile_state}{/is_ancillary}, died on '
        '{decedent_death_date}, and Letters of '
        '{#is_ancillary}Ancillary {/is_ancillary}Administration were issued '
        'to {petitioner_label} on {letters_issued_date}.')

    # 2. Full administration.
    _pleading_para(doc,
        '{petitioner_label} {petitioner_verb_has} fully administered this '
        'estate by making payment, settlement, or other disposition of all '
        'claims and debts that were presented, and by paying or making '
        'provision for the payment of all taxes and expenses of administration.')

    # 3. Estate tax.
    _pleading_para(doc,
        '{petitioner_label} {petitioner_verb_has} '
        '{#estate_tax_return_required}filed all required estate tax returns '
        'with the Internal Revenue Service and {petitioner_verb_has} obtained '
        'and filed with this court evidence of the satisfaction of this '
        'estate’s obligations for federal estate tax.{/estate_tax_return_required}'
        '{^estate_tax_return_required}determined that no federal estate tax '
        'return was required for this estate.{/estate_tax_return_required}')

    # 4. Other interested persons (table).
    _pleading_para(doc,
        'The only persons, other than {petitioner_label}, having an interest '
        'in this proceeding, and their respective addresses are:',
        keep_with_next=True)
    _interested_persons_table(doc)

    # 5. Waivers (lettered list a-i — use level-1 of the pleading numbering).
    _pleading_para(doc,
        '{petitioner_label}, pursuant to Florida Statutes section 731.302, '
        '{petitioner_verb_has} filed herewith waivers and receipts signed by '
        'all interested persons:',
        keep_with_next=True)

    waivers = [
        'acknowledging that they are aware of the right to have a final accounting;',
        'waiving the filing and service of a final accounting;',
        'waiving the inclusion in this petition of the amount of compensation paid or to be paid to the personal representative, attorneys, accountants, appraisers or other agents employed by the personal representative and the manner of determining that compensation;',
        'acknowledging that they have actual knowledge of the amount and manner of determining the compensation of the personal representative, attorneys, accountants, appraisers, or other agents employed by the personal representative, and agreeing to the amount and manner of determining such compensation, and waiving any objections to the payment of such compensation;',
        'waiving the inclusion in this petition of a plan of distribution;',
        'waiving service of this petition and all notice thereof;',
        'waiving all objections to any accounting and to the Petition for Discharge;',
        'acknowledging receipt of complete distribution of the share of the estate to which they are entitled;',
        'consenting to the entry of an order discharging {petitioner_label}, as {pr_label}, without notice, hearing or waiting period and without further accounting.',
    ]
    for w in waivers:
        _pleading_para(doc, w, level=1)

    # WHEREFORE
    indent = Inches(0.5)
    _add_para(doc,
        '{petitioner_label} request{^multiple_petitioners}s{/multiple_petitioners} '
        'that an order be entered discharging {petitioner_label} as {pr_label} '
        'of this estate and releasing the surety on any bond which '
        '{petitioner_label} may have posted in this proceeding from any '
        'liability on it.',
        first_indent=indent, space_before=12, space_after=12)

    _add_para(doc,
        'Under penalties of perjury, {petitioner_label} declare'
        '{^multiple_petitioners}s{/multiple_petitioners} that {petitioner_label} '
        '{petitioner_verb_has} read the foregoing, and the facts alleged are '
        'true, to the best of {petitioner_poss} knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Petition for Discharge')
    _add_miami_dade_ai_certification(doc, 'Petition for Discharge')

    _add_probate_signature_block(doc)

    out_path = os.path.join(TEMPLATE_DIR, 'P5-PETITION-DISCHARGE-FULL-WAIVER.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


def _interested_persons_table(doc):
    """Two-column bordered table for the discharge petition's
    'persons having an interest' list. Repeats over `interested_persons[]`."""
    tbl = _table_with_borders(doc, rows=2, cols=2, col_widths_in=[2.8, 3.7])
    for cell in tbl.rows[0].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(0, 0), 'NAME', bold=True, space_after=0)
    _cell_para(tbl.cell(0, 1), 'ADDRESS', bold=True, space_after=0)
    for cell in tbl.rows[1].cells:
        _clear_cell(cell)
    _cell_para(tbl.cell(1, 0), '{#interested_persons}{ip_name}', space_after=0)
    _cell_para(tbl.cell(1, 1), '{ip_address}{/interested_persons}', space_after=0)
    return tbl


def build_p5_receipt():
    """Receipt — replaces P5-0510 (basic) + P5-0511 (with refunding
    language). Branches on receipt_includes_refunding."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')

    _add_probate_caption(doc)

    _add_para(doc, 'RECEIPT',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '{#receipt_includes_refunding}(with refunding language){/receipt_includes_refunding}',
        align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, space_after=18)

    _add_para(doc,
        'The undersigned, {receipt_recipient_name}, having an interest in '
        'this estate as {receipt_recipient_interest}, hereby acknowledges '
        'having received from the {pr_label} of this estate the following:',
        space_after=12)

    indent = Inches(0.5)
    _add_para(doc, '{receipt_property_received}',
              first_indent=indent, space_after=18)

    # Conditional refunding paragraph (P5-0511 variant).
    _add_para(doc,
        '{#receipt_includes_refunding}A condition of this distribution is '
        'the agreement by the undersigned to return to the {pr_label}, upon '
        'demand, any property improperly received and its income since '
        'distribution or, if the undersigned does not have the property, to '
        'return to the {pr_label} the value of the property at the date of '
        'disposition and its income and gain received. The undersigned shall '
        'have no obligation to return the property unless it was improperly '
        'distributed.{/receipt_includes_refunding}',
        space_after=18)

    _add_broward_ai_certification(doc, 'Receipt')
    _add_miami_dade_ai_certification(doc, 'Receipt')

    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{receipt_recipient_name}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P5-RECEIPT.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


def build_p5_report_dist():
    """Report of Distribution — replaces P5-0700 + P5-0701. Branches on
    multiple_prs."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'REPORT OF DISTRIBUTION',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '({^multiple_prs}single personal representative{/multiple_prs}'
        '{#multiple_prs}multiple personal representatives{/multiple_prs})',
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    _add_para(doc,
        '{petitioner_label}, {pr_names}, {petitioner_verb_alleges}:',
        space_after=12)

    _pleading_para(doc,
        '{petitioner_label}, as the {pr_label} of the above estate, file'
        '{^multiple_petitioners}s{/multiple_petitioners} herewith all '
        'remaining receipts necessary to show that the estate has been '
        'distributed in accordance with the proposed plan of distribution '
        'set forth in the Petition for Discharge.')

    _pleading_para(doc,
        'The claims of all creditors have been paid or otherwise disposed of.')

    indent = Inches(0.5)
    _add_para(doc,
        '{petitioner_label} request{^multiple_petitioners}s{/multiple_petitioners} '
        'that an order be entered discharging {petitioner_label} as '
        '{pr_label} of this estate, and releasing the surety on '
        '{petitioner_poss} bond, if any, from further liability thereon.',
        first_indent=indent, space_before=12, space_after=12)

    _add_para(doc,
        'Under penalties of perjury, {petitioner_label} declare'
        '{^multiple_petitioners}s{/multiple_petitioners} that {petitioner_label} '
        '{petitioner_verb_has} read the foregoing, and the facts alleged are '
        'true, to the best of {petitioner_poss} knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Report of Distribution')
    _add_miami_dade_ai_certification(doc, 'Report of Distribution')

    _add_probate_signature_block(doc)

    # Certificate of Service (FLSSI P5-0700/0701 includes one).
    _add_para(doc, 'CERTIFICATE OF SERVICE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_before=18, space_after=12)
    _add_para(doc,
        'I CERTIFY that a copy of the foregoing Report of Distribution has '
        'been furnished to {cos_served_to} by {cos_service_method} on this '
        '_____ day of __________, 20___.',
        first_indent=indent, space_after=24)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {petitioner_label}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P5-REPORT-DIST.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


def build_p5_order_discharge():
    """Order of Discharge — replaces P5-0800 + P5-0810. Branches on
    multiple_prs. Judge-signed → NO AI certification."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')

    _add_probate_caption(doc)

    _add_para(doc, 'ORDER OF DISCHARGE',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc,
        '({^multiple_prs}single personal representative{/multiple_prs}'
        '{#multiple_prs}multiple personal representatives{/multiple_prs})',
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    _add_para(doc,
        'On the Petition for Discharge of {pr_names}, as {pr_label} of the '
        'estate of {decedent_full_name}, deceased, the court finding that '
        'the estate has been fully administered and properly distributed, '
        'that claims of creditors have been paid or otherwise disposed of '
        'and that the {pr_label} should be discharged, it therefore is',
        space_after=18)

    _add_para(doc,
        'ADJUDGED that the {pr_label} '
        '{^multiple_prs}is{/multiple_prs}{#multiple_prs}are{/multiple_prs} '
        'discharged, and the surety on the {pr_label}’s bond, if any, is '
        'released from further liability.',
        first_indent=Inches(0.5), space_after=18)

    _add_order_signature_block(doc)

    # NO AI certification — judge-signed. Hard rule.

    out_path = os.path.join(TEMPLATE_DIR, 'P5-ORDER-DISCHARGE.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-0740  Notice to Creditors (formal administration)
# ---------------------------------------------------------------------------

def build_p3_0740():
    """Notice to Creditors per F.S. §733.2121(2). Published in a newspaper
    of general circulation in the county where decedent's estate is being
    administered + served on known/reasonably ascertainable creditors."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')

    _add_probate_caption(doc)

    _add_para(doc, 'NOTICE TO CREDITORS',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc,
        'The administration of the estate of {decedent_full_name}, deceased, whose date of death was {decedent_death_date}, is pending in the Circuit Court for {county} County, Florida, Probate Division, the address of which is {court_address}. The names and addresses of the personal representative and the personal representative’s attorney are set forth below.',
        space_after=12)

    _add_para(doc,
        'All creditors of the decedent and other persons having claims or demands against decedent’s estate on whom a copy of this notice is required to be served must file their claims with this court ON OR BEFORE THE LATER OF 3 MONTHS AFTER THE TIME OF THE FIRST PUBLICATION OF THIS NOTICE OR 30 DAYS AFTER THE DATE OF SERVICE OF A COPY OF THIS NOTICE ON THEM.',
        space_after=12)

    _add_para(doc,
        'All other creditors of the decedent and other persons having claims or demands against decedent’s estate must file their claims with this court WITHIN 3 MONTHS AFTER THE DATE OF THE FIRST PUBLICATION OF THIS NOTICE.',
        space_after=12)

    _add_para(doc,
        'The personal representative has no duty to discover whether any property held at the time of the decedent’s death by the decedent or the decedent’s surviving spouse is property to which the Florida Uniform Disposition of Community Property Rights at Death Act as described in ss. 732.216-732.228, applies, or may apply, unless a written demand is made by a creditor as specified under s. 732.2211, Florida Statutes.',
        space_after=12)

    _add_para(doc,
        'ALL CLAIMS NOT FILED WITHIN THE TIME PERIODS SET FORTH IN FLORIDA STATUTES SECTION 733.702 WILL BE FOREVER BARRED.',
        bold=True, space_after=6)
    _add_para(doc,
        'NOTWITHSTANDING THE TIME PERIODS SET FORTH ABOVE, ANY CLAIM FILED TWO (2) YEARS OR MORE AFTER THE DECEDENT’S DATE OF DEATH IS BARRED.',
        bold=True, space_after=18)

    _add_para(doc,
        'The date of first publication of this notice is {publication_date}.',
        space_after=24)

    _add_broward_ai_certification(doc, 'Notice to Creditors')
    _add_miami_dade_ai_certification(doc, 'Notice to Creditors')

    # Two-column signature block: PR (right) and Attorney (left). Use a
    # borderless table for clean alignment.
    sig = _borderless_table(doc, rows=1, cols=2, col_widths_in=[3.25, 3.25])
    _clear_cell(sig.cell(0, 0))
    _cell_para(sig.cell(0, 0), 'Attorney for Personal Representative:', bold=True, space_after=12)
    _cell_para(sig.cell(0, 0), '_______________________________________', space_after=0)
    _cell_para(sig.cell(0, 0), '{attorney_name}', space_after=12)
    _cell_para(sig.cell(0, 0), 'Email Addresses:', space_after=0)
    _cell_para(sig.cell(0, 0), '{attorney_email}', space_after=0)
    _cell_para(sig.cell(0, 0), '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _cell_para(sig.cell(0, 0), 'Florida Bar No. {attorney_bar_no}', space_after=6)
    _cell_para(sig.cell(0, 0), '{attorney_firm}', space_after=0)
    _cell_para(sig.cell(0, 0), '{attorney_address}', space_after=6)
    _cell_para(sig.cell(0, 0), 'Telephone: {attorney_phone}', space_after=0)

    _clear_cell(sig.cell(0, 1))
    _cell_para(sig.cell(0, 1), 'Personal Representative:', bold=True, space_after=12)
    _cell_para(sig.cell(0, 1), '{#prs}', space_after=0)
    _cell_para(sig.cell(0, 1), '_______________________________________', space_after=0)
    _cell_para(sig.cell(0, 1), '{pr_name}', space_after=6)
    _cell_para(sig.cell(0, 1), '{pr_address}', space_after=12)
    _cell_para(sig.cell(0, 1), '{/prs}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-0740.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


# ---------------------------------------------------------------------------
# P3-0900  Inventory
# ---------------------------------------------------------------------------

def build_p3_0900():
    """Verified inventory of estate assets per F.S. §733.604. PR files this
    within 60 days of issuance of letters of administration."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name}')
    _ensure_pleading_numbering(doc)

    _add_probate_caption(doc)

    _add_para(doc, 'INVENTORY',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc,
        'The undersigned {pr_label}{^multiple_prs}{/multiple_prs}{#multiple_prs}s{/multiple_prs} of the estate of {decedent_full_name}, deceased, who died on {decedent_death_date}, submit{^multiple_prs}s{/multiple_prs} this inventory of all the property of the estate that has come into the hands, possession, control, or knowledge of {petitioner_poss}{^multiple_prs}{/multiple_prs} {pr_label}:',
        space_after=18)

    # ---- Real Estate sections ----
    _add_para(doc, 'REAL ESTATE IN FLORIDA — Exempt (Protected) Homestead:',
              bold=True, space_after=6)
    _real_estate_table(doc, 'homestead_real_estate', value_col=False)
    _add_para(doc, '', space_after=6)

    _add_para(doc, 'REAL ESTATE IN FLORIDA — Non-Exempt Homestead:',
              bold=True, space_after=6)
    _real_estate_table(doc, 'nonexempt_homestead_real_estate', value_col=True)
    _add_para(doc,
        '(Whether homestead property is exempt from the claims of creditors, is properly devised and is a probate asset may have to be determined by appropriate proceedings.)',
        italic=True, space_after=12)

    _add_para(doc, 'OTHER REAL ESTATE IN FLORIDA:',
              bold=True, space_after=6)
    _real_estate_table(doc, 'other_real_estate', value_col=True)
    _add_para(doc,
        'Total Real Estate in Florida (except Exempt Homestead): ${total_real_estate}',
        bold=True, space_after=12)

    _add_para(doc, 'PERSONAL PROPERTY WHEREVER LOCATED:',
              bold=True, space_after=6)
    _real_estate_table(doc, 'personal_property', value_col=True)
    _add_para(doc,
        'Total Personal Property — Wherever Located: ${total_personal_property}',
        bold=True, space_after=12)

    _add_para(doc,
        'TOTAL OF ALL PERSONAL PROPERTY AND FLORIDA REAL ESTATE (except Exempt Homestead): ${total_estate}',
        bold=True, space_after=18)

    _add_para(doc,
        'All real estate located outside the State of Florida owned by the decedent of which the {pr_label} {pr_verb_is} aware{#has_out_of_state_property}, if any, is described on a schedule attached hereto{/has_out_of_state_property}{^has_out_of_state_property}: None{/has_out_of_state_property}.',
        space_after=12)

    _add_para(doc,
        'NOTICE: Each residuary beneficiary in a testate estate or heir in an intestate estate has the right to request a written explanation of how the inventory value of any asset was determined, including whether the {pr_label} obtained an independent appraisal for that asset and, if so, a copy of the appraisal. Any other beneficiary may request this information regarding all assets distributed to or proposed to be distributed to that beneficiary.',
        italic=True, space_after=18)

    indent = Inches(0.5)
    _add_para(doc,
        'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true to the best of my knowledge and belief.',
        first_indent=indent, space_after=18)

    _add_broward_ai_certification(doc, 'Inventory')
    _add_miami_dade_ai_certification(doc, 'Inventory')

    # PR signature block (no petitioner — inventory is signed by PR only).
    _add_para(doc, 'Signed on this _____ day of __________, 20___.',
              first_indent=indent, space_after=24)
    _add_para(doc, '{#prs}', space_after=0)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{pr_name}, {pr_label_title}', space_after=24)
    _add_para(doc, '{/prs}', space_after=0)

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for {pr_label_title}', space_after=12)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=0)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone: {attorney_phone}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'P3-0900.docx')
    doc.save(out_path)
    _inject_numbering_part(out_path)
    print(f'Wrote {out_path}')


def _real_estate_table(doc, group_name, *, value_col=True):
    """Bordered table for real-estate / personal-property items with
    repeating row. group_name is the docxtemplater array name (e.g.,
    'homestead_real_estate', 'personal_property')."""
    if value_col:
        tbl = _table_with_borders(doc, rows=2, cols=2, col_widths_in=[4.5, 2.0])
        for cell in tbl.rows[0].cells:
            _clear_cell(cell)
        _cell_para(tbl.cell(0, 0), 'Description', bold=True, space_after=0)
        _cell_para(tbl.cell(0, 1), 'Estimated Fair Market Value', bold=True, space_after=0)
        for cell in tbl.rows[1].cells:
            _clear_cell(cell)
        _cell_para(tbl.cell(1, 0), '{#' + group_name + '}{item_description}', space_after=0)
        _cell_para(tbl.cell(1, 1), '{item_value}{/' + group_name + '}', space_after=0)
    else:
        tbl = _table_with_borders(doc, rows=2, cols=1, col_widths_in=[6.5])
        for cell in tbl.rows[0].cells:
            _clear_cell(cell)
        _cell_para(tbl.cell(0, 0), 'Description', bold=True, space_after=0)
        for cell in tbl.rows[1].cells:
            _clear_cell(cell)
        _cell_para(tbl.cell(1, 0), '{#' + group_name + '}{item_description}{/' + group_name + '}', space_after=0)
    return tbl


# ---------------------------------------------------------------------------
# BW-0060  Affidavit of Heirs (Broward County local)
# ---------------------------------------------------------------------------

def build_bw_0060():
    """Broward 17th Circuit local form. Required for intestate estates per
    Broward Local Procedures (Oct 2023). Different caption from FLSSI —
    'IN THE SEVENTEENTH JUDICIAL CIRCUIT' instead of 'IN THE CIRCUIT COURT
    FOR {county} COUNTY'."""
    doc = Document()
    _apply_page_setup(doc)
    _apply_running_header(doc, 'Estate of {decedent_name} — Affidavit of Heirs')

    # ---- Broward-specific caption ----
    _add_para(doc, 'IN THE SEVENTEENTH JUDICIAL CIRCUIT, IN AND FOR',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'BROWARD COUNTY, FLORIDA',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, 'PROBATE DIVISION',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=12)

    cap = _borderless_table(doc, rows=1, cols=2, col_widths_in=[3.5, 3.0])
    _clear_cell(cap.cell(0, 0))
    _cell_para(cap.cell(0, 0), 'In Re: Estate of {decedent_full_name},', bold=True, space_after=0)
    _cell_para(cap.cell(0, 0), 'Deceased.', bold=True, space_after=0)
    _clear_cell(cap.cell(0, 1))
    _cell_para(cap.cell(0, 1), 'Case No.: {file_no}', bold=True, space_after=0)
    _cell_para(cap.cell(0, 1), 'Judge: {judge_name}', bold=True, space_after=0)
    _add_para(doc, '', space_after=18)

    _add_para(doc, 'AFFIDAVIT OF HEIRS',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=18)

    _add_para(doc,
        'For purposes of this affidavit, you must list ALL RELATIVES of the Decedent, including yourself, if applicable. Please include even the names of relatives who were deceased at the time of the Decedent’s death, indicating that they are deceased and specifying the date of death. If the Decedent never had a relative within a particular category (i.e. the decedent was the only child, and therefore had no siblings), please indicate "None" in that category. If the Decedent’s relatives in a particular category are unknown please specify "Unknown." When applicable, please indicate if the relationship is that of a half-relative (i.e. half-brother or half-sister).',
        italic=True, space_after=18)

    indent = Inches(0.5)

    _add_para(doc,
        '1. The undersigned, {affiant_name}, '
        '{#aoh_has_interest}has an interest in this estate{/aoh_has_interest}'
        '{^aoh_has_interest}does not have an interest in this estate{/aoh_has_interest}.',
        first_indent=indent, space_after=6)

    _add_para(doc,
        'I am '
        '{#aoh_is_related}related to the Decedent as follows: {aoh_relationship}{/aoh_is_related}'
        '{^aoh_is_related}not related to the Decedent{/aoh_is_related}.',
        first_indent=indent, space_after=6)

    _add_para(doc,
        'I have known the Decedent for {aoh_years_known} years.',
        first_indent=indent, space_after=18)

    # The categorical relative sections (2a–9). Each is a labeled
    # textarea-rendered free-text section; the user fills in names/addresses
    # in the questionnaire and the result lands in the .docx as a block.
    sections = [
        ('2a. Spouse of the Decedent.',
         '(Provide the name, age, and address. If the spouse is deceased, provide the name and date of death.)',
         'aoh_spouse_info'),
        ('2b. Decedent’s former spouse(s) (due to death or divorce).',
         '(Provide the name, age, and address. If the former spouse is deceased, provide the name and date of death. If Decedent and former spouse were divorced, provide the name of former spouse and date of divorce.)',
         'aoh_former_spouses'),
        ('3a. Children of the Decedent, or descendants of deceased children.',
         '(Provide the name, age, and address. If any of the children are deceased, provide the name and date of death. In addition, indicate if Decedent has any grandchildren from the predeceased children, include the grandchild’s name here and provide further information at 4.)',
         'aoh_children'),
        ('3b. Non-biological children.',
         '(If any of the children are not biologically related to both the Decedent and Decedent’s spouse at the time of Decedent’s death, provide that child’s name here and the name of that particular child’s other biological parent. If the surviving spouse has children who are not the children of the Decedent, provide their names.)',
         'aoh_non_biological_children'),
        ('4. Grandchildren of the Decedent.',
         '(Provide the name, age, and address. If the grandchild is deceased, indicate the name and date of death.)',
         'aoh_grandchildren'),
        ('5. Parents of the Decedent.',
         '(Provide the name, age, and address. If the parents are deceased, indicate the name and date of death.)',
         'aoh_parents'),
        ('6. Brothers and sisters of the Decedent, or descendants of deceased brothers or sisters.',
         '(You must specify if the relationship is that of a half-relative, i.e., half-brother or half-sister. You must provide the name, age, and current address of the Decedent’s brothers or sisters or half-siblings. If any of the brothers or sisters or half-siblings are deceased, you must provide the name and date of death. In addition, you must list the children of any predeceased brothers or sisters or half-siblings, if any, along with their current addresses.)',
         'aoh_siblings'),
        ('7. Nephews and nieces of the Decedent.',
         '(Provide the name, age, and address. If the nephew or niece is deceased, indicate the name and date of death.)',
         'aoh_nephews_nieces'),
        ('8. Grandparents of the Decedent.',
         '(Provide the name, age, and address. If the grandparents are deceased, indicate the name and date of death.)',
         'aoh_grandparents'),
        ('9. Other relatives.',
         '(If there are any relatives who have survived the Decedent and are not listed in the categories specified above, provide the name, relationship to the Decedent, age, and address (i.e., aunts, uncles, cousins, if applicable). Attach additional pages if necessary.)',
         'aoh_other_relatives'),
    ]
    for label, instruct, tag in sections:
        _add_para(doc, label, bold=True, space_after=0)
        _add_para(doc, instruct, italic=True, space_after=6)
        _add_para(doc, '{' + tag + '}', first_indent=indent, space_after=12)

    _add_para(doc,
        'UNDER PENALTIES OF PERJURY, I DECLARE THAT I HAVE READ THE FOREGOING AFFIDAVIT OF HEIRS AND THE FACTS STATED HEREIN ARE TRUE AND COMPLETE TO THE BEST OF MY KNOWLEDGE AND BELIEF.',
        bold=True, space_after=18)

    _add_broward_ai_certification(doc, 'Affidavit of Heirs')

    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{affiant_name}, Affiant', space_after=12)
    _add_para(doc, '{affiant_address}', space_after=18)

    # Notary block (the form is sworn — Broward requires).
    _add_para(doc, 'STATE OF {notary_state}', space_after=0)
    _add_para(doc, 'COUNTY OF {notary_county}', space_after=12)
    _add_para(doc,
        'Sworn to (or affirmed) and subscribed before me by means of ☐ physical presence or ☐ online notarization, on _____ day of __________, 20___, by {affiant_name}, who is personally known to me or produced ____________________ as identification.',
        space_after=18)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Notary Public or Deputy Clerk', space_after=12)
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, 'Print, Type, or Stamp Commissioned Name', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'BW-0060.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    build_p3_petition()
    build_p3_oath()
    build_p3_order()
    build_p3_letters()
    build_p1_0900()
    build_p1_0400()
    build_p1_formal_notice()
    build_p1_proof_of_service_fn()
    build_p1_0530()
    build_p1_caveat()
    build_p2_petition()
    build_p2_order()
    build_p2_0355()
    # Phase 8c
    build_p1_0100()
    build_p1_0620()
    build_p1_notice_confidential()
    build_p3_curator_petition()
    build_p3_curator_order()
    build_p3_curator_oath()
    build_p3_curator_letters()
    build_p3_oath_witness()
    build_p3_proof_will()
    # Phase 8e
    build_p1_0800()
    # Phase 11 legacy rebuilds (item B)
    build_p3_0740()
    build_p3_0900()
    build_bw_0060()
    # Phase 11 discharge cluster (item E)
    build_p5_petition_discharge_full_waiver()
    build_p5_receipt()
    build_p5_report_dist()
    build_p5_order_discharge()
