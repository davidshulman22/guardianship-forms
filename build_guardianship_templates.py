#!/usr/bin/env python3
"""
Build clean single-column guardianship petition templates that match
Ginsberg Shulman's preferred format (based on Jill's Villareal edits).

Abandons the 2-column FLSSI layout in favor of:
  - Single-column caption via a 2-column table (left: case title, right: division/file)
  - Left-aligned body (no justified text — prevents word-spreading when {field} has line breaks)
  - Real Word table for next-of-kin and property value
  - docxtemplater conditional blocks for has/has-no alternatives, preneed, professional guardian
  - Running header with matter name + Page X of Y on pages 2+
  - Clean signature lines for Petitioner and Attorney

Currently builds:
  G3-025  Petition for Appointment of Plenary Guardian of Property
"""

from docx import Document
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), 'templates')

FONT = 'Times New Roman'
BODY_SIZE = 12
SMALL_SIZE = 11


def _set_run(run, *, size=BODY_SIZE, bold=False, italic=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # East Asia + complex script font for Word compat
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)


def _add_para(doc_or_cell, text='', *, size=BODY_SIZE, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=None, left_indent=None,
              space_before=0, space_after=6, line_spacing=None, keep_with_next=False):
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    if first_indent is not None:
        pf.first_line_indent = first_indent
    if left_indent is not None:
        pf.left_indent = left_indent
    if line_spacing is not None:
        pf.line_spacing = line_spacing
    if keep_with_next:
        pf.keep_with_next = True
    if text:
        run = p.add_run(text)
        _set_run(run, size=size, bold=bold, italic=italic)
    return p


def _add_run(p, text, *, size=BODY_SIZE, bold=False, italic=False):
    run = p.add_run(text)
    _set_run(run, size=size, bold=bold, italic=italic)
    return run


def _borderless_table(doc, rows, cols, *, col_widths_in=None):
    """Create a table with no visible borders (for the caption layout)."""
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    # Nuke borders
    tblPr = tbl._element.find(qn('w:tblPr')) or tbl._element.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return tbl


def _table_with_borders(doc, rows, cols, *, col_widths_in=None):
    tbl = doc.add_table(rows=rows, cols=cols)
    tbl.autofit = False
    if col_widths_in:
        for i, w in enumerate(col_widths_in):
            for row in tbl.rows:
                row.cells[i].width = Inches(w)
    tblPr = tbl._element.find(qn('w:tblPr')) or tbl._element.get_or_add_tblPr()
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        tblBorders.append(b)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)
    return tbl


def _clear_cell(cell):
    """Remove the default empty paragraph a new cell ships with."""
    p = cell.paragraphs[0]
    if not p.runs and not p.text:
        p._element.getparent().remove(p._element)


def _cell_para(cell, text='', **kwargs):
    """Add a paragraph inside a table cell (works like _add_para)."""
    return _add_para(cell, text, **kwargs)


def _add_page_field(paragraph, field_code):
    """Insert a Word field (PAGE / NUMPAGES) into a paragraph."""
    run = paragraph.add_run()
    _set_run(run, size=SMALL_SIZE, italic=True)
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' ' + field_code + ' '
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._element.append(fldChar1)
    run._element.append(instrText)
    run._element.append(fldChar2)


def _set_different_first_page_header(section):
    """Enable 'different first page' so page 1 has no header."""
    sectPr = section._sectPr
    titlePg = sectPr.find(qn('w:titlePg'))
    if titlePg is None:
        titlePg = OxmlElement('w:titlePg')
        sectPr.append(titlePg)


# ---------------------------------------------------------------------------
# G3-025 Petition for Appointment of Plenary Guardian of Property
# ---------------------------------------------------------------------------

def build_g3_025():
    doc = Document()

    # Page setup: 1" margins, Times New Roman 12pt default
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    style = doc.styles['Normal']
    style.font.name = FONT
    style.font.size = Pt(BODY_SIZE)

    # Running header on pages 2+ only
    section = doc.sections[0]
    _set_different_first_page_header(section)
    header = section.header
    # Clear default empty paragraph and build our own
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)
    h_p1 = header.add_paragraph()
    h_p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_p1.paragraph_format.space_after = Pt(0)
    _add_run(h_p1, 'Guardianship of {aip_name}', size=SMALL_SIZE, italic=True)
    h_p2 = header.add_paragraph()
    h_p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    h_p2.paragraph_format.space_after = Pt(12)
    _add_run(h_p2, 'Page ', size=SMALL_SIZE, italic=True)
    _add_page_field(h_p2, 'PAGE')
    _add_run(h_p2, ' of ', size=SMALL_SIZE, italic=True)
    _add_page_field(h_p2, 'NUMPAGES')

    # ---- Caption ----
    _add_para(doc, 'IN THE CIRCUIT COURT FOR {county} COUNTY, FLORIDA',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    caption = _borderless_table(doc, rows=4, cols=2, col_widths_in=[3.5, 3.0])
    # Row 0
    _clear_cell(caption.cell(0, 0)); _cell_para(caption.cell(0, 0), 'IN RE: GUARDIANSHIP OF', space_after=0)
    _clear_cell(caption.cell(0, 1)); _cell_para(caption.cell(0, 1), 'PROBATE DIVISION', space_after=0)
    # Row 1 — AIP name
    _clear_cell(caption.cell(1, 0)); _cell_para(caption.cell(1, 0), '{aip_name_upper},', space_after=0)
    _clear_cell(caption.cell(1, 1)); _cell_para(caption.cell(1, 1), '', space_after=0)
    # Row 2
    _clear_cell(caption.cell(2, 0)); _cell_para(caption.cell(2, 0), 'An alleged incapacitated person', space_after=0)
    _clear_cell(caption.cell(2, 1)); _cell_para(caption.cell(2, 1), 'File No. {file_no}', space_after=0)
    # Row 3
    _clear_cell(caption.cell(3, 0)); _cell_para(caption.cell(3, 0), '', space_after=0)
    _clear_cell(caption.cell(3, 1)); _cell_para(caption.cell(3, 1), 'Division {division}', space_after=0)

    _add_para(doc, '', space_after=18)  # spacer

    # ---- Title ----
    _add_para(doc, 'PETITION FOR APPOINTMENT OF PLENARY GUARDIAN',
              align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, space_after=0)
    _add_para(doc, '(Incapacity - property)',
              align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)

    # ---- Body ----
    indent = Inches(0.5)

    _add_para(doc, 'Petitioner, {petitioner_name}, alleges:', space_after=12)

    # 1. Petitioner residence
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '1.\t')
    _add_run(p, 'Petitioner\u2019s residence is {petitioner_residence} and petitioner\u2019s post office address is {petitioner_address}.')

    # 2. AIP details
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '2.\t')
    _add_run(p, '{aip_name} (the AIP-Alleged Incapacitated Person) is an alleged incapacitated person who is {aip_age} years of age.  The residence of the Ward is {aip_residence}; and the post office address of the Ward is {aip_address}.')

    # 3. Nature of incapacity
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '3.\t')
    _add_run(p, 'The nature of the Ward\u2019s alleged incapacity is {ward_incapacity_nature}.')

    # 4. Alternatives — conditional
    # Build one paragraph with conditional branches
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '4.\t')
    _add_run(p, 'Petitioner {#has_alternatives}has knowledge, information and belief that there are alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. They are described as follows: {alternatives_description}. However, they are insufficient to meet the needs of the Ward because: {alternatives_insufficient_reason}. Thus, it is necessary that a plenary guardian be appointed to exercise all delegable rights of the Ward.{/has_alternatives}{^has_alternatives}has knowledge, information and belief that there are NO alternatives to the appointment of a guardian, such as trust agreements, powers of attorney, designations of health care surrogate, a guardian advocate under Florida Statutes section 744.3085, other type of guardianship or other advance directives. Thus, it is necessary that a plenary guardian be appointed to exercise all delegable rights of the Ward, as no other alternatives exist.{/has_alternatives}')

    # 5. Preneed — conditional
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '5.\t')
    _add_run(p, 'Petitioner {#has_preneed}has knowledge, information and belief that the alleged incapacitated person has a preneed guardian designation. The designated preneed guardian is {preneed_guardian_name}.{/has_preneed}{^has_preneed}has knowledge, information and belief that the alleged incapacitated person has NO preneed guardian designation{#preneed_reason}, {preneed_reason}{/preneed_reason}.{/has_preneed}')

    # 6. Next of kin intro
    p = _add_para(doc, '', first_indent=indent, space_after=6, keep_with_next=True)
    _add_run(p, '6.\t')
    _add_run(p, 'The names and addresses of the next of kin of the Ward are:')

    # Next of kin table
    kin = _table_with_borders(doc, rows=2, cols=3, col_widths_in=[2.3, 2.9, 1.3])
    # Header row
    for cell in kin.rows[0].cells:
        _clear_cell(cell)
    _cell_para(kin.cell(0, 0), 'NAME', bold=True, space_after=0)
    _cell_para(kin.cell(0, 1), 'ADDRESS', bold=True, space_after=0)
    _cell_para(kin.cell(0, 2), 'RELATIONSHIP', bold=True, space_after=0)
    # Repeating row — use paragraphLoop semantics with {#next_of_kin}/{/next_of_kin}
    for cell in kin.rows[1].cells:
        _clear_cell(cell)
    _cell_para(kin.cell(1, 0), '{#next_of_kin}{name}', space_after=0)
    _cell_para(kin.cell(1, 1), '{address}', space_after=0)
    _cell_para(kin.cell(1, 2), '{relationship}{/next_of_kin}', space_after=0)

    _add_para(doc, '', space_after=6)  # spacer

    # 7. Proposed guardian
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '7.\t')
    _add_run(p, 'The proposed guardian, {proposed_guardian_name}, whose residence is {proposed_guardian_residence}, and whose post office address is {proposed_guardian_address}, is sui juris and otherwise qualified under the laws of Florida to act as plenary guardian of the property of the Ward.{#is_professional_guardian} The proposed guardian is a professional guardian and has complied with the registration requirements of Florida Statutes section 744.2002.{/is_professional_guardian} The relationship and previous association of the proposed guardian to the Ward (including any activities designated in Florida Statutes section 744.446(3)) are {proposed_guardian_relationship}. The proposed guardian should be appointed because {appointment_reason}.')

    # 8. Property intro
    p = _add_para(doc, '', first_indent=indent, space_after=6, keep_with_next=True)
    _add_run(p, '8.\t')
    _add_run(p, 'The nature and value of the property subject to guardianship are as follows:')

    # Property table
    prop = _table_with_borders(doc, rows=2, cols=2, col_widths_in=[4.5, 2.0])
    for cell in prop.rows[0].cells:
        _clear_cell(cell)
    _cell_para(prop.cell(0, 0), 'Nature of Property', bold=True, space_after=0)
    _cell_para(prop.cell(0, 1), 'Value', bold=True, space_after=0)
    for cell in prop.rows[1].cells:
        _clear_cell(cell)
    _cell_para(prop.cell(1, 0), '{#property_items}{item_description}', space_after=0)
    _cell_para(prop.cell(1, 1), '{item_value}{/property_items}', space_after=0)

    _add_para(doc, '', space_after=6)

    # 9. Reasonable search
    p = _add_para(doc, '', first_indent=indent, space_after=12)
    _add_run(p, '9.\t')
    _add_run(p, 'Reasonable search has been made for all of the information required by Florida law and by the applicable Florida Probate Rules. Any such information that is not set forth in full above cannot be ascertained without delay that would adversely affect the Ward or the Ward\u2019s property.')

    # Request + oath
    _add_para(doc, 'Petitioner requests that {proposed_guardian_name} be appointed plenary guardian of the property of the Ward.',
              first_indent=indent, space_after=12)
    _add_para(doc, 'Under penalties of perjury, I declare that I have read the foregoing, and the facts alleged are true, to the best of my knowledge and belief.',
              first_indent=indent, space_after=18)

    # Signed on this ___ day
    _add_para(doc, 'Signed on this _____ day of {signing_month} {signing_year}.',
              first_indent=indent, space_after=24)

    # Petitioner signature line
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{petitioner_name}, Petitioner', space_after=24)

    # Attorney signature line
    _add_para(doc, '_______________________________________', space_after=0)
    _add_para(doc, '{attorney_name}, Attorney for Petitioner', space_after=18)

    # Attorney block
    _add_para(doc, 'Email Addresses:', space_after=0)
    _add_para(doc, '{attorney_email}', space_after=0)
    _add_para(doc, '{#attorney_email_secondary}{attorney_email_secondary}{/attorney_email_secondary}', space_after=0)
    _add_para(doc, 'Florida Bar No. {attorney_bar_no}', space_after=12)
    _add_para(doc, '{attorney_firm}', space_after=0)
    _add_para(doc, '{attorney_address}', space_after=12)
    _add_para(doc, 'Telephone {attorney_phone}', space_after=0)

    out_path = os.path.join(TEMPLATE_DIR, 'G3-025.docx')
    doc.save(out_path)
    print(f'Wrote {out_path}')


if __name__ == '__main__':
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    build_g3_025()
